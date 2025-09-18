# Copyright (C) 2025 CHAIB YAHIA
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU General Public License as published by
# the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU General Public License for more details.
#
# You should have received a copy of the GNU General Public License
# along with this program.  If not, see <https://www.gnu.org/licenses/>.

# ================== الجزء الأول: استيراد المكتبات الأساسية ==================
from flask import Flask, jsonify, request, render_template, send_file
import os
import json
import sys
import random
import re
from collections import defaultdict, Counter, deque
import io
from datetime import datetime
import math
import copy
import webbrowser
from waitress import serve
from threading import Timer
from ortools.sat.python import cp_model
import queue
from flask import stream_with_context, Response
from concurrent.futures import ThreadPoolExecutor
import signal
import threading
import time
import uuid
import sqlite3

# --- إضافة جديدة: لاستيراد مكتبة التعامل مع اكسل ---
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
import pandas as pd

from docx import Document
from docx.shared import Cm
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from collections import Counter
import uuid


# ================== الجزء الثاني: الإعدادات الأولية والدوال المساعدة ==================

log_queue = queue.Queue()
executor = ThreadPoolExecutor(max_workers=1)
STOP_EVENT = threading.Event()

def get_correct_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

# ================== بداية الكود الجديد لتحديد مسار البيانات بذكاء ==================
import shutil # تأكد من وجود هذه المكتبة مع بقية المكتبات المستوردة

def get_data_directory():
    """
    يحدد المسار الصحيح لمجلد 'data' بناءً على بيئة التشغيل.
    """
    # التحقق مما إذا كان البرنامج يعمل كملف تنفيذي (مُجمّد)
    if getattr(sys, 'frozen', False):
        executable_path = os.path.dirname(sys.executable)

        # التحقق مما إذا كان يعمل من مجلد نظام محمي مثل Program Files
        if 'program files' in executable_path.lower():
            # الحالة 1: البرنامج مثبت -> استخدم مجلد AppData
            app_name = "ExamGuardScheduler" # اسم فريد لبرنامجك
            app_data_path = os.path.join(os.getenv('APPDATA'), app_name)

            # عند أول تشغيل، قم بنسخ مجلد 'data' من مكان التثبيت إلى AppData
            install_source_data_dir = os.path.join(executable_path, 'data')
            if not os.path.exists(app_data_path) and os.path.exists(install_source_data_dir):
                shutil.copytree(install_source_data_dir, app_data_path)

            return app_data_path
        else:
            # الحالة 2: البرنامج يعمل كـ exe ولكن من مجلد عادي (مثل dist)
            return os.path.join(executable_path, 'data')
    else:
        # الحالة 3: البرنامج يعمل كسكربت .py في بيئة التطوير
        return os.path.join(os.path.dirname(__file__), 'data')

# استدعاء الدالة لتحديد المسارات النهائية
DATA_DIR = get_data_directory()
DB_PATH = os.path.join(DATA_DIR, 'database.db')
# ================== نهاية الكود الجديد ==================

# --- دوال التعامل مع قاعدة البيانات ---

def get_db_connection():
    os.makedirs(DATA_DIR, exist_ok=True)
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS professors (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT NOT NULL UNIQUE
    )''')
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS halls (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT NOT NULL UNIQUE,
        type TEXT NOT NULL
    )''')
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS levels (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT NOT NULL UNIQUE
    )''')
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS subjects (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT NOT NULL,
        level_id INTEGER NOT NULL,
        FOREIGN KEY (level_id) REFERENCES levels (id) ON DELETE CASCADE
    )''')
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS assignments (
        professor_id INTEGER NOT NULL,
        subject_id INTEGER NOT NULL,
        PRIMARY KEY (professor_id, subject_id),
        FOREIGN KEY (professor_id) REFERENCES professors (id) ON DELETE CASCADE,
        FOREIGN KEY (subject_id) REFERENCES subjects (id) ON DELETE CASCADE
    )''')
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS settings (
        key TEXT PRIMARY KEY,
        value TEXT NOT NULL
    )''')
    conn.commit()
    conn.close()
    print("Database initialized successfully.")

# --- دوال مساعدة أخرى ---

def clean_string_for_matching(text):
    if not isinstance(text, str): return text
    text = text.strip()
    return re.sub(r'\s+', ' ', text)

def parse_unique_id(unique_id, all_levels):
    sorted_levels = sorted(all_levels, key=len, reverse=True)
    for level in sorted_levels:
        if unique_id.endswith(f"({level})"):
            return unique_id[:-(len(level) + 2)].strip(), level
    last_paren_index = unique_id.rfind('(')
    if last_paren_index != -1 and unique_id.endswith(')'):
        return unique_id[:last_paren_index].strip(), unique_id[last_paren_index + 1:-1].strip()
    return unique_id, None

def apply_styles_to_cell(cell, font, alignment, border, fill=None):
    cell.font = font
    cell.alignment = alignment
    cell.border = border
    if fill: cell.fill = fill

# ================== دوال الخوارزمية (المنطق الأساسي) ==================
# الصق محتوى الدوال الأصلية هنا

def calculate_balanced_distribution(total_large, total_other, num_profs, w_large, w_other):
    if num_profs == 0: return []
    total_workload = (total_large * w_large) + (total_other * w_other)
    target_workload = total_workload / num_profs
    
    distribution = []
    base_large = total_large // num_profs
    remainder_large = total_large % num_profs
    
    for i in range(num_profs):
        large_count = base_large + 1 if i < remainder_large else base_large
        rem_workload = target_workload - (large_count * w_large)
        other_count = max(0, round(rem_workload / w_other))
        distribution.append({'large': large_count, 'other': other_count, 'total_workload': (large_count * w_large) + (other_count * w_other)})
    return distribution

def generate_balance_report(prof_stats, prof_targets):
    patterns = defaultdict(int)
    for stats in prof_stats.values():
        patterns[(stats['large'], stats['other'])] += 1
        
    target_patterns = defaultdict(int)
    if prof_targets:
        for target in prof_targets.values():
            target_patterns[(target['large'], target['other'])] += 1

    report_details = []
    all_keys = sorted(list(set(patterns.keys()) | set(target_patterns.keys())))
    total_deviation = 0
    
    for key in all_keys:
        actual = patterns.get(key, 0)
        target = target_patterns.get(key, 0)
        deviation = actual - target
        total_deviation += abs(deviation)
        report_details.append({'pattern': f"{key[0]} كبيرة + {key[1]} أخرى", 'target_count': target, 'actual_count': actual, 'deviation': deviation})

    balance_score = max(0, 100 - (total_deviation * 2))
    return {'details': report_details, 'balance_score': round(balance_score)}

def calculate_schedule_balance_score(schedule, all_professors, settings, num_professors):
    if not schedule or not all_professors: return 0.0
    large_hall_weight = float(settings.get('largeHallWeight', 3.0))
    other_hall_weight = float(settings.get('otherHallWeight', 1.0))
    guards_large_hall = int(settings.get('guardsLargeHall', 4))
    enable_custom_targets = settings.get('enableCustomTargets', False)
    custom_target_patterns = settings.get('customTargetPatterns', [])
    
    prof_stats = {prof: {'large': 0, 'other': 0} for prof in all_professors}
    all_exams = [exam for date_slots in schedule.values() for time_slots in date_slots.values() for exam in time_slots]
    for exam in all_exams:
        guards_copy = list(exam.get('guards', []))
        large_guards_needed = sum(guards_large_hall for h in exam.get('halls', []) if h.get('type') == 'كبيرة')
        large_hall_guards, other_hall_guards = guards_copy[:large_guards_needed], guards_copy[large_guards_needed:]
        for guard in large_hall_guards:
            if guard in prof_stats: prof_stats[guard]['large'] += 1
        for guard in other_hall_guards:
            if guard in prof_stats: prof_stats[guard]['other'] += 1
    prof_targets_map = {}
    if enable_custom_targets and custom_target_patterns:
        prof_targets_list = []
        for pattern in custom_target_patterns:
            for _ in range(pattern.get('count', 0)):
                prof_targets_list.append({'large': pattern.get('large', 0), 'other': pattern.get('other', 0)})
        shuffled_profs = list(all_professors); random.shuffle(shuffled_profs)
        prof_targets_map = {prof: prof_targets_list[i] for i, prof in enumerate(shuffled_profs) if i < len(prof_targets_list)}
    else:
        total_large_slots_final = sum(stats['large'] for stats in prof_stats.values())
        total_other_slots_final = sum(stats['other'] for stats in prof_stats.values())
        prof_targets_list = calculate_balanced_distribution(total_large_slots_final, total_other_slots_final, num_professors, large_hall_weight, other_hall_weight)
        if prof_targets_list:
            prof_targets_map = {prof: prof_targets_list[i % len(prof_targets_list)] for i, prof in enumerate(sorted(prof_stats.keys()))}
    if not prof_targets_map: return 0.0
    balance_report = generate_balance_report(prof_stats, prof_targets_map)
    return balance_report.get('balance_score', 0.0)

def is_assignment_valid(prof, exam, prof_assignments, prof_large_counts, settings, date_map):
    """
    دالة مركزية للتحقق مما إذا كان تعيين حارس لامتحان معين صالحاً أم لا.
    """
    # استخلاص الإعدادات من القاموس
    duty_patterns = settings.get('dutyPatterns', {})
    unavailable_days = settings.get('unavailableDays', {})
    max_shifts = int(settings.get('maxShifts', '0')) if settings.get('maxShifts', '0') != '0' else float('inf')
    max_large_hall_shifts = int(settings.get('maxLargeHallShifts', '2')) if settings.get('maxLargeHallShifts', '2') != '0' else float('inf')
    
    # 1. التحقق من التزامن (مشغول في نفس الوقت)
    if any(e['date'] == exam['date'] and e['time'] == exam['time'] for e in prof_assignments.get(prof, [])):
        return False

    # 2. التحقق من أيام الغياب
    if exam['date'] in unavailable_days.get(prof, []):
        return False

    # 3. التحقق من سقف الحصص الإجمالي
    if len(prof_assignments.get(prof, [])) >= max_shifts:
        return False

    # 4. التحقق من سقف حصص القاعة الكبيرة
    is_large_hall_exam = any(h['type'] == 'كبيرة' for h in exam['halls'])
    if is_large_hall_exam and prof_large_counts.get(prof, 0) >= max_large_hall_shifts:
        return False

    # 5. التحقق من نمط الحراسة
    # --- بداية: أسطر التعريفات التي كانت مفقودة ---
    prof_pattern = duty_patterns.get(prof, 'flexible_2_days')
    duties_dates = {d['date'] for d in prof_assignments.get(prof, [])}
    is_new_day = exam['date'] not in duties_dates
    num_duty_days = len(duties_dates)
    # --- نهاية: أسطر التعريفات التي كانت مفقودة ---

    if is_new_day:
        if (prof_pattern == 'one_day_only' and num_duty_days >= 1) or \
           (prof_pattern == 'flexible_2_days' and num_duty_days >= 2) or \
           (prof_pattern == 'flexible_3_days' and num_duty_days >= 3) or \
           (prof_pattern == 'consecutive_strict' and num_duty_days >= 2):
            return False
        elif prof_pattern == 'consecutive_strict' and num_duty_days == 1:
            idx1 = date_map.get(list(duties_dates)[0])
            idx2 = date_map.get(exam['date'])
            if idx1 is None or idx2 is None or abs(idx1 - idx2) != 1:
                return False

    # إذا نجح في كل الاختبارات، يكون التعيين صالحاً
    return True

def is_schedule_valid(schedule, settings, all_professors, duty_patterns, date_map):
    """
    النسخة المحدثة: مع إضافة التحقق من قيد "أزواج الأساتذة".
    """
    unavailable_days = settings.get('unavailableDays', {})
    max_shifts = int(settings.get('maxShifts', '0')) if settings.get('maxShifts', '0') != '0' else float('inf')
    max_large_hall_shifts = int(settings.get('maxLargeHallShifts', '2')) if settings.get('maxLargeHallShifts', '2') != '0' else float('inf')

    prof_assignments = defaultdict(list)
    prof_large_counts = defaultdict(int)
    
    all_exams = [exam for date_slots in schedule.values() for time_slots in date_slots.values() for exam in time_slots]

    # التحقق من القيود الأساسية (حصص متزامنة، غياب، عدد أقصى للحصص)
    for exam in all_exams:
        is_large_exam = any(h['type'] == 'كبيرة' for h in exam.get('halls', []))
        for guard in exam.get('guards', []):
            if guard == "**نقص**":
                return False 
            
            for other_exam in prof_assignments.get(guard, []):
                if other_exam['date'] == exam['date'] and other_exam['time'] == exam['time']:
                    return False 
                    
            if exam['date'] in unavailable_days.get(guard, []):
                return False 
            
            prof_assignments[guard].append(exam)
            if is_large_exam:
                prof_large_counts[guard] += 1

    for prof in all_professors:
        if len(prof_assignments[prof]) > max_shifts:
            return False
        if prof_large_counts[prof] > max_large_hall_shifts:
            return False
            
    # التحقق من نمط أيام الحراسة
    prof_assigned_slots = defaultdict(list)
    for exam in all_exams:
        for guard in exam.get('guards', []):
            if guard != "**نقص**":
                prof_assigned_slots[guard].append((exam['date'], exam['time']))

    for prof, pattern in duty_patterns.items():
        if not prof_assigned_slots.get(prof):
            continue
        duties_dates_indices = sorted(list({date_map.get(d_date) for d_date, d_time in prof_assigned_slots.get(prof, []) if date_map.get(d_date) is not None}))
        if not duties_dates_indices and prof_assigned_slots.get(prof):
             continue
        num_unique_duty_days = len(duties_dates_indices)
        if pattern == 'consecutive_strict':
            if num_unique_duty_days > 0 and (num_unique_duty_days != 2 or (duties_dates_indices[1] - duties_dates_indices[0] != 1)):
                return False
        elif pattern == 'one_day_only':
            if num_unique_duty_days > 1:
                return False
        elif pattern == 'flexible_2_days':
            if num_unique_duty_days > 0 and num_unique_duty_days != 2:
                return False
        elif pattern == 'flexible_3_days':
            if num_unique_duty_days > 0 and (num_unique_duty_days < 2 or num_unique_duty_days > 3):
                return False
    
    # --- بداية: التحقق من قيد أزواج الأساتذة ---
    professor_pairs = settings.get('professorPartnerships', []) # تم تغيير اسم المفتاح
    if professor_pairs:
        # أولاً، قم ببناء قاموس يحتوي على أيام الحراسة لكل أستاذ
        prof_duty_days = defaultdict(set)
        for guard, duties in prof_assigned_slots.items():
            for duty_date, _ in duties:
                prof_duty_days[guard].add(duty_date)

        # ثانياً، قم بالمرور على كل زوج وتحقق من تطابق أيام الحراسة
        for pair in professor_pairs:
            if len(pair) == 2:
                prof1_name, prof2_name = pair[0], pair[1]
                prof1_days = prof_duty_days.get(prof1_name, set())
                prof2_days = prof_duty_days.get(prof2_name, set())
                # إذا كانت مجموعة أيام الأستاذ الأول لا تساوي مجموعة أيام الثاني، فالقيد مكسور
                if prof1_days != prof2_days:
                    return False
    # --- نهاية: التحقق من قيد أزواج الأساتذة ---

    # إذا نجح الجدول في كل الاختبارات
    return True

# ================== الواجهة الرئيسية لتشغيل الخوارزمية (النسخة النهائية مع التحسينات) ==================

def analyze_failure_reason(exam, all_professors, prof_assignments, unavailable_days, max_shifts, max_large_hall_shifts, prof_large_hall_counts, duty_patterns, date_map, current_schedule):
    date, time = exam['date'], exam['time']
    is_large_hall_exam = any(h['type'] == 'كبيرة' for h in exam['halls'])
    
    reasons = Counter()
    for prof in all_professors:
        if any(prof in e['guards'] for e in current_schedule[date][time]):
            reasons['مشغول في نفس الوقت'] += 1
            continue
        if date in unavailable_days.get(prof, []):
            reasons['غير متاح في هذا اليوم'] += 1
            continue
        if len(prof_assignments.get(prof, [])) >= max_shifts:
            reasons['وصل للحد الأقصى من الحصص'] += 1
            continue
        if is_large_hall_exam and prof_large_hall_counts.get(prof, 0) >= max_large_hall_shifts:
            reasons['وصل للحد الأقصى من القاعات الكبيرة'] += 1
            continue
        # (Could add more detailed checks for duty patterns here if needed)

    if not reasons:
        return "سبب غير معروف، قد يكون متعلقاً بقيود نمط الحراسة المعقدة."

    most_common_reason, count = reasons.most_common(1)[0]
    return f"لم يتم العثور على حارس للامتحان '{exam['subject']}' في {date} {time}. السبب الأكثر شيوعًا: {count} أستاذًا كانوا '{most_common_reason}'."

# ===================================================================
# --- START: المرحلة 1.5 (النسخة الكاملة والنهائية) ---
# ===================================================================
def run_subject_optimization_phase(schedule, assignments, all_levels_list, subject_owners, settings, log_q, group_mappings, ideal_guard_days=None, stop_event=None):
    """
    النسخة الكاملة والنهائية:
    - تستقبل "أيام الحراسة المثالية" كـ "تغذية راجعة" لتوجيه التحسين.
    - تستخدم حلقة محاولات محددة لكل أستاذ لتجنب التعليق.
    - تهدف لتجميع المواد في يوم واحد أو يومين متتاليين.
    """
    log_q.put(">>> بدء المرحلة 1.5 (النسخة الكاملة): تحسين تجميع مواد الأساتذة...")
    
    # --- 1. الإعدادات ---
    passes = 3 # عدد جولات التحسين الكاملة
    max_improvement_attempts_per_prof = 25 # "ميزانية" المحاولات لكل أستاذ في كل جولة
    optimized_schedule = copy.deepcopy(schedule)
    
    # --- 2. هياكل بيانات مساعدة ---
    sorted_dates = sorted(optimized_schedule.keys())
    date_map = {date: i for i, date in enumerate(sorted_dates)}

    # --- 3. دالة لحساب التكلفة ---
    def calculate_scatter_cost(current_schedule, prof_to_exams_map):
        cost = 0
        for prof, exams in prof_to_exams_map.items():
            cost += len({e['date'] for e in exams})
        return cost

    # --- 4. حلقة التحسين الرئيسية (عدة جولات) ---
    for p in range(passes):
        if stop_event and stop_event.is_set():
            log_q.put("... [المرحلة 1.5] تم الإيقاف بواسطة المستخدم.")
            break
        # في بداية كل جولة، يتم بناء خريطة الأساتذة والمواد من جديد
        prof_to_exams = defaultdict(list)
        for date, time_slots in optimized_schedule.items():
            for time, exams in time_slots.items():
                for exam in exams:
                    owner = exam.get('professor')
                    if owner and owner != "غير محدد": prof_to_exams[owner].append(exam)
        
        if p == 0:
            initial_cost = calculate_scatter_cost(optimized_schedule, prof_to_exams)
            log_q.put(f"... التكلفة الأولية لتشتت الأيام: {initial_cost}")

        log_q.put(f"... [المرحلة 1.5] بدء جولة التحسين رقم {p + 1}/{passes}...")
        
        # ترتيب الأساتذة للبدء بالأكثر تشتتًا
        sorted_profs = sorted(prof_to_exams.keys(), key=lambda prof: len({e['date'] for e in prof_to_exams[prof]}), reverse=True)

        for prof in sorted_profs:
            if stop_event and stop_event.is_set():
                break
            # --- 5. حلقة التحسين المحددة لكل أستاذ ---
            for improvement_attempt in range(max_improvement_attempts_per_prof):
                prof_exams = prof_to_exams[prof]
                exam_days = set(e['date'] for e in prof_exams)
                if len(exam_days) <= 1: break # تم تحقيق الهدف، ننتقل للأستاذ التالي

                # --- 6. تحديد "اليوم الهدف" بذكاء ---
                anchor_day = None
                prof_ideal_days = ideal_guard_days.get(prof) if ideal_guard_days else None
                
                # الأولوية: اختيار يوم هدف من أيام الحراسة المثالية (التغذية الراجعة)
                if prof_ideal_days:
                    days_in_common = exam_days.intersection(prof_ideal_days)
                    if days_in_common:
                        anchor_day = random.choice(list(days_in_common))
                
                # إذا لم يوجد، نختار اليوم الذي به أكبر عدد من المواد حاليًا
                if not anchor_day:
                    day_counts = Counter(e['date'] for e in prof_exams)
                    anchor_day = day_counts.most_common(1)[0][0]
                
                exams_to_move = [e for e in prof_exams if e['date'] != anchor_day]
                if not exams_to_move: break
                exam_to_move = random.choice(exams_to_move)

                # --- 7. البحث عن تبديل مناسب (للأهداف 1 و 2) ---
                target_day_for_swap = anchor_day
                partner_found = None
                
                # الهدف 1: محاولة النقل إلى اليوم الأساسي
                partners = optimized_schedule.get(target_day_for_swap, {}).get(exam_to_move['time'], [])
                for partner in partners:
                    if partner['level'] == exam_to_move['level']:
                        owner_c = partner.get('professor')
                        if owner_c and owner_c != "غير محدد":
                            current_days_c = {e['date'] for e in prof_to_exams.get(owner_c, [])}
                            new_days_c = (current_days_c - {target_day_for_swap}) | {exam_to_move['date']}
                            if len(new_days_c) > len(current_days_c): continue
                        partner_found = partner
                        break
                
                # الهدف 2: محاولة النقل ليوم مجاور
                if not partner_found:
                    # ... (منطق اختيار الأيام المجاورة)
                    anchor_day_index = date_map.get(anchor_day)
                    adjacent_dates = []
                    if anchor_day_index > 0: adjacent_dates.append(sorted_dates[anchor_day_index - 1])
                    if anchor_day_index < len(sorted_dates) - 1: adjacent_dates.append(sorted_dates[anchor_day_index + 1])
                    
                    for adj_date in adjacent_dates:
                        if adj_date in exam_days: continue
                        target_day_for_swap = adj_date
                        partners = optimized_schedule.get(target_day_for_swap, {}).get(exam_to_move['time'], [])
                        for partner in partners:
                            if partner['level'] == exam_to_move['level']:
                                # ... (نفس التحقق من عدم الإضرار بالأستاذ الآخر)
                                owner_c = partner.get('professor')
                                if owner_c and owner_c != "غير محدد":
                                    current_days_c = {e['date'] for e in prof_to_exams.get(owner_c, [])}
                                    new_days_c = (current_days_c - {target_day_for_swap}) | {exam_to_move['date']}
                                    if len(new_days_c) > len(current_days_c): continue
                                partner_found = partner
                                break
                        if partner_found: break
                
                # --- 8. تنفيذ التبديل الآمن ---
                if partner_found:
                    try:
                        list_b = optimized_schedule[exam_to_move['date']][exam_to_move['time']]
                        list_c = optimized_schedule[target_day_for_swap][partner_found['time']]
                        idx_b = list_b.index(exam_to_move)
                        idx_c = list_c.index(partner_found)
                        
                        list_b[idx_b], list_c[idx_c] = partner_found, exam_to_move
                        exam_to_move['date'], partner_found['date'] = target_day_for_swap, exam_to_move['date']
                    except (ValueError, KeyError):
                        pass

        if stop_event and stop_event.is_set():
            break
    final_cost = calculate_scatter_cost(optimized_schedule, prof_to_exams)
    log_q.put(f"✓ انتهاء المرحلة 1.5. التكلفة النهائية لتشتت الأيام: {final_cost}")
    
    return optimized_schedule
# ===================================================================
# --- END: المرحلة 1.5 ---
# ===================================================================

def run_post_processing_swaps(schedule, prof_assignments, prof_workload, prof_large_counts, settings, all_professors, date_map, swap_attempts, locked_guards=set(), stop_event=None, log_q=None):
    large_hall_weight = float(settings.get('largeHallWeight', 3.0))
    other_hall_weight = float(settings.get('otherHallWeight', 1.0))
    duty_patterns = settings.get('dutyPatterns', {})
    unavailable_days = settings.get('unavailableDays', {})
    max_shifts = int(settings.get('maxShifts', '0')) if settings.get('maxShifts', '0') != '0' else float('inf')
    max_large_hall_shifts = int(settings.get('maxLargeHallShifts', '2')) if settings.get('maxLargeHallShifts', '2') != '0' else float('inf')
    
    temp_schedule = copy.deepcopy(schedule)
    
    # بناء/تحديث القواميس المساعدة من الجدول الحالي لضمان دقتها
    temp_assignments = defaultdict(list)
    temp_large_counts = defaultdict(int)
    temp_workload = defaultdict(float)
    all_exams_flat = [exam for day in temp_schedule.values() for slot in day.values() for exam in slot]
    for exam in all_exams_flat:
        is_large = any(h['type'] == 'كبيرة' for h in exam.get('halls', []))
        duty_weight = large_hall_weight if is_large else other_hall_weight
        for g in exam.get('guards', []):
            if g != "**نقص**":
                temp_assignments[g].append(exam)
                temp_workload[g] += duty_weight
                if is_large:
                    temp_large_counts[g] += 1

    for _ in range(swap_attempts):
        if stop_event and stop_event.is_set():
            if i > 0: log_q.put(f"... [الصقل] تم الإيقاف بعد {i} محاولة تبديل.")
            break
        if not temp_workload or len(temp_workload) < 2: break
        
        most_burdened_prof = max(temp_workload, key=temp_workload.get)
        least_burdened_prof = min(temp_workload, key=temp_workload.get)

        if most_burdened_prof == least_burdened_prof or temp_workload[most_burdened_prof] <= temp_workload[least_burdened_prof]:
            break 
            
        swap_found = False
        
        possible_swaps = [
            exam for exam in temp_assignments.get(most_burdened_prof, [])
            if (exam.get('uuid'), most_burdened_prof) not in locked_guards
        ]
        random.shuffle(possible_swaps)

        for exam in possible_swaps:
            date, time = exam['date'], exam['time']
            is_large_hall_exam = any(h['type'] == 'كبيرة' for h in exam.get('halls', []))
            # تجميع الإعدادات في قاموس واحد
            settings_for_validation = {
                'dutyPatterns': duty_patterns,
                'unavailableDays': unavailable_days,
                'maxShifts': max_shifts,
                'maxLargeHallShifts': max_large_hall_shifts
            }

            # استدعاء واحد يحل محل كل الشروط السابقة
            # لاحظ أننا نتحقق من صلاحية التعيين للأستاذ الأقل عبئاً (least_burdened_prof)
            if is_assignment_valid(least_burdened_prof, exam, temp_assignments, temp_large_counts, settings_for_validation, date_map):
                # --- بداية التعديل الشامل والمقترح ---
                exam_in_schedule = next((e for e in temp_schedule[date][time] if e.get('uuid') == exam.get('uuid')), None)
                if not exam_in_schedule: continue

                # الخطوة 1: الحذف الآمن للحارس القديم من قائمة الحراس
                try:
                    # نجد أول ظهور للحارس ونحذفه. هذا أكثر أمانًا من .remove()
                    guard_index_to_remove = exam_in_schedule['guards'].index(most_burdened_prof)
                    del exam_in_schedule['guards'][guard_index_to_remove]
                    exam_in_schedule['guards'].append(least_burdened_prof)
                except ValueError:
                    # إذا لم يتم العثور على الحارس، نتجاهل هذا التبديل وننتقل للتالي
                    # هذا يعالج الخطأ مباشرة "list.remove(x): x not in list"
                    continue

                # الخطوة 2: تحديث إحصائيات عبء العمل
                duty_weight = large_hall_weight if is_large_hall_exam else other_hall_weight
                temp_workload[most_burdened_prof] -= duty_weight
                temp_workload[least_burdened_prof] += duty_weight
                if is_large_hall_exam:
                    temp_large_counts[most_burdened_prof] -= 1
                    temp_large_counts[least_burdened_prof] = temp_large_counts.get(least_burdened_prof, 0) + 1
                
                # الخطوة 3: الحذف الآمن للامتحان من قائمة مهام الحارس القديم
                exam_to_remove_index = -1
                for i, assigned_exam in enumerate(temp_assignments[most_burdened_prof]):
                    if assigned_exam.get('uuid') == exam.get('uuid'):
                        exam_to_remove_index = i
                        break
                
                if exam_to_remove_index != -1:
                    del temp_assignments[most_burdened_prof][exam_to_remove_index]
                else:
                    # هذا لا يفترض أن يحدث، ولكنه إجراء احترازي لمنع أخطاء مستقبلية
                    continue 

                # الخطوة 4: إضافة الامتحان لقائمة مهام الحارس الجديد
                temp_assignments[least_burdened_prof].append(exam)
                
                swap_found = True
                break
                # --- نهاية التعديل الشامل والمقترح ---
        
        if not swap_found:
            break 
            
    return temp_schedule, temp_assignments, temp_workload, temp_large_counts

# في ملف app.py، قم باستبدال هذه الدالة بالكامل

def run_simulated_annealing(schedule, prof_assignments, prof_workload, prof_large_counts, settings, all_professors, date_map, duty_patterns, annealing_iterations, annealing_temp, annealing_cooling):
    """
    النسخة النهائية والمحسنة.
    تستهدف الأنماط المخصصة إذا كانت مفعلة، وإلا تقوم بالموازنة العامة.
    """
    # --- استخلاص الإعدادات ---
    large_hall_weight = float(settings.get('largeHallWeight', 3.0))
    other_hall_weight = float(settings.get('otherHallWeight', 1.0))
    enable_custom_targets = settings.get('enableCustomTargets', False)
    custom_target_patterns = settings.get('customTargetPatterns', [])
    guards_large_hall = int(settings.get('guardsLargeHall', 4))

    # --- تحديد دالة الطاقة التي سنستخدمها ---
    energy_function = None

    if enable_custom_targets and custom_target_patterns:
        # الهدف: تقليل الانحراف عن الأنماط المستهدفة
        target_counts = Counter((p['large'], p['other']) for p in custom_target_patterns for _ in range(p.get('count', 0)))
        all_target_patterns = set(target_counts.keys())

        def calculate_pattern_deviation_energy(sch):
            prof_stats = {prof: {'large': 0, 'other': 0} for prof in all_professors}
            all_exams = [exam for date_slots in sch.values() for time_slots in date_slots.values() for exam in time_slots]
            for exam in all_exams:
                guards_copy = [g for g in exam.get('guards', []) if g != "**نقص**"]
                large_guards_needed = sum(guards_large_hall for h in exam.get('halls', []) if h.get('type') == 'كبيرة')
                large_hall_guards = guards_copy[:large_guards_needed]
                other_hall_guards = guards_copy[large_guards_needed:]
                for guard in large_hall_guards:
                    if guard in prof_stats: prof_stats[guard]['large'] += 1
                for guard in other_hall_guards:
                    if guard in prof_stats: prof_stats[guard]['other'] += 1
            
            actual_counts = Counter((s['large'], s['other']) for s in prof_stats.values())
            
            total_deviation = 0
            all_current_patterns = set(actual_counts.keys()) | all_target_patterns
            for pattern in all_current_patterns:
                total_deviation += abs(actual_counts.get(pattern, 0) - target_counts.get(pattern, 0))
            
            return total_deviation
        
        energy_function = calculate_pattern_deviation_energy

    else:
        # الهدف الافتراضي: الموازنة العامة
        def calculate_workload_energy(sch):
            workload_dict = defaultdict(float)
            all_exams = [exam for date_slots in sch.values() for time_slots in date_slots.values() for exam in time_slots]
            for exam in all_exams:
                 is_large_exam = any(h['type'] == 'كبيرة' for h in exam.get('halls', []))
                 duty_weight = large_hall_weight if is_large_exam else other_hall_weight
                 for guard in exam.get('guards', []):
                     if guard != "**نقص**" and guard in all_professors:
                         workload_dict[guard] += duty_weight
            
            if not workload_dict: return 0.0
            workloads = list(workload_dict.values())
            return max(workloads) - min(workloads)

        energy_function = calculate_workload_energy

    # --- إعدادات الحالة الأولية ---
    current_schedule = copy.deepcopy(schedule)
    current_energy = energy_function(current_schedule)
    
    best_schedule = copy.deepcopy(current_schedule)
    best_energy = current_energy

    temp = annealing_temp
    
    # --- حلقة التلدين المحاكي الرئيسية ---
    for i in range(annealing_iterations):
        if temp < 0.01: break

        all_duties = [{'exam': ex, 'guard': g, 'idx': g_idx} 
                      for ex_list in current_schedule.values() 
                      for exs in ex_list.values() for ex in exs 
                      for g_idx, g in enumerate(ex.get('guards', [])) if g != "**نقص**"]

        if not all_duties: break

        duty_to_change = random.choice(all_duties)
        exam = duty_to_change['exam']
        prof1 = duty_to_change['guard']
        guard_idx = duty_to_change['idx']

        possible_new_profs = [p for p in all_professors if p != prof1]
        if not possible_new_profs: continue
        prof2 = random.choice(possible_new_profs)
        
        schedule_copy = copy.deepcopy(current_schedule)
        exam_in_copy = None
        for date_key, time_slots in schedule_copy.items():
            if date_key == exam['date']:
                for time_key, exams in time_slots.items():
                    if time_key == exam['time']:
                        for ex_item in exams:
                            if ex_item['subject'] == exam['subject'] and ex_item['level'] == exam['level']:
                                exam_in_copy = ex_item
                                break
                if exam_in_copy: break
        
        if not exam_in_copy: continue
        
        exam_in_copy['guards'][guard_idx] = prof2
        
        if not is_schedule_valid(schedule_copy, settings, all_professors, duty_patterns, date_map):
            continue

        new_energy = energy_function(schedule_copy)
        
        if new_energy < current_energy or random.random() < math.exp((current_energy - new_energy) / temp):
            current_schedule = schedule_copy
            current_energy = new_energy
            
            if current_energy < best_energy:
                best_energy = current_energy
                best_schedule = copy.deepcopy(current_schedule)
        
        temp *= annealing_cooling

    # إعادة بناء بقية بيانات الحل النهائي
    final_assignments = defaultdict(list)
    final_large_counts = defaultdict(int)
    final_workload = defaultdict(float)
    for ex_list in best_schedule.values():
        for exs in ex_list.values():
            for ex in exs:
                is_large = any(h['type'] == 'كبيرة' for h in ex['halls'])
                duty_weight = large_hall_weight if is_large else other_hall_weight
                for g in ex.get('guards', []):
                    if g != "**نقص**":
                        final_assignments[g].append(ex)
                        final_workload[g] += duty_weight
                        if is_large:
                            final_large_counts[g] += 1
    
    return best_schedule, final_assignments, final_workload, final_large_counts

# ===================================================================
# --- START: النسخة النهائية من Tabu Search (مكتملة ومستقرة) ---
# ===================================================================
def run_tabu_search(initial_schedule, settings, all_professors, duty_patterns, date_map, log_q, locked_guards=set(), stop_event=None):
    """
    النسخة النهائية والمصححة بالكامل:
    - كل الحركات (إصلاح وموازنة) تتحقق من صلاحية القيود قبل تنفيذها.
    - تستخدم المنطق الصحيح للبحث المحظور للهروب من الحلول المحلية.
    - عالية الكفاءة بفضل بناء الحالة مرة واحدة لكل دورة.
    """
    log_q.put(">>> تشغيل البحث المحظور (النسخة النهائية الكاملة)...")

    # --- 1. استخلاص الإعدادات ---
    max_iterations = int(settings.get('tabuIterations', 200))
    tabu_tenure = int(settings.get('tabuTenure', 20))
    neighborhood_size = int(settings.get('tabuNeighborhoodSize', 100))

    # --- 2. الإعدادات الأولية للبحث ---
    current_solution = copy.deepcopy(initial_schedule)
    best_solution = copy.deepcopy(current_solution)

    current_cost = calculate_cost(current_solution, settings, all_professors, duty_patterns, date_map)
    best_cost = current_cost
    log_q.put(f"... [Tabu Search] التكلفة الأولية = {format_cost_tuple(best_cost)}")
    tabu_list = deque(maxlen=tabu_tenure)

    # --- 3. حلقة البحث الرئيسية ---
    for i in range(max_iterations):
        if stop_event and stop_event.is_set(): break
        if settings.get('should_stop_event', threading.Event()).is_set(): break
        percent_complete = int(((i + 1) / max_iterations) * 100)
        log_q.put(f"PROGRESS:{percent_complete}")

        # ✅ تصحيح: تم تغيير float('inf') إلى tuple لتجنب خطأ المقارنة
        best_neighbor_in_iteration, best_neighbor_cost_in_iteration, best_move_in_iteration = None, (float('inf'), float('inf'), float('inf'), float('inf')), None

        # --- بناء الحالة الحالية مرة واحدة لكل دورة (لتحقيق أقصى كفاءة) ---
        current_assignments = defaultdict(list); current_large_counts = defaultdict(int)
        all_exams_in_current = [e for d in current_solution.values() for s in d.values() for e in s]
        for e in all_exams_in_current:
            is_large = any(h['type'] == 'كبيرة' for h in e.get('halls', []))
            for g in e.get('guards', []):
                if g != "**نقص**":
                    current_assignments[g].append(e)
                    if is_large: current_large_counts[g] += 1

        # --- استراتيجية الحركة الديناميكية ---
        repair_probability = 0.8 if current_cost[0] > 0 or current_cost[1] > 0 else 0.1

        # --- 4. استكشاف الجوار ---
        for _ in range(neighborhood_size):
            neighbor = None
            move = None

            if random.random() < repair_probability:
                # --- حركة الإصلاح ---
                shortage_slots = [(exam, g_idx) for exam in all_exams_in_current for g_idx, g in enumerate(exam.get('guards',[])) if g == "**نقص**"]
                if not shortage_slots: continue
                exam_to_repair, guard_idx = random.choice(shortage_slots)

                shuffled_profs = list(all_professors); random.shuffle(shuffled_profs)
                prof_to_add = next((p for p in shuffled_profs if is_assignment_valid(p, exam_to_repair, current_assignments, current_large_counts, settings, date_map)), None)

                if prof_to_add:
                    neighbor = copy.deepcopy(current_solution)
                    exam_in_neighbor = next(e for e in neighbor[exam_to_repair['date']][exam_to_repair['time']] if e.get('uuid') == exam_to_repair.get('uuid'))
                    exam_in_neighbor['guards'][guard_idx] = prof_to_add
                    move = (exam_in_neighbor.get('uuid'), guard_idx)
            else:
                # --- حركة الموازنة ---
                all_duties = [(exam, g, d_idx) for exam in all_exams_in_current for d_idx, g in enumerate(exam.get('guards',[])) if g != "**نقص**" and (exam.get('uuid'), g) not in locked_guards]
                if not all_duties: continue

                exam_to_change, prof1, guard_idx = random.choice(all_duties)
                possible_profs = [p for p in all_professors if p != prof1 and p not in exam_to_change.get('guards', [])]
                if not possible_profs: continue
                prof2 = random.choice(possible_profs)

                if not is_assignment_valid(prof2, exam_to_change, current_assignments, current_large_counts, settings, date_map):
                    continue

                neighbor = copy.deepcopy(current_solution)
                exam_in_neighbor = next(e for e in neighbor[exam_to_change['date']][exam_to_change['time']] if e.get('uuid') == exam_to_change.get('uuid'))
                exam_in_neighbor['guards'][guard_idx] = prof2
                move = (exam_in_neighbor.get('uuid'), guard_idx)

            if not neighbor: continue

            # --- 5. تقييم الجار وتطبيق منطق المحظورات ---
            neighbor_cost = calculate_cost(neighbor, settings, all_professors, duty_patterns, date_map)

            is_tabu = move in tabu_list
            if is_tabu:
                if neighbor_cost < best_cost: # Aspiration
                    if neighbor_cost < best_neighbor_cost_in_iteration:
                         best_neighbor_in_iteration, best_neighbor_cost_in_iteration, best_move_in_iteration = neighbor, neighbor_cost, move
            else:
                if neighbor_cost < best_neighbor_cost_in_iteration:
                    best_neighbor_in_iteration, best_neighbor_cost_in_iteration, best_move_in_iteration = neighbor, neighbor_cost, move

        # --- 6. تحديث الحالة للدورة القادمة ---
        if not best_neighbor_in_iteration:
            continue

        current_solution = best_neighbor_in_iteration
        current_cost = best_neighbor_cost_in_iteration # Use the already calculated cost
        tabu_list.append(best_move_in_iteration)

        if current_cost < best_cost:
            best_cost, best_solution = current_cost, copy.deepcopy(current_solution)
            log_q.put(f"... [Tabu] دورة {i+1}: حل أفضل بتكلفة = {format_cost_tuple(best_cost)}")

            if best_cost[0] == 0 and best_cost[1] == 0:
                log_q.put("... [Tabu] تم العثور على حل صالح ومكتمل.")
                if best_cost[2] == 0 and best_cost[3] == 0:
                    log_q.put("... [Tabu] الحل مثالي، إنهاء البحث.")
                    break

    # --- 7. إرجاع أفضل حل تم العثور عليه ---
    final_cost = calculate_cost(best_solution, settings, all_professors, duty_patterns, date_map)
    log_q.put(f"✓ البحث المحظور انتهى بأفضل تكلفة: {format_cost_tuple(final_cost)}")
    return best_solution, None, None, None
# ===================================================================
# --- END: النسخة النهائية من Tabu Search ---
# ===================================================================




# ===================================================================
# --- START: COST FUNCTION V6 (DEVIATION > SOFT CONSTRAINTS) ---
# ===================================================================
def calculate_cost(schedule, settings, all_professors, duty_patterns, date_map):
    """
    (النسخة V6) تعطي الأولوية للانحراف عن التوزيع على القيود المرنة.
    """
    # حساب كل المكونات بشكل منفصل
    
    # 1. نقص الحراسة
    all_exams_flat = [exam for day in schedule.values() for slot in day.values() for exam in slot]
    shortage_component = sum(e.get('guards', []).count("**نقص**") for e in (exam for day in schedule.values() for slot in day.values() for exam in slot))

    # 2. القيود الصارمة
    hard_constraint_component = 1 if not is_schedule_valid(schedule, settings, all_professors, duty_patterns, date_map) else 0

    # 3. القيود المرنة (الهيكلية)
    soft_constraint_component = 0
    prof_subject_days = defaultdict(set)
    prof_guards_days = defaultdict(set)
    for exam in all_exams_flat:
        # ... (نفس منطق حساب الأيام)
        owner = exam.get('professor', "غير محدد")
        if owner != "غير محدد":
            prof_subject_days[owner].add(exam['date'])
        for guard in exam.get('guards', []):
            if guard != "**نقص**":
                prof_guards_days[guard].add(exam['date'])

    for prof in all_professors:
        subject_days = prof_subject_days.get(prof, set())
        guard_days = prof_guards_days.get(prof, set())
        
        # نحسب فقط "الأيام الضائعة": أيام المواد التي لم يتم استغلالها للحراسة
        missed_opportunity_days = len(subject_days - guard_days)
        
        soft_constraint_component += missed_opportunity_days * 10
    
    for prof, days in prof_subject_days.items():
        if len(days) > 2:
            soft_constraint_component += (len(days) - 2) * 5

    # 4. الانحراف عن التوزيع
    deviation_component = 0.0
    # ... (نفس منطق حساب balance_cost)
    # ...
    large_hall_weight = float(settings.get('largeHallWeight', 3.0))
    prof_stats = {prof: {'large': 0, 'other': 0} for prof in all_professors}
    guards_large_hall = int(settings.get('guardsLargeHall', 4))
    for exam in all_exams_flat:
        guards_copy = [g for g in exam.get('guards', []) if g != "**نقص**"]
        large_guards_needed = sum(guards_large_hall for h in exam.get('halls', []) if h.get('type') == 'كبيرة')
        for guard in guards_copy[:large_guards_needed]:
            if guard in prof_stats: prof_stats[guard]['large'] += 1
        for guard in guards_copy[large_guards_needed:]:
            if guard in prof_stats: prof_stats[guard]['other'] += 1
            
    enable_custom_targets = settings.get('enableCustomTargets', False)
    custom_target_patterns = settings.get('customTargetPatterns', [])

    if enable_custom_targets and custom_target_patterns:
        target_counts = Counter((p['large'], p['other']) for p in custom_target_patterns for _ in range(p.get('count', 0)))
        actual_counts = Counter((s['large'], s['other']) for s in prof_stats.values())
        total_deviation = sum(abs(actual_counts.get(p, 0) - target_counts.get(p, 0)) for p in set(target_counts.keys()) | set(actual_counts.keys()))
        deviation_component = total_deviation * 2.0
    else:
        prof_workload = {p: s['large'] * large_hall_weight + s['other'] for p, s in prof_stats.items()}
        if prof_workload:
            workload_values = list(prof_workload.values())
            deviation_component = max(workload_values) - min(workload_values) if workload_values else 0.0

    # إرجاع tuple يحتوي على كل المكونات بالترتيب
    return (shortage_component, hard_constraint_component, deviation_component, soft_constraint_component)
# ===================================================================
# --- END: COST FUNCTION V6 ---
# ===================================================================

def format_cost_tuple(cost_tuple):
    """(النسخة المحدثة) تنسيق تفاصيل التكلفة لطباعتها في السجل حسب الترتيب الجديد."""
    # ✅ التغيير هنا: قمنا بتغيير ترتيب المتغيرات لتطابق الـ tuple الجديد
    s, h, d, f = cost_tuple
    # ✅ التغيير هنا: قمنا بتغيير ترتيب المسميات في النص المطبوع
    return f"(نقص: {s}, قيود صارمة: {h}, انحراف: {d:.2f}, قيود مرنة: {f})"


# =====================================================================
# START: HYPER-HEURISTIC HELPER FUNCTIONS (PORTED FROM PROJECT 1)
# =====================================================================

def get_state_from_failures_dominant(failures, unplaced_count):
    """
    تحول قائمة الأخطاء إلى تمثيل "حالة" مبسط.
    تركز الحالة على المشكلة الأكثر إلحاحًا: النقص أولاً، ثم الأخطاء الصارمة.
    """
    if unplaced_count > 0:
        return "UNPLACED_ITEMS"

    hard_failures = [f for f in failures if f.get('penalty', 0) >= 100]
    if hard_failures:
        # يمكنك جعلها أكثر تفصيلاً بالبحث عن نوع الخطأ الصارم الأكثر تكراراً
        return "HARD_CONSTRAINT_VIOLATION"

    soft_failures = [f for f in failures if 0 < f.get('penalty', 0) < 100]
    if soft_failures:
        return "SOFT_CONSTRAINT_VIOLATION"
        
    return "OPTIMAL_OR_NEAR_OPTIMAL"

def calculate_reward_from_cost(old_cost_tuple, new_cost_tuple):
    """
    تحسب المكافأة بناءً على التحسن الهرمي بين حليّن باستخدام tuple التكلفة.
    (متوافقة مع نظام التكلفة في المشروع الثاني)
    """
    # old_cost = (shortage, hard, deviation, soft)
    # new_cost = (shortage, hard, deviation, soft)
    
    # 1. مقارنة عدد النقص (الأهم)
    if new_cost_tuple[0] < old_cost_tuple[0]: return 1000  # تحسن كبير
    if new_cost_tuple[0] > old_cost_tuple[0]: return -2000 # تدهور كارثي

    # 2. إذا كان النقص متساويًا، نقارن القيود الصارمة
    if new_cost_tuple[1] < old_cost_tuple[1]: return 500   # تحسن ممتاز
    if new_cost_tuple[1] > old_cost_tuple[1]: return -1000 # تدهور كبير

    # 3. إذا كان النقص والقيود الصارمة متساويين، نقارن الانحراف
    if new_cost_tuple[2] < old_cost_tuple[2]: return 100   # تحسن جيد
    if new_cost_tuple[2] > old_cost_tuple[2]: return -50   # تدهور طفيف

    # 4. إذا كان كل ما سبق متساوياً، نقارن القيود المرنة
    if new_cost_tuple[3] < old_cost_tuple[3]: return 20    # تحسن مقبول
    
    # 5. في حالة عدم وجود أي تغيير (ركود)
    return -5

# =====================================================================
# END: HYPER-HEURISTIC HELPER FUNCTIONS
# =====================================================================

# ===================================================================
# --- START: النسخة النهائية المدمجة من LNS ---
# ===================================================================
def run_large_neighborhood_search(
    initial_schedule, settings, all_professors, duty_patterns, date_map, log_q, locked_guards=set(), stop_event=None
):
    """
    النسخة النهائية والمحسّنة من LNS:
    - تجمع بين الإصلاح الشامل لكل حالات النقص (من الدالة الجديدة).
    - تستخدم نظام عبء العمل الموزون الدقيق (من الدالة الأصلية) لاختيار أفضل مرشح.
    - تحتفظ بنسبة التدمير الديناميكية.
    """
    log_q.put(">>> تشغيل LNS (النسخة النهائية المدمجة والمحسّنة)...")

    # --- 1. استخلاص الإعدادات (تبقى كما هي) ---
    iterations = int(settings.get('lnsIterations', 100))
    initial_destroy_fraction = float(settings.get('lnsDestroyFraction', 0.2))
    min_destroy_fraction = 0.05
    destroy_fraction_decay_rate = 0.995
    initial_temp = 10.0
    cooling_rate = 0.99
    
    settings_for_validation = {
        'dutyPatterns': duty_patterns,
        'unavailableDays': settings.get('unavailableDays', {}),
        'maxShifts': settings.get('maxShifts', '0'),
        'maxLargeHallShifts': settings.get('maxLargeHallShifts', '2')
    }
    large_hall_weight = float(settings.get('largeHallWeight', 3.0))
    other_hall_weight = float(settings.get('otherHallWeight', 1.0))

    # --- 2. الحل المبدئي وحساب التكلفة الأولية (تبقى كما هي) ---
    current_solution = copy.deepcopy(initial_schedule)
    best_solution_so_far = copy.deepcopy(current_solution)
    
    current_cost = calculate_cost(current_solution, settings, all_professors, duty_patterns, date_map)
    best_cost_so_far = current_cost
    log_q.put(f"... [LNS] التكلفة الأولية = {format_cost_tuple(current_cost)}")

    temp = initial_temp
    dynamic_destroy_fraction = initial_destroy_fraction

    # --- 3. حلقة LNS الرئيسية (تبقى كما هي) ---
    for i in range(iterations):
        if stop_event and stop_event.is_set(): break
        
            
        percent_complete = int(((i + 1) / iterations) * 100)
        log_q.put(f"PROGRESS:{percent_complete}")
        
        ruined_solution = copy.deepcopy(current_solution)

        # --- 4. مرحلة التدمير (تبقى كما هي) ---
        duties_to_destroy = []
        for day in ruined_solution.values():
            for slot in day.values():
                for exam in slot:
                    for g_idx, guard in enumerate(exam.get('guards', [])):
                        if guard != "**نقص**" and (exam.get('uuid'), guard) not in locked_guards:
                            duties_to_destroy.append({'exam': exam, 'guard_index': g_idx})

        random.shuffle(duties_to_destroy)
        num_to_destroy = int(len(duties_to_destroy) * dynamic_destroy_fraction)
        
        for j in range(min(num_to_destroy, len(duties_to_destroy))):
            duty_info = duties_to_destroy[j]
            duty_info['exam']['guards'][duty_info['guard_index']] = "**نقص**"

        # --- 5. مرحلة الإصلاح الذكي والمستهدف (النسخة المدمجة) ---
        all_exams_in_ruined = [exam for day in ruined_solution.values() for slot in day.values() for exam in slot]
        
        # (## تعديل مدمج ##): 1. حساب الحالة الأولية (العبء الموزون) مرة واحدة قبل حلقة الإصلاح
        prof_assignments = defaultdict(list)
        prof_large_counts = defaultdict(int)
        prof_workload = defaultdict(float)
        for exam in all_exams_in_ruined:
            is_large = any(h['type'] == 'كبيرة' for h in exam.get('halls', []))
            duty_weight = large_hall_weight if is_large else other_hall_weight
            for guard in exam.get('guards', []):
                if guard != "**نقص**":
                    prof_assignments[guard].append(exam)
                    prof_workload[guard] += duty_weight
                    if is_large:
                        prof_large_counts[guard] += 1
        
        # 2. تحديد كل خانات النقص (من الدالة الجديدة)
        shortage_slots = []
        for exam in all_exams_in_ruined:
            for idx, guard in enumerate(exam.get('guards', [])):
                if guard == "**نقص**":
                    shortage_slots.append({'exam': exam, 'index_to_fill': idx})
        random.shuffle(shortage_slots)

        # 3. حلقة الإصلاح التي تستهدف كل خانات النقص مع استخدام العبء الموزون
        for repair_info in shortage_slots:
            exam_to_repair = repair_info['exam']
            
            # (## تعديل مدمج ##): حساب وزن المهمة المطلوب إصلاحها
            is_large_repair_exam = any(h['type'] == 'كبيرة' for h in exam_to_repair.get('halls',[]))
            repair_duty_weight = large_hall_weight if is_large_repair_exam else other_hall_weight
            
            # إيجاد أفضل مرشح صالح
            valid_candidates = []
            for prof in all_professors:
                if prof in exam_to_repair.get('guards', []): continue
                
                if is_assignment_valid(prof, exam_to_repair, prof_assignments, prof_large_counts, settings_for_validation, date_map):
                    # (## تعديل مدمج ##): استخدام العبء الموزون `prof_workload` لاختيار الأفضل
                    valid_candidates.append((prof, prof_workload.get(prof, 0)))

            if valid_candidates:
                # (## تعديل مدمج ##): اختيار الأستاذ صاحب أقل عبء عمل موزون
                best_prof_found, _ = min(valid_candidates, key=lambda item: item[1])
                
                # تعيين الأستاذ في الخانة الفارغة
                exam_to_repair['guards'][repair_info['index_to_fill']] = best_prof_found
                
                # (## تعديل مدمج ##): تحديث حالة الأستاذ ديناميكيًا لتؤثر على الاختيار التالي
                prof_assignments[best_prof_found].append(exam_to_repair)
                prof_workload[best_prof_found] += repair_duty_weight
                if is_large_repair_exam:
                    prof_large_counts[best_prof_found] += 1
        
        # --- 6. مرحلة القبول والتحديثات (تبقى كما هي) ---
        # ✅ الكود الجديد (المصحح)
        repaired_solution = ruined_solution # This line might not be in your code, but the context is the same
        new_cost = calculate_cost(repaired_solution, settings, all_professors, duty_patterns, date_map)
        
        # نحسب "الطاقة" الإجمالية لكل حل كرقم واحد (مجموع موزون)
        # وذلك فقط لاستخدامها في معادلة القبول العشوائي
        weights = (100000, 50000, 10, 1) # أوزان العقوبات
        current_energy = sum(c * w for c, w in zip(current_cost, weights))
        new_energy = sum(c * w for c, w in zip(new_cost, weights))
        
        # الآن نستخدم هذه الطاقة في المعادلة
        if new_cost < current_cost or random.random() < (math.exp((current_energy - new_energy) / temp) if temp > 0 else 0):
            current_solution, current_cost = repaired_solution, new_cost
        
        if current_cost < best_cost_so_far:
            best_cost_so_far = current_cost
            best_solution_so_far = copy.deepcopy(current_solution)
            log_q.put(f"... [LNS] دورة {i+1}: تم إيجاد حل أفضل بتكلفة = {format_cost_tuple(best_cost_so_far)}")
            if best_cost_so_far[0] == 0 and best_cost_so_far[1] == 0:
                # سنطبع الرسالة كعلامة فارقة، لكن لن نوقف البحث
                log_q.put("... [LNS] تم العثور على حل صالح ومكتمل (سنواصل البحث عن تحسينات).")
        
        temp *= cooling_rate
        dynamic_destroy_fraction = max(min_destroy_fraction, dynamic_destroy_fraction * destroy_fraction_decay_rate)

    log_q.put(f"✓ انتهى LNS المحسن بأفضل تكلفة: {format_cost_tuple(best_cost_so_far)}")
    
    # --- 7. إعادة بناء البيانات النهائية (تبقى كما هي) ---
    final_assignments = defaultdict(list)
    final_workload = defaultdict(float)
    final_large_counts = defaultdict(int)
    for day_data in best_solution_so_far.values():
        for slot_data in day_data.values():
            for exam in slot_data:
                is_large_exam = any(h['type'] == 'كبيرة' for h in exam['halls'])
                duty_weight = large_hall_weight if is_large_exam else other_hall_weight
                for guard in exam.get('guards', []):
                    if guard != "**نقص**":
                        final_assignments[guard].append(exam)
                        final_workload[guard] += duty_weight
                        if is_large_exam: final_large_counts[guard] += 1
                        
    return best_solution_so_far, final_assignments, final_workload, final_large_counts

# ===================================================================
# --- END: النسخة النهائية المدمجة من LNS ---
# ===================================================================
# ================== نهاية الكود الجديد والمصحح بالكامل ==================


# ===================================================================
# --- START: النسخة النهائية من VNS (تستهدف النقص) ---
# ===================================================================
def run_variable_neighborhood_search(
    initial_schedule, settings, all_professors, duty_patterns, date_map, log_q, locked_guards=set(), stop_event=None
):
    """
    النسخة النهائية والمحسنة من VNS:
    - مرحلة الإصلاح بعد الهزة تستهدف كل حالات النقص.
    - تستخدم البحث المحلي لتحسين الحلول.
    """
    log_q.put(">>> تشغيل VNS (النسخة النهائية المستهدفة للنقص)...")

    # --- 1. استخلاص الإعدادات ---
    iterations = int(settings.get('vnsIterations', 100))
    k_max = int(settings.get('vnsMaxK', 25)) # زيادة k القصوى لإتاحة تغييرات أكبر
    local_search_swaps = 100 # زيادة عدد محاولات البحث المحلي

    settings_for_validation = {
        'dutyPatterns': duty_patterns,
        'unavailableDays': settings.get('unavailableDays', {}),
        'maxShifts': settings.get('maxShifts', '0'),
        'maxLargeHallShifts': settings.get('maxLargeHallShifts', '2')
    }

    # --- 2. الحل المبدئي والتكلفة الأولية ---
    current_solution = copy.deepcopy(initial_schedule)
    best_solution_so_far = copy.deepcopy(current_solution)

    current_cost = calculate_cost(current_solution, settings, all_professors, duty_patterns, date_map)
    best_cost_so_far = current_cost
    log_q.put(f"... [VNS] التكلفة الأولية = {format_cost_tuple(current_cost)}")

    # --- 3. حلقة VNS الرئيسية ---
    i = 0
    stop_early = False
    while i < iterations:
        if stop_event and stop_event.is_set(): break
        if settings.get('should_stop_event', threading.Event()).is_set(): break

        percent_complete = int(((i + 1) / iterations) * 100)
        log_q.put(f"PROGRESS:{percent_complete}")
        
        k = 1
        while k <= k_max:
            # --- 4أ. مرحلة الهز (Shaking) ---
            shaken_solution = copy.deepcopy(current_solution)
            
            duties_to_destroy = []
            for day in shaken_solution.values():
                for slot in day.values():
                    for exam in slot:
                        for g_idx, guard in enumerate(exam.get('guards', [])):
                            if guard != "**نقص**" and (exam.get('uuid'), guard) not in locked_guards:
                                duties_to_destroy.append({'exam': exam, 'guard_index': g_idx})

            if not duties_to_destroy: break 
            
            random.shuffle(duties_to_destroy)
            
            for j in range(min(k, len(duties_to_destroy))):
                duty_info = duties_to_destroy[j]
                duty_info['exam']['guards'][duty_info['guard_index']] = "**نقص**"

            # --- 4ب. مرحلة الإصلاح الشامل (Repair) ---
            # ✅ --- هذا هو نفس المنطق الذكي المستخدم في LNS --- ✅
            all_exams_in_shaken = [exam for day in shaken_solution.values() for slot in day.values() for exam in slot]
            
            shortage_slots = []
            for exam in all_exams_in_shaken:
                for idx, guard in enumerate(exam.get('guards', [])):
                    if guard == "**نقص**":
                        shortage_slots.append({'exam': exam, 'index_to_fill': idx})
            
            for repair_info in shortage_slots:
                prof_assignments = defaultdict(list)
                prof_large_counts = defaultdict(int)
                for exam in all_exams_in_shaken:
                    is_large = any(h['type'] == 'كبيرة' for h in exam.get('halls', []))
                    for guard in exam.get('guards', []):
                        if guard != "**نقص**":
                            prof_assignments[guard].append(exam)
                            if is_large: prof_large_counts[guard] += 1
                
                exam_to_repair = repair_info['exam']
                
                valid_candidates = []
                for prof in all_professors:
                    if prof in exam_to_repair.get('guards', []): continue
                    if is_assignment_valid(prof, exam_to_repair, prof_assignments, prof_large_counts, settings_for_validation, date_map):
                        valid_candidates.append(prof)
                
                if valid_candidates:
                    exam_to_repair['guards'][repair_info['index_to_fill']] = random.choice(valid_candidates)

            # --- 4ج. مرحلة البحث المحلي (Local Search) ---
            local_search_solution, _, _, _ = run_post_processing_swaps(
                shaken_solution, defaultdict(list), defaultdict(float), defaultdict(int),
                settings, all_professors, date_map, local_search_swaps, locked_guards
            )

            # --- 5. مرحلة التحديث (Move or not) ---
            new_cost = calculate_cost(local_search_solution, settings, all_professors, duty_patterns, date_map)

            if new_cost < current_cost:
                current_solution, current_cost = local_search_solution, new_cost
                log_q.put(f"... [VNS] دورة {i+1}, k={k}: تم العثور على حل أفضل بتكلفة = {format_cost_tuple(current_cost)}")
                
                if new_cost < best_cost_so_far:
                    best_cost_so_far, best_solution_so_far = new_cost, copy.deepcopy(current_solution)
                    if best_cost_so_far[0] == 0 and best_cost_so_far[1] == 0:
                        log_q.put("... [VNS] تم العثور على حل صالح ومكتمل (سنواصل البحث عن تحسينات).")
                k = 1
            else:
                k += 1
        
        if stop_early: break
        i += 1

    log_q.put(f"✓ انتهى VNS بأفضل تكلفة: {format_cost_tuple(best_cost_so_far)}")
    
    # إعادة بناء البيانات النهائية من أفضل حل
    final_assignments = defaultdict(list)
    final_workload = defaultdict(float)
    final_large_counts = defaultdict(int)
    large_hall_weight = float(settings.get('largeHallWeight', 3.0))
    other_hall_weight = float(settings.get('otherHallWeight', 1.0))
    for day in best_solution_so_far.values():
        for slot in day.values():
            for exam in slot:
                 is_large = any(h['type'] == 'كبيرة' for h in exam.get('halls', []))
                 duty_weight = large_hall_weight if is_large else other_hall_weight
                 for guard in exam.get('guards',[]):
                     if guard != "**نقص**":
                         final_assignments[guard].append(exam)
                         final_workload[guard] += duty_weight
                         if is_large:
                             final_large_counts[guard] += 1

    return best_solution_so_far, final_assignments, final_workload, final_large_counts
# ===================================================================
# --- END: النسخة النهائية من VNS ---
# ===================================================================

# --- الدالة الثانية: دالة مساعدة جديدة ---
def calculate_deviation_from_stats(prof_stats, settings, all_professors):
    large_hall_weight = float(settings.get('largeHallWeight', 3.0))
    enable_custom_targets = settings.get('enableCustomTargets', False)
    custom_target_patterns = settings.get('customTargetPatterns', [])

    if enable_custom_targets and custom_target_patterns:
        target_counts = Counter((p['large'], p['other']) for p in custom_target_patterns for _ in range(p.get('count', 0)))
        actual_counts = Counter((s['large'], s['other']) for s in prof_stats.values())
        all_patterns = set(target_counts.keys()) | set(actual_counts.keys())
        total_deviation = sum(abs(actual_counts.get(p, 0) - target_counts.get(p, 0)) for p in all_patterns)
        return total_deviation * 2.0
    else:
        prof_workload = {p: s['large'] * large_hall_weight + s['other'] for p, s in prof_stats.items()}
        if prof_workload:
            workload_values = list(prof_workload.values())
            return max(workload_values) - min(workload_values) if workload_values else 0.0
    return 0.0

# =====================================================================================
# --- START: UNIFIED LNS ALGORITHM V13 (INTELLIGENT SWAPPING) ---
# =====================================================================================

def run_unified_lns_optimizer(initial_schedule, settings, all_professors, assignments, duty_patterns, date_map, all_subjects, log_q, all_levels_list, locked_guards=set(), stop_event=None):
    """
    (النسخة V13) "التبديل الذكي"
    - تعيد تقديم ميزة تبديل المواد ككتلة واحدة (المادة + الأستاذ) بشكل آمن.
    - تستخدم التبديل كأداة لتحسين القيود المرنة (تجميع أيام الأستاذ).
    - أي نقص ناتج عن التبديل يتم إصلاحه فوراً بواسطة دالة الإصلاح الشاملة.
    """
    log_q.put(">>> بدء تشغيل مُحسِّن LNS بالتبديل الذكي (V13)...")

    # --- 1. الإعدادات ---
    iterations = int(settings.get('lnsUnifiedIterations', 300))
    destroy_fraction = float(settings.get('lnsUnifiedDestroyFraction', 0.25))
    initial_temp = 5.0
    cooling_rate = 0.995

    # --- 2. الحل المبدئي ---
    current_solution = complete_schedule_with_guards(
        initial_schedule, settings, all_professors, assignments,
        all_levels_list, duty_patterns, date_map, all_subjects, locked_guards=locked_guards
    )
    best_solution = copy.deepcopy(current_solution)
    current_cost = calculate_cost(current_solution, settings, all_professors, duty_patterns, date_map)
    best_cost = current_cost
    log_q.put(f"... التكلفة الأولية للحل: {format_cost_tuple(best_cost)}")
    
    temp = initial_temp

    # --- 3. الحلقة الرئيسية ---
    for i in range(iterations):
        if stop_event and stop_event.is_set():
            log_q.put("... [المرحلة 1] تم الإيقاف بواسطة المستخدم.")
            break
        log_q.put(f"PROGRESS:{int(((i+1)/iterations)*100)}")
        if settings.get('should_stop_event', threading.Event()).is_set(): break

        neighbor_solution = copy.deepcopy(current_solution)
        was_changed = False # Flag to indicate if repair is needed

        # --- 4. بناء هياكل البيانات المساعدة ---
        prof_guard_days = defaultdict(set); prof_subject_days = defaultdict(set)
        prof_stats = {p: {'large': 0, 'other': 0} for p in all_professors}
        prof_duties = defaultdict(list); prof_assignments = defaultdict(list); prof_large_counts = defaultdict(int)
        all_exams_flat = [exam for day in neighbor_solution.values() for slot in day.values() for exam in slot]
        
        guards_large_hall = int(settings.get('guardsLargeHall', 4))
        guards_medium_hall = int(settings.get('guardsMediumHall', 2))
        guards_small_hall = int(settings.get('guardsSmallHall', 1))

        for exam in all_exams_flat:
            owner = exam.get('professor', "غير محدد")
            if owner != "غير محدد": prof_subject_days[owner].add(exam['date'])
            is_large = any(h['type'] == 'كبيرة' for h in exam.get('halls',[]))
            large_guards_needed = sum(guards_large_hall for h in exam.get('halls', []) if h.get('type') == 'كبيرة')
            for idx, guard in enumerate(exam.get('guards', [])):
                if guard != "**نقص**":
                    prof_guard_days[guard].add(exam['date'])
                    prof_duties[guard].append({'exam': exam, 'guard_index': idx})
                    prof_assignments[guard].append(exam)
                    if idx < large_guards_needed: prof_stats[guard]['large'] += 1
                    else: prof_stats[guard]['other'] += 1
                    if is_large: prof_large_counts[guard] += 1

        # --- 5. تحديد المرحلة واختيار الأداة ---
        is_currently_valid = (current_cost[0] == 0 and current_cost[1] == 0)

        if not is_currently_valid:
            # --- المرحلة 1: وضع الإزالة (نفس منطق V12) ---
            # ... (الكود يبقى كما هو)
            duties_to_destroy = [d for g, duties in prof_duties.items() for d in duties if (d['exam'].get('uuid'), g) not in locked_guards]
            if duties_to_destroy:
                random.shuffle(duties_to_destroy)
                num_to_destroy = int(len(duties_to_destroy) * destroy_fraction)
                for duty_info in duties_to_destroy[:num_to_destroy]:
                    exam_ref = duty_info['exam']
                    neighbor_solution[exam_ref['date']][exam_ref['time']][
                        [e['uuid'] for e in neighbor_solution[exam_ref['date']][exam_ref['time']]].index(exam_ref['uuid'])
                    ]['guards'][duty_info['guard_index']] = "**نقص**"
                was_changed = True
        else:
            # --- المرحلة 2: وضع التحسين (مع إضافة التبديل الذكي) ---
            tool_choice = random.random()
            if tool_choice < 0.6: # 60% فرصة لمحاولة تحسين الانحراف
                # ... (أداة التبرع من V12)
                # ... (الكود يبقى كما هو)
                custom_patterns = settings.get('customTargetPatterns', [])
                target_counts = Counter((p['large'], p['other']) for p in custom_patterns for _ in range(p.get('count', 0)))
                actual_counts = Counter((s['large'], s['other']) for s in prof_stats.values())
                over_patterns = {p for p, a in actual_counts.items() if a > target_counts.get(p, 0)}
                donors = [p for p, s in prof_stats.items() if (s['large'], s['other']) in over_patterns]
                recipients = all_professors
                random.shuffle(donors); random.shuffle(recipients)
                move_made = False
                for prof_donor in donors:
                    donatable_duties = [d for d in prof_duties[prof_donor] if (d['exam'].get('uuid'), prof_donor) not in locked_guards]
                    random.shuffle(donatable_duties)
                    for duty_to_donate in donatable_duties:
                        for prof_recipient in recipients:
                            if prof_donor == prof_recipient: continue
                            exam_to_reassign = duty_to_donate['exam']
                            if is_assignment_valid(prof_recipient, exam_to_reassign, prof_assignments, prof_large_counts, settings, date_map):
                                neighbor_solution[exam_to_reassign['date']][exam_to_reassign['time']][
                                    [e['uuid'] for e in neighbor_solution[exam_to_reassign['date']][exam_to_reassign['time']]].index(exam_to_reassign['uuid'])
                                ]['guards'][duty_to_donate['guard_index']] = prof_recipient
                                move_made = True
                                break
                        if move_made: break
                    if move_made: break
            else: # 40% فرصة لمحاولة تبديل المواد لتحسين القيود المرنة
                # ✅ --- الأداة الجديدة: التبديل الذكي للمواد ---
                improvement_candidates = [p for p, s_days in prof_subject_days.items() if not s_days.issubset(prof_guard_days.get(p, set()))]
                if improvement_candidates:
                    prof_to_fix = random.choice(improvement_candidates)
                    guard_days = prof_guard_days.get(prof_to_fix, set())
                    non_guard_subject_days = prof_subject_days.get(prof_to_fix, set()) - guard_days
                    
                    if non_guard_subject_days and guard_days:
                        day_from = random.choice(list(non_guard_subject_days))
                        day_to = random.choice(list(guard_days))
                        
                        exam_A = next((e for e in all_exams_flat if e['date'] == day_from and e['professor'] == prof_to_fix), None)
                        
                        # ابحث عن شريك تبديل في اليوم المستهدف بنفس المستوى
                        partners = [e for e in all_exams_flat if e['date'] == day_to and e['level'] == exam_A.get('level')]
                        if exam_A and partners:
                            exam_B = random.choice(partners)

                            # تبديل معلومات الامتحان ككتلة واحدة
                            props_A = {'subject': exam_A['subject'], 'professor': exam_A['professor'], 'halls': exam_A['halls']}
                            props_B = {'subject': exam_B['subject'], 'professor': exam_B['professor'], 'halls': exam_B['halls']}
                            exam_A.update(props_B)
                            exam_B.update(props_A)

                            # الآن، تعديل قوائم الحراس لتناسب المتطلبات الجديدة
                            for exam in [exam_A, exam_B]:
                                needed = (sum(guards_large_hall for h in exam.get('halls',[]) if h['type']=='كبيرة') +
                                          sum(guards_medium_hall for h in exam.get('halls',[]) if h['type']=='متوسطة') +
                                          sum(guards_small_hall for h in exam.get('halls',[]) if h['type']=='صغيرة'))
                                
                                current_guards = [g for g in exam.get('guards', []) if g != "**نقص**"]
                                if len(current_guards) > needed:
                                    exam['guards'] = current_guards[:needed]
                                else:
                                    exam['guards'] = current_guards + ["**نقص**"] * (needed - len(current_guards))
                            
                            was_changed = True # تفعيل الإصلاح الشامل

        # --- 6. الإصلاح (فقط إذا حدث تغيير يتطلب ذلك) ---
        final_neighbor = neighbor_solution
        if was_changed:
            final_neighbor = complete_schedule_with_guards(neighbor_solution, settings, all_professors, assignments, all_levels_list, duty_patterns, date_map, all_subjects, locked_guards=locked_guards)

        # --- 7. التقييم ومنطق القبول (نفس منطق V12) ---
        new_cost = calculate_cost(final_neighbor, settings, all_professors, duty_patterns, date_map)
        
        accepted = False
        if not is_currently_valid:
            if new_cost[0] < current_cost[0] or (new_cost[0] == current_cost[0] and new_cost[1] < current_cost[1]):
                accepted = True
        else:
            if new_cost[0] > 0 or new_cost[1] > 0:
                accepted = False
            else:
                cost_diff = sum(w * (n - c) for w, n, c in zip((10, 1), new_cost[2:], current_cost[2:]))
                if cost_diff < 0 or random.random() < math.exp(-cost_diff / temp if temp > 0 else float('-inf')):
                    accepted = True

        if accepted:
            current_solution, current_cost = final_neighbor, new_cost
            if new_cost < best_cost:
                best_solution, best_cost = copy.deepcopy(current_solution), new_cost
                log_q.put(f"... [مُحسِّن LNS v13] دورة {i+1}: حل أفضل بتكلفة = {format_cost_tuple(best_cost)}")
                if best_cost[0] == 0 and best_cost[1] == 0:
                    log_q.put("... تم الوصول إلى حل صحيح! التركيز الآن على التحسين.")

        temp = max(0.1, temp * cooling_rate)

    log_q.put(f"✓ انتهى مُحسِّن LNS v13 بأفضل تكلفة: {format_cost_tuple(best_cost)}")
    return best_solution, True

# =====================================================================================
# --- END: UNIFIED LNS ALGORITHM V13 ---
# =====================================================================================




# =====================================================================================
# --- START: FINAL COMPLETE GENETIC ALGORITHM (WITH ALL CONSTRAINTS) ---
# =====================================================================================


def run_genetic_algorithm(fixed_subject_schedule, settings, all_professors, assignments, all_levels_list, all_halls, exam_schedule_settings, all_subjects, level_hall_assignments, date_map, log_q, locked_guards=set(), stop_event=None):
    """
    النسخة النهائية والمكتملة من الخوارزمية الجينية مع جميع القيود الصارمة.
    """
    log_q.put(">>> [Genetic Alg v3] بدء الخوارزمية الجينية المحدثة...")
    # --- استخلاص الإعدادات ---
    pop_size = int(settings.get('geneticPopulation', 100))
    num_generations = int(settings.get('geneticGenerations', 500))
    crossover_rate = 0.8
    mutation_rate = float(settings.get('geneticMutation', 0.15))
    elitism_count = int(settings.get('geneticElitism', 4))
    
    # --- 1. إعطاء كل امتحان مُعرّف فريد (UUID) وتحضير خانات الحراسة ---
    schedule_with_ids = copy.deepcopy(fixed_subject_schedule)
    all_exams_flat = [exam for slots in schedule_with_ids.values() for exams in slots.values() for exam in exams]
    for exam in all_exams_flat:
        exam['uuid'] = str(uuid.uuid4())

    duty_slots = []
    guards_large_hall = int(settings.get('guardsLargeHall', 4))
    guards_medium_hall = int(settings.get('guardsMediumHall', 2))
    guards_small_hall = int(settings.get('guardsSmallHall', 1))

    for exam in all_exams_flat:
        num_guards_needed = (sum(guards_large_hall for h in exam.get('halls', []) if h.get('type') == 'كبيرة') +
                           sum(guards_medium_hall for h in exam.get('halls', []) if h.get('type') == 'متوسطة') +
                           sum(guards_small_hall for h in exam.get('halls', []) if h.get('type') == 'صغيرة'))
        for _ in range(num_guards_needed):
            duty_slots.append(exam)

    if not duty_slots:
        return fixed_subject_schedule, True

    # --- 2. تعريف دوال مساعدة ---
    
    def build_schedule_from_chromosome(chromosome):
        schedule_copy = copy.deepcopy(schedule_with_ids)
        exam_map = {ex['uuid']: ex for slots in schedule_copy.values() for exams in slots.values() for ex in exams}
        for ex in exam_map.values():
            ex['guards'] = []

        for i, guard in enumerate(chromosome):
            exam_ref_uuid = duty_slots[i]['uuid']
            exam_in_copy = exam_map.get(exam_ref_uuid)
            if exam_in_copy:
                exam_in_copy['guards'].append(guard)
        return schedule_copy

    def create_random_chromosome():
        chromosome = [None] * len(duty_slots)
        prof_assignments = {prof: [] for prof in all_professors}
        prof_large_counts = defaultdict(int)
        
        unavailable_days = settings.get('unavailableDays', {})
        duty_patterns = settings.get('dutyPatterns', {})
        max_shifts = int(settings.get('maxShifts', '0')) if settings.get('maxShifts', '0') != '0' else float('inf')
        max_large_hall_shifts = int(settings.get('maxLargeHallShifts', '2')) if settings.get('maxLargeHallShifts', '2') != '0' else float('inf')
        assign_owner_as_guard = settings.get('assignOwnerAsGuard', False)

        # ## <<< بداية: إضافة منطق تعيين أستاذ المادة الإلزامي >>>
        if assign_owner_as_guard:
            subject_owners = { (clean_string_for_matching(s['name']), clean_string_for_matching(s['level'])): clean_string_for_matching(prof) for prof, uids in assignments.items() for uid in uids for s in all_subjects if f"{s['name']} ({s['level']})" == uid }
            prof_last_exam = {}
            for exam in all_exams_flat:
                owner = subject_owners.get((clean_string_for_matching(exam['subject']), clean_string_for_matching(exam['level'])))
                if owner:
                    exam_date_time = (exam['date'], exam['time'])
                    if owner not in prof_last_exam or exam_date_time > prof_last_exam[owner]['datetime']:
                        prof_last_exam[owner] = {'exam': exam, 'datetime': exam_date_time}

            for owner, data in prof_last_exam.items():
                exam_to_assign = data['exam']
                
                # إيجاد خانة حراسة فارغة لهذا الامتحان
                duty_index_to_fill = -1
                for i, slot_exam in enumerate(duty_slots):
                    if slot_exam['uuid'] == exam_to_assign['uuid'] and chromosome[i] is None:
                        duty_index_to_fill = i
                        break
                
                if duty_index_to_fill != -1:
                    # التحقق من صلاحية هذا التعيين الإلزامي
                    is_large = any(h['type'] == 'كبيرة' for h in exam_to_assign.get('halls', []))
                    if exam_to_assign['date'] not in unavailable_days.get(owner, []) and \
                       len(prof_assignments[owner]) < max_shifts and \
                       (not is_large or prof_large_counts[owner] < max_large_hall_shifts):
                        
                        chromosome[duty_index_to_fill] = owner
                        prof_assignments[owner].append(exam_to_assign)
                        if is_large:
                            prof_large_counts[owner] += 1
        # ## <<< نهاية: منطق التعيين الإلزامي >>>

        shuffled_indices = list(range(len(duty_slots)))
        random.shuffle(shuffled_indices)

        for i in shuffled_indices:
            if chromosome[i] is not None: continue # تخطي الخانات التي تم ملؤها

            exam = duty_slots[i]
            is_large_exam = any(h['type'] == 'كبيرة' for h in exam.get('halls', []))
            
            shuffled_profs = list(all_professors)
            random.shuffle(shuffled_profs)

            assigned_prof = None
            for prof in shuffled_profs:
                # التحقق من أن الأستاذ غير معين بالفعل في هذه الخانة
                if any(chromosome[j] == prof for j, slot in enumerate(duty_slots) if slot['uuid'] == exam['uuid']):
                    continue

                is_busy = any(a_exam['date'] == exam['date'] and a_exam['time'] == exam['time'] for a_exam in prof_assignments[prof])
                if is_busy: continue
                if exam['date'] in unavailable_days.get(prof, []): continue
                if len(prof_assignments[prof]) >= max_shifts: continue
                if is_large_exam and prof_large_counts[prof] >= max_large_hall_shifts: continue

                prof_pattern = duty_patterns.get(prof, 'flexible_2_days')
                current_prof_dates = {e['date'] for e in prof_assignments.get(prof, [])}
                is_new_day = exam['date'] not in current_prof_dates
                
                is_pattern_violated = False
                if is_new_day:
                    num_current_days = len(current_prof_dates)
                    if prof_pattern == 'one_day_only' and num_current_days >= 1: is_pattern_violated = True
                    elif prof_pattern == 'flexible_2_days' and num_current_days >= 2: is_pattern_violated = True
                    elif prof_pattern == 'flexible_3_days' and num_current_days >= 3: is_pattern_violated = True
                    elif prof_pattern == 'consecutive_strict':
                        if num_current_days >= 2: is_pattern_violated = True
                        elif num_current_days == 1:
                            existing_day_idx = date_map.get(list(current_prof_dates)[0])
                            new_day_idx = date_map.get(exam['date'])
                            if existing_day_idx is None or new_day_idx is None or abs(new_day_idx - existing_day_idx) != 1:
                                is_pattern_violated = True
                if is_pattern_violated: continue
                
                assigned_prof = prof
                break
            
            if assigned_prof:
                chromosome[i] = assigned_prof
                prof_assignments[assigned_prof].append(exam)
                if is_large_exam:
                    prof_large_counts[assigned_prof] += 1
            else:
                chromosome[i] = "**نقص**"
                
        return chromosome

    def calculate_fitness(chromosome):
        """
        (النسخة المحدثة) تحسب درجة اللياقة بناءً على التكلفة المفصلة.
        """
        schedule = build_schedule_from_chromosome(chromosome)
        
        # استدعاء دالة التكلفة الموحدة
        cost_tuple = calculate_cost(schedule, settings, all_professors, settings.get('dutyPatterns', {}), date_map)
        shortage, hard_violations, deviation, soft_violations = cost_tuple

        # تحويل التكلفة إلى درجة لياقة (رقم واحد، كلما كان أعلى كان أفضل)
        # نبدأ برقم كبير ونخصم منه العقوبات بأوزان مختلفة
        fitness_score = 1_000_000
        fitness_score -= shortage * 10000
        fitness_score -= hard_violations * 5000
        fitness_score -= deviation * 10
        fitness_score -= soft_violations * 1
        
        # نضمن أن اللياقة لا تقل عن صفر
        return max(0.0, fitness_score)

    def crossover(parent1, parent2):
        point = random.randint(1, len(parent1) - 1)
        return parent1[:point] + parent2[point:], parent2[:point] + parent1[point:]

    def mutation(chromosome):
        if len(chromosome) < 2: return chromosome
        
        # اختيار خانتين مختلفتين للتبديل
        idx1, idx2 = random.sample(range(len(chromosome)), 2)
        
        # عمل نسخة من الكروموسوم لتجنب التعديل المباشر
        mutated_chromosome = list(chromosome)
        mutated_chromosome[idx1], mutated_chromosome[idx2] = mutated_chromosome[idx2], mutated_chromosome[idx1]
        
        # التحقق من صلاحية الطفرة
        temp_schedule = build_schedule_from_chromosome(mutated_chromosome)
        if is_schedule_valid(temp_schedule, settings, all_professors, settings.get('dutyPatterns', {}), date_map):
            return mutated_chromosome # إرجاع الحل الجديد الصالح
        return chromosome # التراجع عن الطفرة إذا كانت ستنتج حلاً غير صالح

    # --- 3. بدء الخوارزمية ---
    log_queue.put(">>> [Genetic Alg v2] بدء الخوارزمية الجينية المحسنة...")
    log_queue.put("... بناء المجتمع الأولي على جدول المواد الثابت.")
    
    try:
        population = [create_random_chromosome() for _ in range(pop_size)]
    except Exception as e:
        import traceback
        log_queue.put(f"!!! Error during initial population creation: {e}")
        log_queue.put(traceback.format_exc())
        return fixed_subject_schedule, False
    
    best_chromosome, best_fitness = None, -1.0

    for gen in range(num_generations):
        if stop_event and stop_event.is_set():
            log_q.put("... [Genetic Alg] تم الإيقاف بواسطة المستخدم.")
            break
        # =================> بداية الإضافة الجديدة <=================
        # إرسال التقدم بناءً على الجيل الحالي
        percent_complete = int(((gen + 1) / num_generations) * 100)
        log_queue.put(f"PROGRESS:{percent_complete}")
        # =================> نهاية الإضافة الجديدة <=================
        population_with_fitness = [(chrom, calculate_fitness(chrom)) for chrom in population]
        
        valid_population_with_fitness = [item for item in population_with_fitness if item[1] > 0]
        if not valid_population_with_fitness:
             log_queue.put(f"... [Generation {gen+1}] Population collapsed. Rebuilding...")
             population = [create_random_chromosome() for _ in range(pop_size)]
             continue

        current_best_chrom, current_best_fitness = max(valid_population_with_fitness, key=lambda item: item[1])
        if current_best_fitness > best_fitness:
            best_fitness, best_chromosome = current_best_fitness, current_best_chrom
            log_queue.put(f"... [Generation {gen+1}/{num_generations}] New best fitness: {best_fitness:.2f}%")

        if best_fitness >= 99.9:
            log_queue.put("✓ تم العثور على حل مثالي. إنهاء البحث.")
            break
        
        valid_population_with_fitness.sort(key=lambda item: item[1], reverse=True)
        next_generation = [item[0] for item in valid_population_with_fitness[:elitism_count]]
        
        while len(next_generation) < pop_size:
            parents = random.choices(valid_population_with_fitness, weights=[f for _, f in valid_population_with_fitness], k=2)
            p1, p2 = parents[0][0], parents[1][0]

            if random.random() < crossover_rate:
                c1, c2 = crossover(p1, p2)
                next_generation.append(mutation(c1) if random.random() < mutation_rate else c1)
                if len(next_generation) < pop_size:
                    next_generation.append(mutation(c2) if random.random() < mutation_rate else c2)
            else:
                next_generation.append(p1)
                if len(next_generation) < pop_size:
                    next_generation.append(p2)
        population = next_generation

    log_queue.put(f">>> [Genetic Alg v2] Finished. Best fitness score: {best_fitness:.2f}%")
    
    if not best_chromosome:
        log_queue.put("!!! لم تتمكن الخوارزمية الجينية من إيجاد حل صالح.")
        any_valid = [chrom for chrom, fit in population_with_fitness if fit > 0]
        if any_valid:
            best_chromosome = any_valid[0]
        else:
            return fixed_subject_schedule, False

    final_schedule = build_schedule_from_chromosome(best_chromosome)
    return final_schedule, True

# --- END: FINAL COMPLETE GENETIC ALGORITHM ---
# =====================================================================================


# ===================================================================
# --- START: النسخة المحسنة (مبدأ الأستاذ الأقل حيوية) ---
# ===================================================================
def complete_schedule_with_guards(subject_schedule, settings, all_professors, assignments, all_levels_list, duty_patterns, date_map, all_subjects, locked_guards=set(), stop_event=None, log_q=None):
    """
    النسخة المحسّنة: تستخدم "مبدأ الأستاذ الأقل حيوية" لاختيار الحارس
    الأمثل، مما يقلل بشكل كبير من احتمالية حدوث نقص في الحراسة.
    """
    schedule = copy.deepcopy(subject_schedule)
    
    # --- إعدادات وقواميس مساعدة (نفس السابق) ---
    guards_large_hall = int(settings.get('guardsLargeHall', 4))
    guards_medium_hall = int(settings.get('guardsMediumHall', 2))
    guards_small_hall = int(settings.get('guardsSmallHall', 1))
    
    settings_for_validation = {
        'dutyPatterns': duty_patterns,
        'unavailableDays': settings.get('unavailableDays', {}),
        'maxShifts': settings.get('maxShifts', '0'),
        'maxLargeHallShifts': settings.get('maxLargeHallShifts', '2')
    }

    # --- 1. تحديد كل الخانات الفارغة وتطبيق المهام المقفلة (نفس السابق) ---
    duties_to_fill = []
    all_scheduled_exams_flat = [exam for day in schedule.values() for slot in day.values() for exam in slot]
    for exam in all_scheduled_exams_flat:
        if 'uuid' not in exam: exam['uuid'] = str(uuid.uuid4())
        
        locked_profs_for_exam = {p for e_uuid, p in locked_guards if e_uuid == exam.get('uuid')}
        if 'guards' not in exam: exam['guards'] = []
        for prof in locked_profs_for_exam:
            if prof not in exam['guards']:
                exam['guards'].append(prof)

        num_needed = (sum(guards_large_hall for h in exam.get('halls',[]) if h.get('type')=='كبيرة') +
                      sum(guards_medium_hall for h in exam.get('halls',[]) if h.get('type')=='متوسطة') +
                      sum(guards_small_hall for h in exam.get('halls',[]) if h.get('type')=='صغيرة'))
        
        num_to_add = num_needed - len(exam.get('guards', []))
        for _ in range(num_to_add):
            duties_to_fill.append(exam)
    
    # --- 2. الحلقة الديناميكية: في كل خطوة، جدد التفكير (نفس السابق) ---
    while duties_to_fill:
        if stop_event and stop_event.is_set():
            log_q.put("... [توزيع الحراس] تم الإيقاف بواسطة المستخدم.")
            break
        # --- 2أ. إعادة بناء الحالة الحالية للحراس (نفس السابق) ---
        prof_assignments = defaultdict(list)
        prof_large_counts = defaultdict(int)
        for exam in all_scheduled_exams_flat:
            is_large = any(h['type'] == 'كبيرة' for h in exam.get('halls', []))
            for guard in exam.get('guards', []):
                if guard in all_professors:
                    prof_assignments[guard].append(exam)
                    if is_large: prof_large_counts[guard] += 1
        
        # --- 2ب. إعادة تحليل صعوبة كل المهام المتبقية الآن (نفس السابق) ---
        duties_with_candidate_count = []
        # إنشاء خريطة للمهام المتبقية حسب التوقيت لتسهيل البحث لاحقًا
        remaining_duties_by_slot = defaultdict(list)
        for duty_exam in duties_to_fill:
            if stop_event and stop_event.is_set(): break
            candidate_count = 0
            for prof in all_professors:
                if prof in duty_exam.get('guards', []): continue
                if is_assignment_valid(prof, duty_exam, prof_assignments, prof_large_counts, settings_for_validation, date_map):
                    candidate_count += 1
            duties_with_candidate_count.append({'exam': duty_exam, 'candidates': candidate_count})
            remaining_duties_by_slot[(duty_exam['date'], duty_exam['time'])].append(duty_exam)

        if not duties_with_candidate_count: break

        # --- 2ج. تحديد المهمة الأصعب حاليًا (نفس السابق) ---
        hardest_duty_info = min(duties_with_candidate_count, key=lambda x: x['candidates'])
        hardest_duty_exam = hardest_duty_info['exam']
        
        # --- 2د. إيجاد أفضل حارس للمهمة الأصعب (المنطق الجديد والمحسّن) ---
        valid_candidates_with_scores = []
        
        # الحصول على قائمة المهام الأخرى المتزامنة مع المهمة الأصعب
        concurrent_duties = [
            d for d in remaining_duties_by_slot[(hardest_duty_exam['date'], hardest_duty_exam['time'])]
            if d is not hardest_duty_exam
        ]
        
        # إيجاد قائمة المرشحين الصالحين للمهمة الصعبة أولاً
        valid_candidates_for_hardest = []
        for prof in all_professors:
            if prof in hardest_duty_exam.get('guards', []): continue
            if is_assignment_valid(prof, hardest_duty_exam, prof_assignments, prof_large_counts, settings_for_validation, date_map):
                valid_candidates_for_hardest.append(prof)
        
        # حساب درجة الحيوية لكل مرشح صالح
        for prof in valid_candidates_for_hardest:
            if stop_event and stop_event.is_set(): break
            criticality_score = 0
            # احسب كم مرة هذا الأستاذ هو مرشح "حيوي" للمهام المتزامنة الأخرى
            for other_duty in concurrent_duties:
                # احسب عدد المرشحين الكلي للمهمة الأخرى
                other_duty_candidates = 0
                for other_prof in all_professors:
                    if other_prof in other_duty.get('guards',[]): continue
                    if is_assignment_valid(other_prof, other_duty, prof_assignments, prof_large_counts, settings_for_validation, date_map):
                        other_duty_candidates += 1
                
                # إذا كان أستاذنا الحالي مرشحاً لهذه المهمة الصعبة، زد درجة حيويته
                if other_duty_candidates <= 2 and is_assignment_valid(prof, other_duty, prof_assignments, prof_large_counts, settings_for_validation, date_map):
                    criticality_score += 1
            
            workload = len(prof_assignments.get(prof, []))
            # نخزن (درجة الحيوية، عبء العمل، اسم الأستاذ)
            valid_candidates_with_scores.append((criticality_score, workload, prof))
        
        
        best_prof_found = None
        if valid_candidates_with_scores:
            # فرز المرشحين: أولاً حسب أقل درجة حيوية، ثم حسب أقل عبء عمل
            valid_candidates_with_scores.sort(key=lambda x: (x[0], x[1]))
            best_prof_found = valid_candidates_with_scores[0][2]
            hardest_duty_exam['guards'].append(best_prof_found)
        else:
            hardest_duty_exam['guards'].append("**نقص**")
        
        # --- 2هـ. إزالة المهمة التي تم حلها من القائمة (نفس السابق) ---
        try:
            duties_to_fill.remove(hardest_duty_exam)
        except ValueError:
             for i, item in enumerate(duties_to_fill):
                 if item is hardest_duty_exam:
                     del duties_to_fill[i]
                     break

    return schedule
# ===================================================================
# --- END: النسخة المحسنة ---
# ===================================================================

# في ملف app.py، استبدل الدالة بالكامل بهذه النسخة النهائية (الإصدار التاسع - المستقر)
def run_constraint_solver(original_schedule, settings, all_professors, assignments, all_levels_list, subject_owners, last_day_restriction, sorted_dates, duty_patterns, date_map, log_q, stop_event=None):
    """
    النسخة المحدثة: مع إضافة قيد "أزواج الأساتذة" الرياضي.
    """
    model = cp_model.CpModel()
    all_scheduled_exams = [exam for date_exams in original_schedule.values() for time_exams in date_exams.values() for exam in time_exams]

    # --- استخلاص الإعدادات ---
    large_hall_weight = int(settings.get('largeHallWeight', 3))
    other_hall_weight = int(settings.get('otherHallWeight', 1))
    max_shifts = int(settings.get('maxShifts', '0')) if settings.get('maxShifts', '0') != '0' else float('inf')
    max_large_hall_shifts = int(settings.get('maxLargeHallShifts', '2')) if settings.get('maxLargeHallShifts', '2') != '0' else float('inf')
    enable_custom_targets = settings.get('enableCustomTargets', False)
    custom_target_patterns = settings.get('customTargetPatterns', [])
    solver_timelimit = int(settings.get('solverTimelimit', 30))
    prof_map = {name: i for i, name in enumerate(all_professors)}
    num_professors = len(all_professors)
    guards_large_hall = int(settings.get('guardsLargeHall', 4))
    guards_medium_hall = int(settings.get('guardsMediumHall', 2))
    guards_small_hall = int(settings.get('guardsSmallHall', 1))
    unavailable_days = settings.get('unavailableDays', {})
    
    # --- بناء كل المتغيرات والقيود الصارمة (بما في ذلك قيد الأزواج) ---
    duties = []
    for exam in all_scheduled_exams:
        num_guards_needed = sum(guards_large_hall for h in exam.get('halls', []) if h.get('type') == 'كبيرة') + sum(guards_medium_hall for h in exam.get('halls', []) if h.get('type') == 'متوسطة') + sum(guards_small_hall for h in exam.get('halls', []) if h.get('type') == 'صغيرة')
        for _ in range(num_guards_needed): duties.append({'exam': exam, 'is_large': any(h['type'] == 'كبيرة' for h in exam['halls'])})

    x = {(p_idx, d_idx): model.NewBoolVar(f'x_{p_idx}_{d_idx}') for p_idx in range(num_professors) for d_idx in range(len(duties))}
    for d_idx in range(len(duties)): model.AddExactlyOne(x[p_idx, d_idx] for p_idx in range(num_professors))
    
    slots = defaultdict(list)
    for d_idx, duty in enumerate(duties): slots[(duty['exam']['date'], duty['exam']['time'])].append(d_idx)
    for p_idx in range(num_professors):
        for slot_duties in slots.values(): model.Add(sum(x[p_idx, d] for d in slot_duties) <= 1)
    
    if max_shifts != float('inf'):
        for p_idx in range(num_professors): model.Add(sum(x[p_idx, d_idx] for d_idx in range(len(duties))) <= max_shifts)
    if max_large_hall_shifts != float('inf'):
        large_duty_indices = [d_idx for d_idx, duty in enumerate(duties) if duty['is_large']]
        for p_idx in range(num_professors): model.Add(sum(x[p_idx, d_idx] for d_idx in large_duty_indices) <= max_large_hall_shifts)
    for prof, dates in unavailable_days.items():
        if prof in prof_map:
            p_idx = prof_map[prof]
            for d_idx, duty in enumerate(duties):
                if duty['exam']['date'] in dates: model.Add(x[p_idx, d_idx] == 0)

    is_duty_day, prof_has_any_duty = {}, {}
    for p_idx in range(num_professors):
        prof_has_any_duty[p_idx] = model.NewBoolVar(f'prof_has_duty_{p_idx}')
        duties_for_prof = [x[p_idx, d_idx] for d_idx in range(len(duties))]
        model.Add(sum(duties_for_prof) > 0).OnlyEnforceIf(prof_has_any_duty[p_idx]); model.Add(sum(duties_for_prof) == 0).OnlyEnforceIf(prof_has_any_duty[p_idx].Not())
        for day_idx in range(len(sorted_dates)):
            is_duty_day[p_idx, day_idx] = model.NewBoolVar(f'is_duty_day_{p_idx}_{day_idx}')
            duties_in_this_day = [x[p_idx, d_idx] for d_idx, duty in enumerate(duties) if date_map.get(duty['exam']['date']) == day_idx]
            if duties_in_this_day:
                model.AddBoolOr(duties_in_this_day).OnlyEnforceIf(is_duty_day[p_idx, day_idx]); model.Add(sum(duties_in_this_day) == 0).OnlyEnforceIf(is_duty_day[p_idx, day_idx].Not())
            else: model.Add(is_duty_day[p_idx, day_idx] == False)

    for prof_name, pattern in duty_patterns.items():
        if prof_name in prof_map:
            p_idx = prof_map[prof_name]
            num_unique_duty_days = sum(is_duty_day[p_idx, day_idx] for day_idx in range(len(sorted_dates)))
            if pattern == 'consecutive_strict':
                model.Add(num_unique_duty_days == 2).OnlyEnforceIf(prof_has_any_duty[p_idx])
                is_start_of_consecutive_pair = [model.NewBoolVar(f'is_start_{p_idx}_{d}') for d in range(len(sorted_dates) - 1)]
                for day_idx in range(len(sorted_dates) - 1):
                    model.AddBoolAnd([is_duty_day[p_idx, day_idx], is_duty_day[p_idx, day_idx+1]]).OnlyEnforceIf(is_start_of_consecutive_pair[day_idx])
                    model.AddBoolOr([is_duty_day[p_idx, day_idx].Not(), is_duty_day[p_idx, day_idx+1].Not()]).OnlyEnforceIf(is_start_of_consecutive_pair[day_idx].Not())
                model.Add(sum(is_start_of_consecutive_pair) == 1).OnlyEnforceIf(prof_has_any_duty[p_idx])
            elif pattern == 'one_day_only':
                model.Add(num_unique_duty_days <= 1).OnlyEnforceIf(prof_has_any_duty[p_idx])
            elif pattern == 'flexible_2_days': model.Add(num_unique_duty_days == 2).OnlyEnforceIf(prof_has_any_duty[p_idx])
            elif pattern == 'flexible_3_days':
                model.Add(num_unique_duty_days >= 2).OnlyEnforceIf(prof_has_any_duty[p_idx]); model.Add(num_unique_duty_days <= 3).OnlyEnforceIf(prof_has_any_duty[p_idx])
    
    # --- بداية: إضافة قيد أزواج الأساتذة ---
    professor_pairs = settings.get('professorPartnerships', []) # تم تغيير اسم المفتاح
    if professor_pairs:
        for pair in professor_pairs:
            if len(pair) == 2 and pair[0] in prof_map and pair[1] in prof_map:
                prof1_idx = prof_map[pair[0]]
                prof2_idx = prof_map[pair[1]]
                # لكل يوم في الجدول، يجب أن يكونا إما كلاهما يعمل أو كلاهما لا يعمل
                for day_idx in range(len(sorted_dates)):
                    model.Add(is_duty_day[prof1_idx, day_idx] == is_duty_day[prof2_idx, day_idx])
    # --- نهاية: إضافة قيد أزواج الأساتذة ---

    # --- الهدف (Objective) ---
    prof_large_duties = [model.NewIntVar(0, 100, f'large_{p}') for p in range(num_professors)]
    prof_other_duties = [model.NewIntVar(0, 100, f'other_{p}') for p in range(num_professors)]
    # ... (بقية الدالة تبقى كما هي دون تغيير) ...
    # ...
    #
    for p_idx in range(num_professors):
        model.Add(prof_large_duties[p_idx] == sum(x[p_idx, d_idx] for d_idx, duty in enumerate(duties) if duty['is_large']))
        model.Add(prof_other_duties[p_idx] == sum(x[p_idx, d_idx] for d_idx, duty in enumerate(duties) if not duty['is_large']))

    if enable_custom_targets and custom_target_patterns:
        log_q.put("... استخدام نموذج تقليل الانحراف مع معاقبة الأنماط غير المستهدفة")
        target_counts = Counter((p['large'], p['other']) for p in custom_target_patterns for _ in range(p.get('count', 0)))
        actual_counts = {}
        for l, o in target_counts.keys():
            prof_matches_pattern = [model.NewBoolVar(f'p_{p}_m_{l}_{o}') for p in range(num_professors)]
            for p_idx in range(num_professors):
                model.Add(prof_large_duties[p_idx] == l).OnlyEnforceIf(prof_matches_pattern[p_idx])
                model.Add(prof_other_duties[p_idx] == o).OnlyEnforceIf(prof_matches_pattern[p_idx])
            actual_counts[(l,o)] = sum(prof_matches_pattern)
        deviation_terms = []
        for pattern, target_count in target_counts.items():
            deviation = model.NewIntVar(0, num_professors, f'dev_{pattern}')
            model.AddAbsEquality(deviation, actual_counts[pattern] - target_count)
            deviation_terms.append(deviation)
        total_deviation = sum(deviation_terms)
        num_untracked = model.NewIntVar(0, num_professors, 'num_untracked')
        total_tracked_profs = sum(actual_counts.values())
        model.Add(num_untracked == num_professors - total_tracked_profs)
        model.Minimize(total_deviation + 10 * num_untracked)
    else:
        log_q.put("... استخدام دالة الهدف الافتراضية (الموازنة العامة)")
        prof_workload = [model.NewIntVar(0, 1000, f'workload_{p}') for p in range(num_professors)]
        for p_idx in range(num_professors):
            model.Add(prof_workload[p_idx] == prof_large_duties[p_idx] * large_hall_weight + prof_other_duties[p_idx] * other_hall_weight)
        min_w, max_w = model.NewIntVar(0, 1000, 'min_w'), model.NewIntVar(0, 1000, 'max_w')
        if prof_workload:
            model.AddMinEquality(min_w, prof_workload); model.AddMaxEquality(max_w, prof_workload)
        model.Minimize(max_w - min_w)

    # --- حل النموذج ---
    solver = cp_model.CpSolver()
    solver.parameters.max_time_in_seconds = float(solver_timelimit)
    if stop_event:
        class StopSearchCallback(cp_model.CpSolverSolutionCallback):
            def __init__(self, event):
                cp_model.CpSolverSolutionCallback.__init__(self)
                self._event = event
            def on_solution_callback(self):
                if self._event.is_set():
                    self.StopSearch()
        
        status = solver.Solve(model, StopSearchCallback(stop_event))
    else:
        status = solver.Solve(model)
    

    if status == cp_model.OPTIMAL or status == cp_model.FEASIBLE:
        log_q.put("✓ تم العثور على حل باستخدام البرمجة بالقيود.")
        new_schedule_list = copy.deepcopy(all_scheduled_exams)
        for exam in new_schedule_list: exam['guards'] = []
        exam_map_by_props = {(e['subject'], e['level'], e['date'], e['time']): e for e in new_schedule_list}
        for p_idx in range(num_professors):
            for d_idx, duty in enumerate(duties):
                if solver.Value(x[p_idx, d_idx]):
                    exam_props = duty['exam']
                    exam_key = (exam_props['subject'], exam_props['level'], exam_props['date'], exam_props['time'])
                    if exam_key in exam_map_by_props: exam_map_by_props[exam_key]['guards'].append(all_professors[p_idx])
        final_schedule = defaultdict(lambda: defaultdict(list))
        for exam in new_schedule_list: final_schedule[exam['date']][exam['time']].append(exam)
        return final_schedule, True
    else:
        log_q.put("✗ فشل! لم يتم العثور على أي حل صالح. هذا يؤكد وجود تضارب في القيود نفسها.")
        return original_schedule, False

# =====================================================================
# START: HYPER-HEURISTIC FRAMEWORK (CORRECTED & ENHANCED FOR PROJECT 2)
# =====================================================================
def run_hyper_heuristic(
    log_q, initial_schedule, settings, all_professors, assignments, all_levels_list, all_subjects,
    duty_patterns, date_map, all_halls, exam_schedule_settings, level_hall_assignments,
    locked_guards=set(), stop_event=None
):
    log_q.put("--- بدء تشغيل النظام الخبير (Hyper-Heuristic) ---")

    # --- 1. استخلاص إعدادات النظام الخبير ---
    hyper_settings = settings.get('hyperHeuristicSettings', {})
    iterations = int(hyper_settings.get('iterations', 50))
    tabu_tenure = int(hyper_settings.get('tabuTenure', 3))
    # ✨ تعديل: قراءة حد الركود الزمني
    stagnation_time_limit = int(hyper_settings.get('stagnationTimeLimit', 15))
    selected_llh = hyper_settings.get('selectedLLH', ['unified_lns', 'tabu_search', 'lns', 'vns'])

    learning_rate, discount_factor, epsilon, epsilon_decay = 0.1, 0.9, 0.5, 0.995

    # --- 2. تحديد الخوارزميات المتاحة (LLHs) ---
    all_available_llh = {
        "unified_lns": run_unified_lns_optimizer,
        "tabu_search": run_tabu_search,
        "lns": run_large_neighborhood_search,
        "vns": run_variable_neighborhood_search
    }
    
    actions = [name for name in selected_llh if name in all_available_llh]
    if not actions:
        log_q.put("خطأ فادح: لم يتم اختيار أي خوارزميات صالحة للنظام الخبير.")
        return None, False

    # --- 3. تحميل/تهيئة جدول الخبرة (Q-Table) ---
    Q_TABLE_PATH = os.path.join(DATA_DIR, 'q_table_hyper_heuristic.json')
    q_table = defaultdict(lambda: {action: 0.0 for action in actions})
    if os.path.exists(Q_TABLE_PATH):
        try:
            with open(Q_TABLE_PATH, 'r', encoding='utf-8') as f:
                saved_q_table = json.load(f)
                for state, action_values in saved_q_table.items():
                    q_table[state] = {action: action_values.get(action, 0.0) for action in actions}
                log_q.put("   - تم تحميل ذاكرة الخبرة بنجاح.")
        except Exception as e:
            log_q.put(f"   - تحذير: فشل تحميل ذاكرة الخبرة. (الخطأ: {e})")

    # --- 4. إنشاء حل مبدئي للإنطلاق منه ---
    log_q.put("   - جاري إنشاء حل مبدئي عالي الجودة...")
    # ✅ تصحيح: استخدام المتغير `initial_schedule` الذي تم تمريره
    current_solution = complete_schedule_with_guards(
        initial_schedule, settings, all_professors, assignments,
        all_levels_list, duty_patterns, date_map, all_subjects, locked_guards
    )
    
    current_cost = calculate_cost(current_solution, settings, all_professors, duty_patterns, date_map)
    best_cost_so_far = current_cost
    best_solution_so_far = copy.deepcopy(current_solution)
    log_q.put(f"   - الحل المبدئي: التكلفة = {format_cost_tuple(current_cost)}")

    # --- 5. تهيئة متغيرات حلقة التعلم ---
    tabu_list = deque(maxlen=tabu_tenure)
    # ✨ تعديل: استخدام متغيرات للوقت بدلاً من العداد
    last_improvement_time = time.time()

    # --- 6. حلقة التعلم والتحسين الرئيسية ---
    for i in range(iterations):
        if stop_event and stop_event.is_set(): break
        if settings.get('should_stop_event', threading.Event()).is_set(): break
        log_q.put(f"--- [دورة النظام الخبير {i+1}/{iterations}] | أفضل تكلفة: {format_cost_tuple(best_cost_so_far)} ---")
        
        # ✨ تعديل: التحقق من الركود الزمني
        if (time.time() - last_improvement_time) > stagnation_time_limit:
            log_q.put(f"   - ⚠️ تم كشف الركود لأكثر من {stagnation_time_limit} ثانية. سيتم زيادة الاستكشاف.")
            epsilon = min(0.9, epsilon + 0.2)
            last_improvement_time = time.time() # إعادة ضبط المؤقت

        # ... (بقية منطق تحديد الحالة واختيار الخوارزمية يبقى كما هو) ...
        failures_list = [] 
        current_state = get_state_from_failures_dominant(failures_list, current_cost[0])
        available_actions = [action for action in actions if action not in tabu_list]
        if not available_actions: available_actions = actions
        
        action_to_run = random.choice(available_actions) if random.random() < epsilon else max(available_actions, key=lambda act: q_table[current_state].get(act, 0.0))
        
        log_q.put(f"   - الحالة: '{current_state}' | القرار: تشغيل خوارزمية '{action_to_run}'...")
        tabu_list.append(action_to_run)
        
        new_solution = None
        initial_solution_for_llh = copy.deepcopy(current_solution)

        if action_to_run == "unified_lns":
            schedule_without_guards = copy.deepcopy(initial_solution_for_llh)
            for day in schedule_without_guards.values():
                for slot in day.values():
                    for exam in slot: exam['guards'] = []
            new_solution, _ = run_unified_lns_optimizer(schedule_without_guards, settings, all_professors, assignments, duty_patterns, date_map, all_subjects, log_q, all_levels_list, locked_guards)
        else:
            if action_to_run == "tabu_search":
                new_solution, _, _, _ = run_tabu_search(initial_solution_for_llh, settings, all_professors, duty_patterns, date_map, log_q, locked_guards)
            elif action_to_run == "lns":
                new_solution, _, _, _ = run_large_neighborhood_search(initial_solution_for_llh, settings, all_professors, duty_patterns, date_map, log_q, locked_guards)
            elif action_to_run == "vns":
                new_solution, _, _, _ = run_variable_neighborhood_search(initial_solution_for_llh, settings, all_professors, duty_patterns, date_map, log_q, locked_guards)
        
        if not new_solution:
            log_q.put(f"   - تحذير: خوارزمية '{action_to_run}' لم ترجع حلاً.")
            continue
            
        new_cost = calculate_cost(new_solution, settings, all_professors, duty_patterns, date_map)
        reward = calculate_reward_from_cost(current_cost, new_cost)
        
        new_failures_list = []
        new_state = get_state_from_failures_dominant(new_failures_list, new_cost[0])

        old_q_value = q_table[current_state].get(action_to_run, 0.0)
        next_max = max(q_table[new_state].values()) if q_table[new_state] else 0.0
        new_q_value = old_q_value + learning_rate * (reward + discount_factor * next_max - old_q_value)
        q_table[current_state][action_to_run] = new_q_value
        
        log_q.put(f"   - نتيجة '{action_to_run}': التكلفة {format_cost_tuple(new_cost)} | المكافأة: {reward:.1f}")
        log_q.put(f"   - تحديث الخبرة: Q('{current_state}', '{action_to_run}') -> {new_q_value:.2f}")

        current_solution = new_solution
        current_cost = new_cost
        
        if current_cost < best_cost_so_far:
            best_cost_so_far = current_cost
            best_solution_so_far = copy.deepcopy(current_solution)
            log_q.put(f"   >>> ✅ تم العثور على أفضل حل شامل جديد!")
            last_improvement_time = time.time() # ✨ تعديل: إعادة ضبط مؤقت الركود
        
        epsilon = max(0.1, epsilon * epsilon_decay)

    # ... (بقية الدالة لحفظ النتائج تبقى كما هي) ...
    log_q.put("--- انتهى عمل النظام الخبير. ---")
    try:
        q_table_to_save = {k: v for k, v in q_table.items()}
        with open(Q_TABLE_PATH, 'w', encoding='utf-8') as f:
            json.dump(q_table_to_save, f, ensure_ascii=False, indent=4)
        log_q.put("   - تم حفظ ذاكرة الخبرة للمستقبل.")
    except Exception as e:
        log_q.put(f"   - تحذير: فشل حفظ ذاكرة الخبرة. (الخطأ: {e})")

    return best_solution_so_far, True

# =====================================================================
# END: HYPER-HEURISTIC FRAMEWORK
# =====================================================================

def _run_initial_subject_placement(settings, all_subjects, all_levels_list, subject_owners, all_halls):
    """
    تقوم هذه الدالة بتشغيل المرحلة الأولى فقط: توزيع المواد في الخانات الزمنية.
    """
    base_subject_schedule = defaultdict(lambda: defaultdict(list))
    exam_schedule_settings = settings.get('examSchedule', {})
    
    # --- منطق توزيع المواد (المنسوخ من الدالة الرئيسية) ---
    all_subjects_to_schedule = {(clean_string_for_matching(s['name']), clean_string_for_matching(s['level'])) for s in all_subjects}
    group_mappings = {}
    primary_slots_by_group = defaultdict(list)
    reserve_slots = []
    sorted_dates = sorted(exam_schedule_settings.keys())
    for date in sorted_dates:
        for s in exam_schedule_settings.get(date, []):
            slot_with_date = s.copy(); slot_with_date['date'] = date
            slot_time, slot_type = s.get('time'), s.get('type')
            if slot_type == 'primary':
                primary_slots_by_group[slot_time].append(slot_with_date)
                for level in s.get('levels', []): group_mappings[level] = slot_time
            elif slot_type == 'reserve': reserve_slots.append(slot_with_date)
    
    subjects_by_group = defaultdict(set)
    for subject in all_subjects_to_schedule:
        group_id = group_mappings.get(subject[1])
        if group_id: subjects_by_group[group_id].add(subject)
    
    leftovers_by_group = defaultdict(set)
    available_halls_by_slot = defaultdict(lambda: {h['name'] for h in all_halls})
    level_hall_assignments = settings.get('levelHallAssignments', {})

    def schedule_exam_internal(subject, date, time, available_halls):
        subject_name, level_key = subject
        level_name_found = next((lvl for lvl in all_levels_list if clean_string_for_matching(lvl) == level_key), level_key)
        halls_for_level = set(level_hall_assignments.get(level_name_found, []))
        if not halls_for_level or not halls_for_level.issubset(available_halls): return False
        
        halls_details = [h for h in all_halls if h['name'] in halls_for_level]
        exam = {"date": date, "time": time, "subject": subject_name, "level": level_name_found, "professor": subject_owners.get(subject, "غير محدد"), "halls": halls_details, "guards": []}
        base_subject_schedule[date][time].append(exam)
        for hall_name in halls_for_level: available_halls.remove(hall_name)
        return True

    for group_id, subjects_pool in subjects_by_group.items():
        slots_pool = primary_slots_by_group.get(group_id, [])
        current_leftovers = set(subjects_pool)
        for subject in sorted(list(current_leftovers)):
            for slot in slots_pool:
                if subject[1] in slot.get('levels', []):
                    if schedule_exam_internal(subject, slot['date'], slot['time'], available_halls_by_slot[(slot['date'], slot['time'])]):
                        current_leftovers.remove(subject)
                        break
        leftovers_by_group[group_id] = current_leftovers

    total_leftovers = sum(len(s) for s in leftovers_by_group.values())
    if total_leftovers > 0:
        reserve_slot_claims = {}
        for group_id, subjects_left_over in sorted(leftovers_by_group.items(), key=lambda item: len(item[1]), reverse=True):
            subjects_to_remove = set()
            for subject in sorted(list(subjects_left_over)):
                for slot in reserve_slots:
                    slot_key = (slot['date'], slot['time'])
                    claimed_by = reserve_slot_claims.get(slot_key)
                    if claimed_by and claimed_by != group_id: continue
                    if schedule_exam_internal(subject, slot['date'], slot['time'], available_halls_by_slot[slot_key]):
                        subjects_to_remove.add(subject)
                        if not claimed_by: reserve_slot_claims[slot_key] = group_id
                        break
            leftovers_by_group[group_id] -= subjects_to_remove
            
    return base_subject_schedule

# ================== الواجهة الرئيسية لتشغيل الخوارزمية ==================
def _run_schedule_logic_in_background(settings, log_q, stop_event):
    try:
        stop_event.clear()
        # --- تحميل البيانات الأساسية (نفس السابق) ---
        conn = get_db_connection()
        all_levels_list = [row['name'] for row in conn.execute("SELECT name FROM levels").fetchall()]
        all_professors = [row['name'] for row in conn.execute("SELECT name FROM professors").fetchall()]
        num_professors = len(all_professors)
        all_subjects_rows = conn.execute("SELECT s.name, l.name as level FROM subjects s JOIN levels l ON s.level_id = l.id").fetchall()
        all_subjects = [dict(row) for row in all_subjects_rows]
        all_halls = [dict(row) for row in conn.execute("SELECT name, type FROM halls").fetchall()]
        assignments_rows = conn.execute('SELECT p.name as prof_name, s.name as subj_name, l.name as level_name FROM assignments a JOIN professors p ON a.professor_id = p.id JOIN subjects s ON a.subject_id = s.id JOIN levels l ON s.level_id = l.id').fetchall()
        
        # --- ✨ تعديل: قراءة الجدول المثبت من قاعدة البيانات ---
        pinned_schedule_row = conn.execute("SELECT value FROM settings WHERE key = 'pinned_subject_schedule'").fetchone()
        conn.close() # أغلق الاتصال هنا

        # --- استخلاص الإعدادات (من الكود الأصلي) ---
        intensive_search = settings.get('intensiveSearch', False)
        try:
            num_iterations = int(settings.get('iterations', '200'))
        except (ValueError, TypeError):
            num_iterations = 200
        if not intensive_search:
            num_iterations = 1
        
        best_result = {'schedule': None, 'failures': [], 'scheduling_report': [], 'unfilled_slots': float('inf'), 'unscheduled_subjects': [], 'detailed_error': None, 'prof_report': [], 'chart_data': {}, 'balance_report': {}, 'stats_dashboard': {}, 'best_cost_tuple': (float('inf'), float('inf'), float('inf'), float('inf'))}
        
        exam_schedule_settings = copy.deepcopy(settings.get('examSchedule', {}))
        
        balancing_strategy = settings.get('balancingStrategy', 'advanced')
        swap_attempts = int(settings.get('swapAttempts', 50))
        polishing_swaps = int(settings.get('polishingSwaps', 15))
        annealing_temp = float(settings.get('annealingTemp', 1000.0))
        annealing_cooling = float(settings.get('annealingCooling', 0.995))
        annealing_iterations = int(settings.get('annealingIterations', 1000))
        solver_timelimit = int(settings.get('solverTimelimit', 30))
        
        level_hall_assignments = settings.get('levelHallAssignments', {})
        duty_patterns = settings.get('dutyPatterns', {})
        assign_owner_as_guard = settings.get('assignOwnerAsGuard', False)
        max_shifts_str = settings.get('maxShifts', '0')
        max_shifts = int(max_shifts_str) if max_shifts_str != '0' else float('inf')
        last_day_restriction = settings.get('lastDayRestriction', 'none')
        unavailable_days = settings.get('unavailableDays', {})
        max_large_hall_shifts_str = settings.get('maxLargeHallShifts', '2')
        max_large_hall_shifts = int(max_large_hall_shifts_str) if max_large_hall_shifts_str != '0' else float('inf')
        enable_custom_targets = settings.get('enableCustomTargets', False)
        custom_target_patterns = settings.get('customTargetPatterns', [])
        large_hall_weight = float(settings.get('largeHallWeight', 3.0))
        other_hall_weight = float(settings.get('otherHallWeight', 1.0))
        guards_large_hall_val = int(settings.get('guardsLargeHall', 4))
        guards_medium_hall_val = int(settings.get('guardsMediumHall', 2))
        guards_small_hall_val = int(settings.get('guardsSmallHall', 1))

        sorted_dates = sorted(exam_schedule_settings.keys())
        if not sorted_dates: 
            log_q.put("خطأ: لا توجد أيام امتحانات محددة.")
            log_q.put("DONE" + json.dumps({"success": False, "message": "لا توجد أيام امتحانات محددة."}))
            return

        date_map = {date: i for i, date in enumerate(sorted_dates)}
        
        last_exam_day = sorted_dates[-1] if sorted_dates else None
        restricted_slots_on_last_day = []
        if last_exam_day and last_day_restriction != 'none':
            try:
                last_day_all_slots = sorted(exam_schedule_settings.get(last_exam_day, []), key=lambda x: x['time'])
                num_to_restrict = int(last_day_restriction.split('_')[1])
                restricted_slots_on_last_day = [s['time'] for s in last_day_all_slots[-num_to_restrict:]]
                log_q.put(f"INFO: Last day restriction '{last_day_restriction}' applied. Restricted slots: {restricted_slots_on_last_day}")
            except (ValueError, IndexError):
                log_q.put(f"WARNING: Could not parse last_day_restriction: {last_day_restriction}")


        assignments = defaultdict(list)
        for row in assignments_rows:
            assignments[row['prof_name']].append(f"{row['subj_name']} ({row['level_name']})")
        subject_owners = { (clean_string_for_matching(s['name']), clean_string_for_matching(s['level'])): clean_string_for_matching(prof) for prof, uids in assignments.items() for uid in uids for s in all_subjects if f"{s['name']} ({s['level']})" == uid }
        
        for i in range(num_iterations):
            if stop_event.is_set():
                log_q.put(f"... [محاولة {i+1}] تم اكتشاف إشارة توقف. إنهاء البحث المكثف.")
                break
            log_q.put(f">>> [Iteration {i+1}/{num_iterations}] بدء محاولة جديدة...")

            best_schedule_from_refinement = None
            best_cost_from_refinement = (float('inf'), float('inf'), float('inf'), float('inf'))
            
            # --- ✨ الخطوة 1: تحديد جدول المواد المبدئي ---
            optimized_subject_schedule = None
            if pinned_schedule_row and pinned_schedule_row['value']:
                log_q.put("--- تم العثور على جدول مواد مثبت، سيتم استخدامه كنقطة بداية. ---")
                optimized_subject_schedule = json.loads(pinned_schedule_row['value'])
                # مسح الجدول المثبت بعد استخدامه لضمان عدم استخدامه في المحاولة التالية للبحث المكثف
                # conn_temp = get_db_connection()
                # conn_temp.execute("DELETE FROM settings WHERE key = 'pinned_subject_schedule'")
                # conn_temp.commit()
                # conn_temp.close()
            else:
                log_q.put(">>> بدء المرحلة الأولى (توزيع المواد الأولي التلقائي)...")
                # نستدعي الدالة المستقلة التي أنشأناها
                optimized_subject_schedule = _run_initial_subject_placement(settings, all_subjects, all_levels_list, subject_owners, all_halls)
                log_q.put("✓ انتهى توزيع المواد الأولي.")

            # --- ✨ الخطوة 2: بدء حلقة تحسين الحراسة (refinement pass) ---
            # (هذا الجزء كان مفقوداً جزئياً في الكود الذي أرسلته وهو مهم جداً)
            ideal_guard_days = defaultdict(set)
            group_mappings = {}
            for date in sorted_dates:
                 for s in exam_schedule_settings.get(date, []):
                     if s.get('type') == 'primary':
                         for level in s.get('levels', []): 
                             group_mappings[level] = s.get('time')
            refinement_passes = int(settings.get('refinementPasses', 3))

            for refinement_pass in range(refinement_passes):
                if stop_event.is_set():
                    log_q.put(f"... [Refinement Pass {refinement_pass + 1}] تم الإيقاف قبل بدء الجولة.")
                    break
                log_q.put(f"--- بدء جولة التحسين التكراري رقم {refinement_pass + 1}/{refinement_passes} ---")
                
                # الخطوة 2: تحسين جدول المواد (يعمل على نسخته الخاصة)
                if settings.get('groupSubjects', False):
                    optimized_subject_schedule = run_subject_optimization_phase(
                        optimized_subject_schedule, assignments, all_levels_list, subject_owners, settings, log_q, group_mappings, ideal_guard_days, stop_event=stop_event
                    )
                
                # الخطوة 3: تشغيل خوارزمية الحراسة على نسخة نظيفة من جدول المواد المحسن

                # الخطوة 01: تشغيل خوارزمية الحراسة على نسخة نظيفة من جدول المواد المحسن
                schedule_for_this_pass = copy.deepcopy(optimized_subject_schedule)
                
                # الخطوة 02: تجهيز الجدول وإعطاء معرفات فريدة للامتحانات (مهم للقفل)
                all_exams_in_pass = [exam for day in schedule_for_this_pass.values() for slot in day.values() for exam in slot]
                for exam in all_exams_in_pass:
                    if 'uuid' not in exam: exam['uuid'] = str(uuid.uuid4())

                # الخطوة 03: حساب التعيينات المقفلة بشكل مركزي قبل أي استراتيجية
                locked_guards = set()
                if assign_owner_as_guard:
                    log_q.put("... تطبيق قيد تعيين أستاذ المادة (قفل)...")
                    prof_last_exam = {}
                    for exam in all_exams_in_pass:
                        owner = subject_owners.get((clean_string_for_matching(exam['subject']), clean_string_for_matching(exam['level'])))
                        if owner:
                            exam_date_time_str = f"{exam['date']} {exam['time'].split('-')[0]}"
                            if owner not in prof_last_exam or exam_date_time_str > prof_last_exam[owner]['datetime_str']:
                                prof_last_exam[owner] = {'exam': exam, 'datetime_str': exam_date_time_str}
                    
                    for owner, data in prof_last_exam.items():
                        exam_to_lock = data['exam']
                        if exam_to_lock['date'] not in unavailable_days.get(owner, []):
                            locked_guards.add((exam_to_lock['uuid'], owner))
                            log_q.put(f"    - قفل: الأستاذ '{owner}' في امتحان '{exam_to_lock['subject']}'")

                # الخطوة 04: الآن قم بتشغيل الاستراتيجية المختارة مع تمرير القيد
                temp_schedule = None
                strategy_success = False

                if balancing_strategy == 'hyper_heuristic':
                    log_q.put(">>> [جولة تحسين] تشغيل النظام الخبير (Hyper-Heuristic)...")
                    # ✅ تصحيح: النظام الخبير يحتاج جدول المواد (قبل توزيع الحراس)
                    # `schedule_for_this_pass` هو المتغير الصحيح هنا
                    temp_schedule, strategy_success = run_hyper_heuristic(
                        log_q, schedule_for_this_pass, settings, all_professors, assignments, 
                        all_levels_list, all_subjects, duty_patterns, date_map, all_halls, 
                        exam_schedule_settings, level_hall_assignments, locked_guards=locked_guards, stop_event=stop_event
                    )
                

                elif balancing_strategy == 'unified_lns':
                    log_q.put(">>> [جولة تحسين] تشغيل مُحسِّن LNS التشخيصي ...")
                    temp_schedule, strategy_success = run_unified_lns_optimizer(
                        schedule_for_this_pass, settings, all_professors, assignments,
                        duty_patterns, date_map, all_subjects, log_q, all_levels_list,
                        locked_guards=locked_guards, stop_event=stop_event
                    )
                
                elif balancing_strategy == 'genetic':
                    log_q.put(">>> [جولة تحسين] تشغيل خوارزمية الجينات...")
                    temp_schedule, strategy_success = run_genetic_algorithm(schedule_for_this_pass, settings, all_professors, assignments, all_levels_list, all_halls, exam_schedule_settings, all_subjects, level_hall_assignments, date_map, log_q, stop_event=stop_event)
                
                elif balancing_strategy == 'constraint_solver':
                    log_q.put(">>> [جولة تحسين] تشغيل البرمجة بالقيود...")
                    temp_schedule, strategy_success = run_constraint_solver(schedule_for_this_pass, settings, all_professors, assignments, all_levels_list, subject_owners, last_day_restriction, sorted_dates, duty_patterns, date_map, log_q, stop_event=stop_event)

                elif balancing_strategy in ['lns', 'vns', 'tabu_search']:
                    log_q.put(f">>> [جولة تحسين] بدء استراتيجية ({balancing_strategy.upper()})...")
                    
                    initial_solution = complete_schedule_with_guards(schedule_for_this_pass, settings, all_professors, assignments, all_levels_list, duty_patterns, date_map, all_subjects, locked_guards=locked_guards, stop_event=stop_event)
                    
                    if balancing_strategy == 'lns':
                        temp_schedule, _, _, _ = run_large_neighborhood_search(initial_solution, settings, all_professors, duty_patterns, date_map, log_q, locked_guards, stop_event=stop_event)
                    elif balancing_strategy == 'vns':
                        temp_schedule, _, _, _ = run_variable_neighborhood_search(initial_solution, settings, all_professors, duty_patterns, date_map, log_q, locked_guards, stop_event=stop_event)
                    elif balancing_strategy == 'tabu_search':
                        temp_schedule, _, _, _ = run_tabu_search(initial_solution, settings, all_professors, duty_patterns, date_map, log_q, locked_guards, stop_event=stop_event)
                    
                    strategy_success = True
                
                else: # Fallback for simple strategies (advanced, phased, etc.)
                    log_q.put(f">>> [جولة تحسين] تشغيل استراتيجية: {balancing_strategy}...")
                    
                    temp_schedule = complete_schedule_with_guards(
                        schedule_for_this_pass, settings, all_professors, assignments, 
                        all_levels_list, duty_patterns, date_map, all_subjects, locked_guards=locked_guards, stop_event=stop_event
                    )
                    
                    if balancing_strategy in ['advanced', 'phased_polished']:
                        log_q.put("... تطبيق مرحلة الصقل والتحسين...")
                        swap_count = swap_attempts if balancing_strategy == 'advanced' else polishing_swaps
                        temp_schedule, _, _, _ = run_post_processing_swaps(
                            temp_schedule, defaultdict(list), defaultdict(float), defaultdict(int), 
                            settings, all_professors, date_map, swap_count, locked_guards=locked_guards, stop_event=stop_event
                        )

                    elif balancing_strategy == 'annealing':
                        temp_schedule, _, _, _ = run_simulated_annealing(
                            temp_schedule, defaultdict(list), defaultdict(float), defaultdict(int),
                            settings, all_professors, date_map, duty_patterns, 
                            annealing_iterations, annealing_temp, annealing_cooling
                        )

                    strategy_success = True
                # =========================================================

                # الخطوة 4: التقييم والتغذية الراجعة
                if temp_schedule:
                    # أولاً، احسب تكلفة الحل الذي تم العثور عليه في هذه الجولة الحالية
                    final_cost_tuple = calculate_cost(temp_schedule, settings, all_professors, duty_patterns, date_map)
                    log_q.put(f"--- نهاية الجولة {refinement_pass + 1}: التكلفة = {format_cost_tuple(final_cost_tuple)}")

                    # الآن، قارن التكلفة الجديدة بأفضل تكلفة تم العثور عليها حتى الآن
                    if final_cost_tuple < best_cost_from_refinement:
                        best_cost_from_refinement = final_cost_tuple
                        # قم بتحديث أفضل جدول تم العثور عليه
                        best_schedule_from_refinement = copy.deepcopy(temp_schedule)
                        log_q.put("^^^ تم العثور على أفضل حل شامل حتى الآن في هذه الجولة.")

                    # تجهيز التغذية الراجعة للجولة القادمة (هذا الجزء مهم ويجب أن يبقى)
                    ideal_guard_days.clear()
                    all_exams_in_pass = [exam for day in temp_schedule.values() for slot in day.values() for exam in slot]
                    for exam in all_exams_in_pass:
                        for guard in exam.get('guards',[]):
                            if guard != "**نقص**":
                                ideal_guard_days[guard].add(exam['date'])
                else:
                    log_q.put(f"!!! فشلت خوارزمية الحراسة في إرجاع جدول في الجولة {refinement_pass + 1}")

            # في نهاية الجولة، الحل الذي سيتم تمريره للمرحلة التالية هو أفضل حل تم العثور عليه
            final_schedule_from_strategy = best_schedule_from_refinement
            # =====================================================================================
            # --- END: حلقة التحسين التكراري ---
            # =====================================================================================

            if not final_schedule_from_strategy:
                if stop_event.is_set():
                    log_q.put("... تم إيقاف العملية، تخطي تحديث أفضل نتيجة.")
                    continue
                log_q.put(f"فشلت الاستراتيجية '{balancing_strategy}' في إيجاد حل في هذه المحاولة. الانتقال للمحاولة التالية...")
                continue
            
            all_exams_in_final_schedule_flat = [exam for date_exams in final_schedule_from_strategy.values() for time_slots_in_day in date_exams.values() for exam in time_slots_in_day]
            
            for exam in all_exams_in_final_schedule_flat:
                num_guards_needed_for_logic = (sum(guards_large_hall_val for h in exam.get('halls', []) if h.get('type') == 'كبيرة') + 
                                            sum(guards_medium_hall_val for h in exam.get('halls', []) if h.get('type') == 'متوسطة') + 
                                            sum(guards_small_hall_val for h in exam.get('halls', []) if h.get('type') == 'صغيرة'))
                
                current_guards = exam.get('guards', [])
                exam['guards'] = [g for g in current_guards if g != "**نقص**"]

                while len(exam['guards']) < num_guards_needed_for_logic:
                    exam['guards'].append("**نقص**")
                
                exam['guards'] = exam['guards'][:num_guards_needed_for_logic]
            
            temp_unfilled_slots_count = 0
            temp_current_failures = []
            temp_current_scheduling_report = []

            for exam in all_exams_in_final_schedule_flat:
                if "**نقص**" in exam.get('guards', []):
                    num_guards_needed_for_logic = (sum(guards_large_hall_val for h in exam.get('halls', []) if h.get('type') == 'كبيرة') + sum(guards_medium_hall_val for h in exam.get('halls', []) if h.get('type') == 'متوسطة') + sum(guards_small_hall_val for h in exam.get('halls', []) if h.get('type') == 'صغيرة'))
                    exam['guards_incomplete'] = True
                    num_actual_shortage = exam['guards'].count("**نقص**") 
                    temp_unfilled_slots_count += num_actual_shortage
                    num_filled = num_guards_needed_for_logic - num_actual_shortage
                    temp_current_scheduling_report.append({"subject": exam['subject'], "level": exam['level'], "reason": f"نقص حراس: {num_filled}/{num_guards_needed_for_logic} (مع {num_actual_shortage} خانة نقص)"})

            current_prof_assignments_for_report = defaultdict(list)
            for exam_report in all_exams_in_final_schedule_flat:
                for guard in exam_report.get('guards', []):
                    if guard != "**نقص**":
                        current_prof_assignments_for_report[guard].append({'date': exam_report['date'], 'time': exam_report['time']})
            
            for prof_name_check, pattern in duty_patterns.items():
                if prof_name_check not in all_professors: continue
                duties_dates_indices = sorted(list({date_map.get(d_item['date']) for d_item in current_prof_assignments_for_report.get(prof_name_check, []) if date_map.get(d_item['date']) is not None}))
                num_unique_duty_days = len(duties_dates_indices)
                if num_unique_duty_days == 0:
                    continue

                if pattern == 'consecutive_strict':
                    if num_unique_duty_days != 2 or (len(duties_dates_indices) > 1 and duties_dates_indices[1] - duties_dates_indices[0] != 1):
                        temp_current_failures.append({"name": prof_name_check, "reason": "فشل في تحقيق شرط اليومين المتتاليين الإلزامي."})
                elif pattern == 'flexible_2_days':
                    if num_unique_duty_days != 2:
                        temp_current_failures.append({"name": prof_name_check, "reason": "فشل في تحقيق شرط اليومين (مرن)."})
                elif pattern == 'flexible_3_days':
                    if num_unique_duty_days not in [2, 3]:
                        temp_current_failures.append({"name": prof_name_check, "reason": "فشل في تحقيق شرط يومين أو 3 أيام (مرن)."})

            # الشرط المحدث لمقارنة أفضل الحلول (بدون المتغير المحذوف)
            if final_schedule_from_strategy:
                current_cost_tuple = calculate_cost(final_schedule_from_strategy, settings, all_professors, duty_patterns, date_map)

                # نقارن التكلفة الحالية بأفضل تكلفة وجدناها حتى الآن في كل المحاولات
                if best_result['schedule'] is None or current_cost_tuple < best_result['best_cost_tuple']:
                    log_q.put(f"✓ [Iteration {i+1}] Found a better overall solution! Cost: {format_cost_tuple(current_cost_tuple)}")

                    # تحديث أفضل نتيجة تم العثور عليها
                    best_result['schedule'] = copy.deepcopy(final_schedule_from_strategy)
                    best_result['best_cost_tuple'] = current_cost_tuple

                    # --- START: إعادة حساب كل التقارير بناءً على الحل الأفضل الجديد ---
                    all_exams_in_final_schedule_flat = [exam for date_exams in best_result['schedule'].values() for time_slots_in_day in date_exams.values() for exam in time_slots_in_day]

                    temp_unfilled_slots_count = sum(e.get('guards',[]).count("**نقص**") for e in all_exams_in_final_schedule_flat)
                    temp_current_failures = []
                    temp_current_scheduling_report = []

                    for exam in all_exams_in_final_schedule_flat:
                        if "**نقص**" in exam.get('guards', []):
                            num_guards_needed_for_logic = (sum(guards_large_hall_val for h in exam.get('halls', []) if h.get('type') == 'كبيرة') + sum(guards_medium_hall_val for h in exam.get('halls', []) if h.get('type') == 'متوسطة') + sum(guards_small_hall_val for h in exam.get('halls', []) if h.get('type') == 'صغيرة'))
                            exam['guards_incomplete'] = True
                            num_actual_shortage = exam['guards'].count("**نقص**")
                            num_filled = num_guards_needed_for_logic - num_actual_shortage
                            temp_current_scheduling_report.append({"subject": exam['subject'], "level": exam['level'], "reason": f"نقص حراس: {num_filled}/{num_guards_needed_for_logic} (مع {num_actual_shortage} خانة نقص)"})

                    current_prof_assignments_for_report = defaultdict(list)
                    for exam_report in all_exams_in_final_schedule_flat:
                        for guard in exam_report.get('guards', []):
                            if guard != "**نقص**":
                                current_prof_assignments_for_report[guard].append({'date': exam_report['date'], 'time': exam_report['time']})

                    for prof_name_check, pattern in duty_patterns.items():
                        if prof_name_check not in all_professors: continue
                        duties_dates_indices = sorted(list({date_map.get(d_item['date']) for d_item in current_prof_assignments_for_report.get(prof_name_check, []) if date_map.get(d_item['date']) is not None}))
                        num_unique_duty_days = len(duties_dates_indices)
                        if num_unique_duty_days == 0:
                            continue

                        if pattern == 'consecutive_strict':
                            if num_unique_duty_days != 2 or (len(duties_dates_indices) > 1 and duties_dates_indices[1] - duties_dates_indices[0] != 1):
                                temp_current_failures.append({"name": prof_name_check, "reason": "فشل في تحقيق شرط اليومين المتتاليين الإلزامي."})
                        elif pattern == 'flexible_2_days':
                            if num_unique_duty_days != 2:
                                temp_current_failures.append({"name": prof_name_check, "reason": "فشل في تحقيق شرط اليومين (مرن)."})
                        elif pattern == 'flexible_3_days':
                            if num_unique_duty_days not in [2, 3]:
                                temp_current_failures.append({"name": prof_name_check, "reason": "فشل في تحقيق شرط يومين أو 3 أيام (مرن)."})

                    best_result['failures'] = list(temp_current_failures)
                    best_result['scheduling_report'] = list(temp_current_scheduling_report)
                    best_result['unfilled_slots'] = temp_unfilled_slots_count
                    best_result['unscheduled_subjects'] = [] # Phase 1 ensures this is empty

                    prof_stats = {prof_name: {'large': 0, 'other': 0} for prof_name in all_professors}
                    for exam_stat in all_exams_in_final_schedule_flat:
                        guards_for_this_exam = exam_stat.get('guards', [])
                        large_guards_needed_for_exam = sum(guards_large_hall_val for h in exam_stat.get('halls', []) if h.get('type') == 'كبيرة')
                        large_hall_guards = guards_for_this_exam[:large_guards_needed_for_exam]
                        other_hall_guards = guards_for_this_exam[large_guards_needed_for_exam:]
                        for guard in large_hall_guards:
                            if guard != "**نقص**" and guard in prof_stats: prof_stats[guard]['large'] += 1
                        for guard in other_hall_guards:
                            if guard != "**نقص**" and guard in prof_stats: prof_stats[guard]['other'] += 1

                    prof_targets_map = {}
                    if num_professors > 0:
                        if enable_custom_targets and custom_target_patterns:
                            prof_targets_list = []
                            for pattern in custom_target_patterns:
                                for _ in range(pattern.get('count', 0)): prof_targets_list.append({'large': pattern.get('large', 0), 'other': pattern.get('other', 0)})
                            num_to_fill = num_professors - len(prof_targets_list)
                            if num_to_fill > 0:
                                rem_large = sum(s['large'] for s in prof_stats.values()) - sum(p['large'] for p in prof_targets_list)
                                rem_other = sum(s['other'] for s in prof_stats.values()) - sum(p['other'] for p in prof_targets_list)
                                if rem_large >= 0 and rem_other >= 0:
                                    prof_targets_list.extend(calculate_balanced_distribution(rem_large, rem_other, num_to_fill, large_hall_weight, other_hall_weight))
                            shuffled_profs = list(prof_stats.keys()); random.shuffle(shuffled_profs)
                            prof_targets_map = {prof: prof_targets_list[i] for i, prof in enumerate(shuffled_profs) if i < len(prof_targets_list)}
                        else:
                            total_large_slots_final = sum(s['large'] for s in prof_stats.values())
                            total_other_slots_final = sum(s['other'] for s in prof_stats.values())
                            prof_targets_list = calculate_balanced_distribution(total_large_slots_final, total_other_slots_final, num_professors, large_hall_weight, other_hall_weight)
                            if prof_targets_list: prof_targets_map = {prof: prof_targets_list[i % len(prof_targets_list)] for i, prof in enumerate(sorted(prof_stats.keys()))}

                    balance_report = generate_balance_report(prof_stats, prof_targets_map) if prof_targets_map else {'balance_score': 0}

                    prof_stats_report = []
                    chart_data = { 'labels': [], 'datasets': [ {'label': 'حصص القاعات الأخرى', 'data': [], 'backgroundColor': 'rgba(54, 162, 235, 0.7)'}, {'label': 'حصص القاعة الكبيرة', 'data': [], 'backgroundColor': 'rgba(255, 99, 132, 0.7)'}] }

                    sorted_prof_names_for_stats = sorted(prof_stats.keys())
                    for prof_name in sorted_prof_names_for_stats:
                        stats = prof_stats[prof_name]
                        total = stats['large'] + stats['other']
                        prof_stats_report.append(f"{prof_name}: {total} حصص (قاعة كبيرة: {stats['large']}، قاعات أخرى: {stats['other']})")
                        chart_data['labels'].append(prof_name)
                        chart_data['datasets'][0]['data'].append(stats['other'])
                        chart_data['datasets'][1]['data'].append(stats['large'])

                    stats_dashboard = {}
                    dashboard_total_large = sum(stats['large'] for stats in prof_stats.values())
                    dashboard_total_other = sum(stats['other'] for stats in prof_stats.values())

                    stats_dashboard['total_large_duties'] = dashboard_total_large
                    stats_dashboard['total_other_duties'] = dashboard_total_other
                    stats_dashboard['total_duties'] = dashboard_total_large + dashboard_total_other
                    stats_dashboard['avg_duties_per_prof'] = stats_dashboard['total_duties'] / num_professors if num_professors > 0 else 0

                    duties_per_day = defaultdict(int)
                    for exam_day_count in all_exams_in_final_schedule_flat:
                        for guard in exam_day_count.get('guards', []):
                            if guard != "**نقص**": duties_per_day[exam_day_count['date']] += 1

                    if duties_per_day:
                        busiest_day_date = max(duties_per_day, key=duties_per_day.get)
                        stats_dashboard['busiest_day'] = {'date': busiest_day_date, 'duties': duties_per_day[busiest_day_date]}
                    else:
                        stats_dashboard['busiest_day'] = {'date': 'N/A', 'duties': 0}

                    prof_workload_for_dashboard = {p: (s['large'] * large_hall_weight) + (s['other'] * other_hall_weight) for p, s in prof_stats.items()}
                    sorted_profs_by_workload = sorted(prof_workload_for_dashboard.items(), key=lambda item: item[1])
                    stats_dashboard['least_burdened_profs'] = [{'name': p[0], 'workload': round(p[1], 2)} for p in sorted_profs_by_workload[:3]]
                    stats_dashboard['most_burdened_profs'] = [{'name': p[0], 'workload': round(p[1], 2)} for p in sorted_profs_by_workload[-3:]][::-1]

                    shortage_report_for_dashboard = []
                    for report_item in temp_current_scheduling_report:
                        if "نقص" in report_item.get("reason", ""):
                            shortage_report_for_dashboard.append(f"{report_item['subject']} ({report_item['level']})")
                    stats_dashboard['shortage_reports'] = shortage_report_for_dashboard
                    stats_dashboard['unscheduled_subjects_report'] = best_result['unscheduled_subjects']

                    best_result['prof_report'] = prof_stats_report
                    best_result['chart_data'] = chart_data
                    best_result['balance_report'] = balance_report
                    best_result['stats_dashboard'] = stats_dashboard

            # if best_result.get('unfilled_slots', float('inf')) == 0 and not best_result.get('failures') and not best_result.get('unscheduled_subjects'):
            #     log_q.put(f">>> [Iteration {i+1}] تم العثور على حل مثالي. إنهاء البحث المكثف.")
            #     break 

        if not best_result.get('schedule'):
            log_q.put("فشل البرنامج في إيجاد أي حل ضمن المحاولات المحددة. الرجاء مراجعة القيود.")
            log_q.put("DONE" + json.dumps({"success": False, "message": "فشل إنشاء أي حل. حاول تخفيف القيود أو زيادة عدد محاولات البحث."}))
            return

        log_q.put("DONE" + json.dumps({ 
            "success": True, 
            "schedule": best_result['schedule'], 
            "failures": best_result.get('failures', []), 
            "scheduling_report": best_result.get('scheduling_report', []),
            "prof_report": best_result.get('prof_report', []),
            "chart_data": best_result.get('chart_data', {}),
            "balance_report": best_result.get('balance_report', {}),
            "stats_dashboard": best_result.get('stats_dashboard', {})
        }, ensure_ascii=False))

    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        log_q.put(f"خطأ فادح في معالجة الخلفية: {str(e)}")
        log_q.put(error_details)
        log_q.put("DONE" + json.dumps({"success": False, "message": f"خطأ فادح: {e}"}))

    finally:
        stop_event.clear()




# يمكن إضافتها بعد دوال الخوارزمية وقبل مسارات API

# استبدل هذه الدالة بالكامل في ملف app.py
def create_word_document_with_table(doc, title, headers, data_grid):
    """
    يضيف عنواناً وجدولاً إلى مستند وورد موجود مع دعم كامل للغة العربية (RTL).
    """
    # --- ✅ بداية: الكود الجديد لتصحيح اتجاه العنوان ---
    heading = doc.add_heading(level=2)
    heading.clear() # مسح أي محتوى افتراضي
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # إضافة خاصية اتجاه الفقرة من اليمين لليسار
    pPr = heading._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)
    
    # إضافة النص كـ "run" مع تفعيل خاصية RTL للخط
    run = heading.add_run(title)
    font = run.font
    font.rtl = True
    font.name = 'Arial'
    # --- ✅ نهاية: الكود الجديد لتصحيح اتجاه العنوان ---

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.autofit = False

    tbl_pr = table._element.xpath('w:tblPr')[0]
    bidi_visual_element = OxmlElement('w:bidiVisual')
    tbl_pr.append(bidi_visual_element)

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell_paragraph = hdr_cells[i].paragraphs[0]
        cell_paragraph.text = ""
        run = cell_paragraph.add_run(header)
        font = run.font
        font.rtl = True
        font.name = 'Arial'
        cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_paragraph.paragraph_format.rtl = True

    for row_data in data_grid:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            cell_paragraph = row_cells[i].paragraphs[0]
            cell_paragraph.text = ""
            lines = str(cell_data).split('\n')
            for idx, line in enumerate(lines):
                if idx > 0:
                    cell_paragraph.add_run().add_break()
                run = cell_paragraph.add_run(line)
                font = run.font
                font.rtl = True
                font.name = 'Arial'
            cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cell_paragraph.paragraph_format.rtl = True
            
    doc.add_page_break()

# ================== الجزء الثالث: إعداد تطبيق فلاسك والمسارات (Routes) ==================
app = Flask( __name__, template_folder=get_correct_path('templates'), static_folder=get_correct_path('static') )
app.config['JSON_AS_ASCII'] = False
app.config['JSON_SORT_KEYS'] = False

@app.route('/')
def index(): return render_template('index.html')

@app.route('/api/professors', methods=['GET'])
def get_professors():
    conn = get_db_connection()
    professors = conn.execute('SELECT name FROM professors ORDER BY name').fetchall()
    conn.close()
    return jsonify([dict(row) for row in professors])

@app.route('/api/subjects', methods=['GET'])
def get_subjects():
    conn = get_db_connection()
    subjects = conn.execute('SELECT s.name, l.name as level FROM subjects s JOIN levels l ON s.level_id = l.id ORDER BY l.name, s.name').fetchall()
    conn.close()
    return jsonify([dict(row) for row in subjects])

@app.route('/api/levels', methods=['GET'])
def get_levels():
    conn = get_db_connection()
    levels = conn.execute('SELECT name FROM levels ORDER BY name').fetchall()
    conn.close()
    return jsonify([row['name'] for row in levels])

@app.route('/api/halls', methods=['GET'])
def get_halls():
    conn = get_db_connection()
    halls = conn.execute('SELECT name, type FROM halls ORDER BY name').fetchall()
    conn.close()
    return jsonify([dict(row) for row in halls])

@app.route('/api/assignments', methods=['GET'])
def get_assignments():
    conn = get_db_connection()
    rows = conn.execute('SELECT p.name as professor_name, s.name as subject_name, l.name as level_name FROM assignments a JOIN professors p ON a.professor_id = p.id JOIN subjects s ON a.subject_id = s.id JOIN levels l ON s.level_id = l.id').fetchall()
    conn.close()
    assignments_dict = defaultdict(list)
    for row in rows:
        unique_id = f"{row['subject_name']} ({row['level_name']})"
        assignments_dict[row['professor_name']].append(unique_id)
    return jsonify(assignments_dict)

@app.route('/api/settings', methods=['GET'])
def get_settings():
    conn = get_db_connection()
    setting_row = conn.execute("SELECT value FROM settings WHERE key = 'main_settings'").fetchone()
    conn.close()
    if setting_row: return jsonify(json.loads(setting_row['value']))
    return jsonify({})

@app.route('/api/settings', methods=['POST'])
def save_settings():
    settings_data = request.get_json()
    conn = get_db_connection()
    conn.execute("INSERT OR REPLACE INTO settings (key, value) VALUES (?, ?)", ('main_settings', json.dumps(settings_data, ensure_ascii=False)))
    conn.commit()
    conn.close()
    return jsonify({"success": True, "message": "تم حفظ الإعدادات بنجاح."})

@app.route('/api/professors/bulk', methods=['POST'])
def add_professors_bulk():
    data = request.get_json()
    new_names = data.get('names', [])
    if not new_names: return jsonify({"error": "البيانات غير صالحة"}), 400
    conn = get_db_connection()
    cursor = conn.cursor()
    added_count = 0
    for name in new_names:
        try:
            cursor.execute("INSERT INTO professors (name) VALUES (?)", (name,))
            added_count += 1
        except sqlite3.IntegrityError: pass
    conn.commit()
    conn.close()
    return jsonify({"success": True, "message": f"تمت إضافة {added_count} أساتذة بنجاح."}), 201

@app.route('/api/halls/bulk', methods=['POST'])
def add_halls_bulk():
    data = request.get_json()
    new_halls = data.get('halls', [])
    if not new_halls: return jsonify({"error": "البيانات غير صالحة"}), 400
    conn = get_db_connection()
    cursor = conn.cursor()
    added_count = 0
    for hall in new_halls:
        try:
            cursor.execute("INSERT INTO halls (name, type) VALUES (?, ?)", (hall['name'], hall['type']))
            added_count += 1
        except sqlite3.IntegrityError: pass
    conn.commit()
    conn.close()
    return jsonify({"success": True, "message": f"تمت إضافة {added_count} قاعات بنجاح."}), 201

@app.route('/api/levels/bulk', methods=['POST'])
def add_levels_bulk():
    data = request.get_json()
    new_names = data.get('names', [])
    if not new_names: return jsonify({"error": "بيانات غير صالحة"}), 400
    conn = get_db_connection()
    cursor = conn.cursor()
    added_count = 0
    for name in new_names:
        try:
            cursor.execute("INSERT INTO levels (name) VALUES (?)", (name,))
            added_count += 1
        except sqlite3.IntegrityError: pass
    conn.commit()
    conn.close()
    return jsonify({"success": True, "message": f"تمت إضافة {added_count} مستويات بنجاح."}), 201

@app.route('/api/subjects/bulk', methods=['POST'])
def add_subjects_bulk():
    data = request.get_json()
    new_subject_names = data.get('subjects', [])
    level_name = data.get('level')
    if not new_subject_names or not level_name: return jsonify({"error": "البيانات غير مكتملة"}), 400
    conn = get_db_connection()
    cursor = conn.cursor()
    level_row = cursor.execute("SELECT id FROM levels WHERE name = ?", (level_name,)).fetchone()
    if not level_row:
        conn.close()
        return jsonify({"error": f"المستوى '{level_name}' غير موجود."}), 404
    level_id = level_row['id']
    added_count = 0
    for subj_name in new_subject_names:
        exists = cursor.execute("SELECT 1 FROM subjects WHERE name = ? AND level_id = ?", (subj_name, level_id)).fetchone()
        if not exists:
            cursor.execute("INSERT INTO subjects (name, level_id) VALUES (?, ?)", (subj_name, level_id))
            added_count += 1
    conn.commit()
    conn.close()
    return jsonify({"success": True, "message": f"تمت إضافة {added_count} مواد بنجاح."}), 201

@app.route('/api/assign-subjects/bulk', methods=['POST'])
def assign_subjects_bulk():
    data = request.get_json()
    prof_name = data.get('professor')
    subj_unique_ids = data.get('subjects', [])
    if not prof_name or not subj_unique_ids: return jsonify({"error": "بيانات غير كافية"}), 400
    
    conn = get_db_connection()
    cursor = conn.cursor()
    
    prof_row = cursor.execute("SELECT id FROM professors WHERE name = ?", (prof_name,)).fetchone()
    if not prof_row:
        conn.close()
        return jsonify({"error": f"خطأ: الأستاذ '{prof_name}' غير موجود في قاعدة البيانات."}), 404
    prof_id = prof_row['id']

    all_levels = [row['name'] for row in cursor.execute("SELECT name FROM levels").fetchall()]
    added_count = 0
    
    for unique_id in subj_unique_ids:
        subj_name, level_name = parse_unique_id(unique_id, all_levels)
        if not subj_name or not level_name:
            conn.close()
            return jsonify({"error": f"خطأ في تحليل اسم المادة: '{unique_id}'"}), 400
            
        subj_row = cursor.execute('SELECT s.id FROM subjects s JOIN levels l ON s.level_id = l.id WHERE s.name = ? AND l.name = ?', (subj_name, level_name)).fetchone()
        
        if subj_row:
            subj_id = subj_row['id']
            try:
                cursor.execute("INSERT INTO assignments (professor_id, subject_id) VALUES (?, ?)", (prof_id, subj_id))
                added_count += 1
            except sqlite3.IntegrityError: pass
        else:
            # ==> هذا هو التعديل المهم: إرجاع رسالة خطأ بدلاً من التجاهل <==
            conn.close()
            return jsonify({"error": f"خطأ: المادة '{subj_name}' في المستوى '{level_name}' غير موجودة. تأكد من تطابق الأسماء."}), 404

    conn.commit()
    conn.close()
    return jsonify({"success": True, "message": f"تم تخصيص {added_count} مواد للأستاذ '{prof_name}'"})

@app.route('/api/professors', methods=['DELETE'])
def delete_professor():
    data = request.get_json()
    name_to_delete = data.get('name')
    if not name_to_delete: return jsonify({"error": "اسم الأستاذ مفقود"}), 400
    conn = get_db_connection()
    conn.execute("PRAGMA foreign_keys = ON")
    conn.execute("DELETE FROM professors WHERE name = ?", (name_to_delete,))
    conn.commit()
    conn.close()
    return jsonify({"success": True, "message": "تم حذف الأستاذ."})

@app.route('/api/halls', methods=['DELETE'])
def delete_hall():
    data = request.get_json()
    name_to_delete = data.get('name')
    if not name_to_delete: return jsonify({"error": "اسم القاعة مفقود"}), 400
    conn = get_db_connection()
    conn.execute("DELETE FROM halls WHERE name = ?", (name_to_delete,))
    conn.commit()
    conn.close()
    return jsonify({"success": True, "message": "تم حذف القاعة."})

@app.route('/api/levels', methods=['DELETE'])
def delete_level():
    data = request.get_json()
    name_to_delete = data.get('name')
    if not name_to_delete: return jsonify({"error": "اسم المستوى مفقود"}), 400
    conn = get_db_connection()
    conn.execute("PRAGMA foreign_keys = ON")
    conn.execute("DELETE FROM levels WHERE name = ?", (name_to_delete,))
    conn.commit()
    conn.close()
    return jsonify({"success": True, "message": "تم حذف المستوى."})

@app.route('/api/subjects', methods=['DELETE'])
def delete_subject():
    data = request.get_json()
    name_to_delete = data.get('name')
    level_to_delete = data.get('level')
    if not name_to_delete or not level_to_delete: return jsonify({"error": "بيانات المادة غير مكتملة"}), 400
    conn = get_db_connection()
    conn.execute("PRAGMA foreign_keys = ON")
    conn.execute('DELETE FROM subjects WHERE name = ? AND level_id = (SELECT id FROM levels WHERE name = ?)', (name_to_delete, level_to_delete))
    conn.commit()
    conn.close()
    return jsonify({"success": True, "message": "تم حذف المادة."})

@app.route('/api/unassign-subject', methods=['POST'])
def unassign_subject():
    data = request.get_json()
    unique_id_to_unassign = data.get('subject')
    if not unique_id_to_unassign: return jsonify({"error": "لم يتم تحديد المادة"}), 400
    conn = get_db_connection()
    cursor = conn.cursor()
    all_levels = [row['name'] for row in cursor.execute("SELECT name FROM levels").fetchall()]
    subj_name, level_name = parse_unique_id(unique_id_to_unassign, all_levels)
    if subj_name and level_name:
        cursor.execute('DELETE FROM assignments WHERE subject_id = (SELECT s.id FROM subjects s JOIN levels l ON s.level_id = l.id WHERE s.name = ? AND l.name = ?)', (subj_name, level_name))
        conn.commit()
    conn.close()
    return jsonify({"success": True, "message": f"تم إلغاء إسناد المادة '{unique_id_to_unassign}' بنجاح."})

@app.route('/api/professors/edit', methods=['POST'])
def edit_professor():
    data = request.get_json()
    old_name, new_name = data.get('old_name'), data.get('new_name')
    if not old_name or not new_name: return jsonify({"error": "البيانات غير كافية"}), 400
    conn = get_db_connection()
    conn.execute("UPDATE professors SET name = ? WHERE name = ?", (new_name, old_name))
    conn.commit()
    conn.close()
    return jsonify({"success": True, "message": "تم تعديل اسم الأستاذ بنجاح."})

@app.route('/api/halls/edit', methods=['POST'])
def edit_hall():
    data = request.get_json()
    old_name, new_name, new_type = data.get('old_name'), data.get('new_name'), data.get('new_type')
    if not old_name or not new_name or not new_type: return jsonify({"error": "البيانات غير كافية"}), 400
    conn = get_db_connection()
    conn.execute("UPDATE halls SET name = ?, type = ? WHERE name = ?", (new_name, new_type, old_name))
    conn.commit()
    conn.close()
    return jsonify({"success": True, "message": "تم تعديل القاعة بنجاح."})

@app.route('/api/subjects/edit', methods=['POST'])
def edit_subject():
    data = request.get_json()
    old_name, new_name = data.get('old_name'), data.get('new_name')
    old_level, new_level = data.get('old_level'), data.get('new_level')
    if not all([old_name, new_name, old_level, new_level]): return jsonify({"error": "البيانات غير كافية"}), 400
    conn = get_db_connection()
    new_level_id_row = conn.execute("SELECT id FROM levels WHERE name = ?", (new_level,)).fetchone()
    if not new_level_id_row:
        conn.close()
        return jsonify({"error": "المستوى الجديد غير موجود"}), 404
    new_level_id = new_level_id_row['id']
    conn.execute("UPDATE subjects SET name = ?, level_id = ? WHERE name = ? AND level_id = (SELECT id FROM levels WHERE name = ?)", (new_name, new_level_id, old_name, old_level))
    conn.commit()
    conn.close()
    return jsonify({"success": True, "message": "تم تعديل المادة بنجاح."})

@app.route('/api/levels/edit', methods=['POST'])
def edit_level():
    data = request.get_json()
    old_name, new_name = data.get('old_name'), data.get('new_name')
    if not old_name or not new_name: return jsonify({"error": "البيانات غير كافية"}), 400
    conn = get_db_connection()
    conn.execute("UPDATE levels SET name = ? WHERE name = ?", (new_name, old_name))
    conn.commit()
    conn.close()
    return jsonify({"success": True, "message": "تم تعديل المستوى بنجاح."})

# --- Routes for Backup, Restore, Export ---
@app.route('/api/backup', methods=['GET'])
def backup_data():
    conn = get_db_connection()
    professors = [dict(row) for row in conn.execute("SELECT name FROM professors").fetchall()]
    halls = [dict(row) for row in conn.execute("SELECT name, type FROM halls").fetchall()]
    levels = [row['name'] for row in conn.execute("SELECT name FROM levels").fetchall()]
    subjects_rows = conn.execute("SELECT s.name, l.name as level FROM subjects s JOIN levels l ON s.level_id = l.id").fetchall()
    subjects = [dict(row) for row in subjects_rows]
    assignments_rows = conn.execute('SELECT p.name as prof_name, s.name as subj_name, l.name as level_name FROM assignments a JOIN professors p ON a.professor_id = p.id JOIN subjects s ON a.subject_id = s.id JOIN levels l ON s.level_id = l.id').fetchall()
    assignments = defaultdict(list)
    for row in assignments_rows:
        assignments[row['prof_name']].append(f"{row['subj_name']} ({row['level_name']})")
    settings_row = conn.execute("SELECT value FROM settings WHERE key = 'main_settings'").fetchone()
    settings = json.loads(settings_row['value']) if settings_row else {}
    conn.close()
    all_data = {'professors': professors, 'halls': halls, 'levels': levels, 'subjects': subjects, 'assignments': assignments, 'settings': settings}
    json_string = json.dumps(all_data, ensure_ascii=False, indent=4)
    buffer = io.BytesIO(json_string.encode('utf-8'))
    return send_file(buffer, as_attachment=True, download_name="backup_sqlite.json", mimetype="application/json")

@app.route('/api/restore', methods=['POST'])
def restore_data():
    backup_data = request.get_json()
    required_keys = ['professors', 'halls', 'levels', 'subjects', 'assignments', 'settings']
    if not all(key in backup_data for key in required_keys): return jsonify({"error": "ملف النسخة الاحتياطية غير صالح أو تالف."}), 400
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        for table in ['assignments', 'subjects', 'levels', 'halls', 'professors', 'settings']:
            cursor.execute(f"DELETE FROM {table}")
        cursor.executemany("INSERT INTO professors (name) VALUES (?)", [(p['name'],) for p in backup_data['professors']])
        cursor.executemany("INSERT INTO halls (name, type) VALUES (?, ?)", [(h['name'], h['type']) for h in backup_data['halls']])
        cursor.executemany("INSERT INTO levels (name) VALUES (?)", [(l,) for l in backup_data['levels']])
        level_map = {row['name']: row['id'] for row in cursor.execute("SELECT id, name FROM levels").fetchall()}
        subjects_to_insert = [(s['name'], level_map[s['level']]) for s in backup_data['subjects'] if s['level'] in level_map]
        cursor.executemany("INSERT INTO subjects (name, level_id) VALUES (?, ?)", subjects_to_insert)
        prof_map = {row['name']: row['id'] for row in cursor.execute("SELECT id, name FROM professors").fetchall()}
        subj_map = {(row['name'], row['level']): row['id'] for row in cursor.execute("SELECT s.id, s.name, l.name as level FROM subjects s JOIN levels l ON s.level_id = l.id").fetchall()}
        assignments_to_insert = []
        all_levels_list = backup_data['levels']
        for prof_name, unique_ids in backup_data['assignments'].items():
            if prof_name in prof_map:
                prof_id = prof_map[prof_name]
                for uid in unique_ids:
                    subj_name, level_name = parse_unique_id(uid, all_levels_list)
                    if (subj_name, level_name) in subj_map:
                        subj_id = subj_map[(subj_name, level_name)]
                        assignments_to_insert.append((prof_id, subj_id))
        cursor.executemany("INSERT OR IGNORE INTO assignments (professor_id, subject_id) VALUES (?, ?)", assignments_to_insert)
        cursor.execute("INSERT INTO settings (key, value) VALUES (?, ?)", ('main_settings', json.dumps(backup_data['settings'])))
        conn.commit()
        conn.close()
    except Exception as e: return jsonify({"error": f"حدث خطأ أثناء استعادة البيانات: {e}"}), 500
    return jsonify({"success": True, "message": "تم استعادة البيانات بنجاح. سيتم إعادة تحميل الصفحة."})

@app.route('/api/reset-all', methods=['POST'])
def reset_all_data():
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        for table in ['assignments', 'subjects', 'levels', 'halls', 'professors', 'settings']:
            cursor.execute(f"DELETE FROM {table}")
        conn.commit()
        conn.close()
        return jsonify({"success": True, "message": "تم مسح جميع البيانات بنجاح. سيتم إعادة تحميل الصفحة."})
    except Exception as e: return jsonify({"error": f"حدث خطأ أثناء مسح البيانات: {e}"}), 500

def build_schedule_table_for_level(ws, start_row, level, all_dates, all_times, day_names, schedule_data, styles, settings):
    header_font, default_font, level_font = styles['header_font'], styles['default_font'], styles['level_font']
    border, header_fill = styles['border'], styles['header_fill']
    guards_large, guards_medium, guards_small = settings['guards_large'], settings['guards_medium'], settings['guards_small']
    current_row = start_row
    ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=len(all_dates) + 1)
    title_cell = ws.cell(row=current_row, column=1, value=f"جدول امتحانات: {level}")
    apply_styles_to_cell(title_cell, level_font, Alignment(horizontal='center', vertical='center'), border)
    current_row += 1
    header_row_index = current_row
    headers = ["الفترة"] + [f"{day_names[datetime.strptime(d, '%Y-%m-%d').isoweekday() % 7]}\n{d}" for d in all_dates]
    for col_num, header_value in enumerate(headers, 1):
        cell = ws.cell(row=header_row_index, column=col_num, value=header_value)
        apply_styles_to_cell(cell, header_font, Alignment(horizontal='center', vertical='center', wrap_text=True), border, header_fill)
    ws.row_dimensions[header_row_index].height = 40
    current_row += 1
    for time in all_times:
        data_row_index = current_row
        max_lines_in_row = 1
        time_cell = ws.cell(row=data_row_index, column=1, value=time)
        apply_styles_to_cell(time_cell, header_font, Alignment(horizontal='center', vertical='top', wrap_text=True), border)
        for col_idx, date in enumerate(all_dates, 2):
            exam = next((e for e in schedule_data.get(date, {}).get(time, []) if e['level'] == level), None)
            content = ""
            if exam:
                content = f"{exam['subject']}\nأستاذ المادة: {exam.get('professor', 'N/A')}\n\nالحراسة:"
                halls_by_type = defaultdict(list)
                for h in exam.get('halls', []): halls_by_type[h['type']].append(h['name'])
                guards = exam.get('guards', [])[:]
                for t, n, p in [('القاعة الكبيرة', 'كبيرة', guards_large), ('القاعات المتوسطة', 'متوسطة', guards_medium), ('القاعات الصغيرة', 'صغيرة', guards_small)]:
                    if halls_by_type.get(n):
                        g_list = guards[:len(halls_by_type[n]) * p]
                        guards = guards[len(halls_by_type[n]) * p:]
                        hall_names = ", ".join(halls_by_type[n])
                        guard_text = '\n'.join(g_list) if g_list else '(لا يوجد حراس)'
                        content += f"\n{t}: {hall_names}\n{guard_text}"
            num_lines = content.count('\n') + 1
            max_lines_in_row = max(max_lines_in_row, num_lines)
            data_cell = ws.cell(row=data_row_index, column=col_idx, value=content)
            apply_styles_to_cell(data_cell, default_font, Alignment(horizontal='right', vertical='top', wrap_text=True), border)
        ws.row_dimensions[data_row_index].height = max(40, max_lines_in_row * 15)
        current_row += 1
    ws.column_dimensions['A'].width = 15
    for i in range(2, len(all_dates) + 2):
        ws.column_dimensions[get_column_letter(i)].width = 35
    return current_row

@app.route('/api/export-schedule', methods=['POST'])
def export_schedule():
    schedule_data = request.get_json()
    if not schedule_data: return jsonify({"error": "No schedule data provided"}), 400
    conn = get_db_connection()
    settings_row = conn.execute("SELECT value FROM settings WHERE key = 'main_settings'").fetchone()
    conn.close()
    settings_data = json.loads(settings_row['value']) if settings_row else {}
    settings_for_export = {
        'guards_large': int(settings_data.get('guardsLargeHall', 4)),
        'guards_medium': int(settings_data.get('guardsMediumHall', 2)),
        'guards_small': int(settings_data.get('guardsSmallHall', 1))
    }
    wb = Workbook()
    if 'Sheet' in wb.sheetnames: wb.remove(wb['Sheet'])
    styles = { 'header_font': Font(bold=True, size=12, name='Calibri'), 'default_font': Font(size=11, name='Calibri'), 'level_font': Font(bold=True, size=16, name='Calibri'), 'border': Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin')), 'header_fill': PatternFill(fill_type="solid", start_color="D9D9D9") }
    all_dates = sorted(schedule_data.keys())
    all_times = sorted({time for date_slots in schedule_data.values() for time in date_slots})
    all_levels = sorted({exam['level'] for date_slots in schedule_data.values() for time_slots in date_slots.values() for exam in time_slots})
    day_names = ["الأحد", "الاثنين", "الثلاثاء", "الأربعاء", "الخميس", "الجمعة", "السبت"]
    summary_ws = wb.create_sheet(title="الجداول المجمعة", index=0)
    summary_ws.sheet_view.rightToLeft = True
    summary_current_row = 1
    for level in all_levels:
        summary_current_row = build_schedule_table_for_level(summary_ws, summary_current_row, level, all_dates, all_times, day_names, schedule_data, styles, settings_for_export)
        summary_current_row += 2 
        individual_ws = wb.create_sheet(title=level[:31])
        individual_ws.sheet_view.rightToLeft = True
        build_schedule_table_for_level(individual_ws, 1, level, all_dates, all_times, day_names, schedule_data, styles, settings_for_export)
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name="الجداول_المجمعة.xlsx", mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

@app.route('/api/export-prof-schedules', methods=['POST'])
def export_prof_schedules():
    schedule_data = request.get_json()
    if not schedule_data: return jsonify({"error": "No schedule data provided"}), 400
    conn = get_db_connection()
    all_professors = sorted([p['name'] for p in conn.execute("SELECT name FROM professors").fetchall()])
    assignments_rows = conn.execute('SELECT p.name as prof_name, s.name as subj_name, l.name as level_name FROM assignments a JOIN professors p ON a.professor_id = p.id JOIN subjects s ON a.subject_id = s.id JOIN levels l ON s.level_id = l.id').fetchall()
    all_levels_list = [row['name'] for row in conn.execute("SELECT name FROM levels").fetchall()]
    conn.close()
    cleaned_assignments = defaultdict(set)
    for row in assignments_rows:
        cleaned_assignments[row['prof_name']].add((clean_string_for_matching(row['subj_name']), clean_string_for_matching(row['level_name'])))
    wb = Workbook()
    if 'Sheet' in wb.sheetnames: wb.remove(wb['Sheet'])
    header_font = Font(bold=True, name='Calibri', size=11)
    default_font = Font(name='Calibri', size=10)
    bold_font = Font(bold=True, name='Calibri', size=10)
    title_font = Font(bold=True, name='Calibri', size=14)
    header_fill = PatternFill(start_color="D3D3D3", fill_type="solid")
    assigned_subject_fill = PatternFill(start_color="FFEB9C", fill_type="solid")
    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
    all_dates = sorted(list(schedule_data.keys()))
    all_times = sorted(list(set(time for date_slots in schedule_data.values() for time in date_slots.keys())))
    day_names = ["الأحد", "الاثنين", "الثلاثاء", "الأربعاء", "الخميس", "الجمعة", "السبت"]
    exam_map = {}
    for date, time_slots in schedule_data.items():
        for time, exams in time_slots.items(): exam_map[(date, time)] = exams
    for prof_name in all_professors:
        ws = wb.create_sheet(title=prof_name[:31])
        ws.sheet_view.rightToLeft = True
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(all_times) + 1)
        title_cell = ws.cell(row=1, column=1, value=f"جدول الحراسة الخاص بالأستاذ: {prof_name}")
        apply_styles_to_cell(title_cell, title_font, Alignment(horizontal='center', vertical='center'), thin_border)
        ws.row_dimensions[1].height = 30
        headers = ["اليوم/التاريخ"] + all_times
        ws.append(headers)
        header_row_num = ws.max_row
        for cell in ws[header_row_num]: apply_styles_to_cell(cell, header_font, Alignment(horizontal='center', vertical='center', wrap_text=True), thin_border, header_fill)
        for date in all_dates:
            day_name = day_names[datetime.strptime(date, '%Y-%m-%d').isoweekday() % 7]
            row_num = ws.max_row + 1
            date_cell = ws.cell(row=row_num, column=1, value=f"{day_name}\n{date}")
            apply_styles_to_cell(date_cell, header_font, Alignment(horizontal='center', vertical='center', wrap_text=True), thin_border)
            for col_idx, time in enumerate(all_times, 2):
                cell = ws.cell(row=row_num, column=col_idx)
                cell_parts, is_teaching_in_slot, processed_subjects = [], False, set()
                exams_in_slot = exam_map.get((date, time), [])
                prof_assigned_subjects_set = cleaned_assignments.get(prof_name, set())
                for exam in exams_in_slot:
                    exam_tuple = (clean_string_for_matching(exam.get('subject', '')), clean_string_for_matching(exam.get('level', '')))
                    if exam_tuple in prof_assigned_subjects_set:
                        is_teaching_in_slot = True
                        is_guarding = prof_name in exam.get('guards', [])
                        hall_names = ", ".join(sorted([h.get('name', '') for h in exam.get('halls', [])]))
                        if is_guarding: cell_parts.append(f"حراسة: {exam.get('subject')} ({exam.get('level')})\nقاعة: {hall_names}")
                        else: cell_parts.append(f"{exam.get('subject')} ({exam.get('level')})\n(دون حراسة)")
                        processed_subjects.add(exam.get('subject'))
                for exam in exams_in_slot:
                    is_guarding = prof_name in exam.get('guards', [])
                    if is_guarding and exam.get('subject') not in processed_subjects:
                        hall_names = ", ".join(sorted([h.get('name', '') for h in exam.get('halls', [])]))
                        cell_parts.append(f"حراسة: {exam.get('subject')} ({exam.get('level')})\nقاعة: {hall_names}")
                cell_content, cell_fill, cell_font = "", None, default_font
                if cell_parts:
                    cell_content = f"\n{'-'*20}\n".join(cell_parts)
                    if is_teaching_in_slot: cell_fill = assigned_subject_fill
                    if "حراسة:" in cell_content: cell_font = bold_font
                cell.value = cell_content
                apply_styles_to_cell(cell, cell_font, Alignment(horizontal='center', vertical='center', wrap_text=True), thin_border, cell_fill)
        ws.column_dimensions['A'].width = 15
        for i in range(2, len(all_times) + 2): ws.column_dimensions[get_column_letter(i)].width = 30
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name="جداول_الحراسة_للأساتذة.xlsx", mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

@app.route('/api/data-template', methods=['GET'])
def export_data_template():
    """
    ينشئ ويصدر ملف Excel يحتوي على البيانات الحالية (أساتذة، قاعات، مواد، مستويات).
    إذا كانت قاعدة البيانات فارغة، سيتم تصدير قالب بالرؤوس فقط.
    --- نسخة محدثة مع أعمدة أوسع ---
    """
    try:
        conn = get_db_connection()
        # 1. استعلام لجلب كل البيانات الموجودة من قاعدة البيانات
        professors_data = pd.read_sql_query("SELECT name FROM professors", conn)
        halls_data = pd.read_sql_query("SELECT name, type FROM halls", conn)
        levels_data = pd.read_sql_query("SELECT name FROM levels", conn)
        subjects_data = pd.read_sql_query("SELECT s.name, l.name as level FROM subjects s JOIN levels l ON s.level_id = l.id", conn)
        conn.close()

        # 2. إعادة تسمية الأعمدة لتطابق رؤوس القالب
        professors_data.rename(columns={'name': 'اسم الأستاذ'}, inplace=True)
        halls_data.rename(columns={'name': 'اسم القاعة', 'type': 'نوع القاعة (صغيرة، متوسطة، كبيرة)'}, inplace=True)
        levels_data.rename(columns={'name': 'اسم المستوى الدراسي'}, inplace=True)
        subjects_data.rename(columns={'name': 'اسم المادة', 'level': 'المستوى الدراسي الخاص بها'}, inplace=True)

        # 3. كتابة البيانات إلى ملف إكسل باستخدام Pandas
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            professors_data.to_excel(writer, sheet_name='الأساتذة', index=False)
            halls_data.to_excel(writer, sheet_name='القاعات', index=False)
            levels_data.to_excel(writer, sheet_name='المستويات', index=False)
            subjects_data.to_excel(writer, sheet_name='المواد', index=False)

            # 4. تطبيق التنسيقات (اتجاه الصفحة وعرض الأعمدة)
            workbook = writer.book
            for sheet_name in workbook.sheetnames:
                worksheet = workbook[sheet_name]
                # تفعيل واجهة من اليمين لليسار
                worksheet.sheet_view.rightToLeft = True
                
                # --- ✅  توسيع عرض الأعمدة ---
                worksheet.column_dimensions['A'].width = 40
                worksheet.column_dimensions['B'].width = 35
                worksheet.column_dimensions['C'].width = 25

        output.seek(0)
        # تغيير اسم الملف المصدّر ليعكس المحتوى
        return send_file(output, as_attachment=True, download_name='بيانات_الحراسة_الحالية.xlsx', mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

    except Exception as e:
        # طباعة الخطأ في الطرفية للمساعدة في التصحيح
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"فشل إنشاء الملف: {e}"}), 500

@app.route('/api/import-data', methods=['POST'])
def import_data_from_file():
    if 'file' not in request.files: return jsonify({"error": "لم يتم العثور على ملف."}), 400
    file = request.files['file']
    if file.filename == '': return jsonify({"error": "لم يتم تحديد أي ملف."}), 400
    try:
        xls = pd.read_excel(file, sheet_name=None)
        conn = get_db_connection()
        cursor = conn.cursor()
        if 'الأساتذة' in xls:
            df_profs = xls['الأساتذة'].dropna(how='all')
            for index, row in df_profs.iterrows():
                try: cursor.execute("INSERT INTO professors (name) VALUES (?)", (str(row['اسم_الأستاذ']).strip(),))
                except sqlite3.IntegrityError: pass
        if 'القاعات' in xls:
            df_halls = xls['القاعات'].dropna(how='all')
            for index, row in df_halls.iterrows():
                try: cursor.execute("INSERT INTO halls (name, type) VALUES (?, ?)", (str(row['اسم_القاعة']).strip(), str(row['نوع_القاعة']).strip()))
                except sqlite3.IntegrityError: pass
        if 'المستويات' in xls:
            df_levels = xls['المستويات'].dropna(how='all')
            for index, row in df_levels.iterrows():
                try: cursor.execute("INSERT INTO levels (name) VALUES (?)", (str(row['اسم_المستوى_الدراسي']).strip(),))
                except sqlite3.IntegrityError: pass
        conn.commit()
        if 'المواد' in xls:
            level_map = {row['name']: row['id'] for row in cursor.execute("SELECT id, name FROM levels").fetchall()}
            df_subjects = xls['المواد'].dropna(how='all')
            for index, row in df_subjects.iterrows():
                name = str(row['اسم_المادة']).strip()
                level = str(row['المستوى_الدراسي_التابع_له']).strip()
                if name and level and level in level_map:
                    level_id = level_map[level]
                    exists = cursor.execute("SELECT 1 FROM subjects WHERE name = ? AND level_id = ?", (name, level_id)).fetchone()
                    if not exists: cursor.execute("INSERT INTO subjects (name, level_id) VALUES (?, ?)", (name, level_id))
        conn.commit()
        conn.close()
        return jsonify({"success": True, "message": "تم استيراد البيانات بنجاح. سيتم إعادة تحميل الصفحة."}), 201
    except Exception as e:
        return jsonify({"error": f"فشل تحليل الملف. تأكد من أن أسماء الأوراق والأعمدة متطابقة مع القالب. الخطأ: {e}"}), 500

@app.route('/api/generate-guard-schedule', methods=['POST'])
def generate_guard_schedule():
    settings = request.get_json()
    log_queue.put("بدء عملية إنشاء الجدول في الخلفية...")
    executor.submit(_run_schedule_logic_in_background, settings, log_queue, STOP_EVENT)
    return jsonify({"success": True, "message": "بدأت عملية إنشاء الجدول. يرجى متابعة السجل الحي."})

@app.route('/stream-logs')
def stream_logs():
    def generate():
        while True:
            message = log_queue.get()
            if "DONE" in message: 
                yield f"data: {message}\n\n"
                break
            yield f"data: {message}\n\n"
    return Response(stream_with_context(generate()), mimetype='text/event-stream')

@app.route('/shutdown', methods=['POST'])
def shutdown():
    def do_shutdown():
        time.sleep(1)
        print("Shutdown request received. Terminating server.")
        executor.shutdown(wait=False, cancel_futures=True) 
        os.kill(os.getpid(), signal.SIGINT)
    threading.Thread(target=do_shutdown).start()
    return jsonify({"success": True, "message": "Shutdown signal sent."})

# أضف هذا الكود في نهاية ملف app.py قبل قسم التشغيل


# استبدل هذه الدالة بالكامل في ملف app.py
# استبدل هذه الدالة بالكامل في ملف app.py
@app.route('/api/export/word/all-exams', methods=['POST'])
def export_exams_word():
    schedule_data = request.get_json()
    if not schedule_data: return jsonify({"error": "No schedule data provided"}), 400
    
    conn = get_db_connection()
    assignments_rows = conn.execute('SELECT s.name as subj_name, l.name as level_name, p.name as prof_name FROM assignments a JOIN subjects s ON a.subject_id = s.id JOIN levels l ON s.level_id = l.id JOIN professors p ON a.professor_id = p.id').fetchall()
    settings_row = conn.execute("SELECT value FROM settings WHERE key = 'main_settings'").fetchone()
    conn.close()
    
    settings_data = json.loads(settings_row['value']) if settings_row else {}
    guards_large = int(settings_data.get('guardsLargeHall', 4))
    guards_medium = int(settings_data.get('guardsMediumHall', 2))
    guards_small = int(settings_data.get('guardsSmallHall', 1))

    subject_owners = {(row['subj_name'], row['level_name']): row['prof_name'] for row in assignments_rows}
    
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    margin = Cm(0.5)
    section.top_margin, section.bottom_margin, section.left_margin, section.right_margin = margin, margin, margin, margin

    all_dates = sorted(schedule_data.keys())
    all_times = sorted({time for date_slots in schedule_data.values() for time in date_slots})
    all_levels = sorted({exam['level'] for slots in schedule_data.values() for exams in slots.values() for exam in exams})
    day_names = ["الأحد", "الاثنين", "الثلاثاء", "الأربعاء", "الخميس", "الجمعة", "السبت"]
    
    headers = ["الفترة"] + [f"{day_names[datetime.strptime(d, '%Y-%m-%d').isoweekday() % 7]}\n{d}" for d in all_dates]

    for level in all_levels:
        data_grid = []
        for time in all_times:
            row_data = [time]
            for date in all_dates:
                exam = next((e for e in schedule_data.get(date, {}).get(time, []) if e['level'] == level), None)
                content = ""
                if exam:
                    owner = subject_owners.get((exam['subject'], exam['level']), "غير محدد")
                    content = f"{exam['subject']}\nأستاذ المادة: {owner}\n\nالحراسة:"

                    halls_by_type = defaultdict(list)
                    for h in exam.get('halls', []): halls_by_type[h['type']].append(h['name'])
                    
                    guards_copy = [g for g in exam.get('guards', []) if g != "**نقص**"]

                    if halls_by_type.get('كبيرة'):
                        num_guards_needed = len(halls_by_type['كبيرة']) * guards_large
                        g_list = guards_copy[:num_guards_needed]
                        guards_copy = guards_copy[num_guards_needed:]
                        hall_names = ", ".join(halls_by_type['كبيرة'])
                        guard_text = '\n'.join(g_list) if g_list else '(لا يوجد)'
                        content += f"\nالقاعة الكبيرة: {hall_names}\n{guard_text}"
                    
                    other_hall_names = halls_by_type.get('متوسطة', []) + halls_by_type.get('صغيرة', [])
                    if other_hall_names:
                        guard_text = '\n'.join(guards_copy) if guards_copy else '(لا يوجد)'
                        content += f"\nالقاعات الأخرى: {', '.join(other_hall_names)}\n{guard_text}"
                
                row_data.append(content)
            data_grid.append(row_data)
        
        create_word_document_with_table(doc, f"جدول امتحانات: {level}", headers, data_grid)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name="جداول_الامتحانات.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# استبدل هذه الدالة بالكامل في ملف app.py
@app.route('/api/export/word/all-profs', methods=['POST'])
def export_profs_word():
    schedule_data = request.get_json()
    if not schedule_data: return jsonify({"error": "No schedule data provided"}), 400

    conn = get_db_connection()
    all_professors = sorted([p['name'] for p in conn.execute("SELECT name FROM professors").fetchall()])
    assignments_rows = conn.execute('SELECT p.name as prof_name, s.name as subj_name, l.name as level_name FROM assignments a JOIN professors p ON a.professor_id = p.id JOIN subjects s ON a.subject_id = s.id JOIN levels l ON s.level_id = l.id').fetchall()
    conn.close()

    prof_owned_subjects = defaultdict(set)
    for row in assignments_rows:
        prof_owned_subjects[row['prof_name']].add((row['subj_name'], row['level_name']))

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    margin = Cm(0.5)
    section.top_margin, section.bottom_margin, section.left_margin, section.right_margin = margin, margin, margin, margin

    all_dates = sorted(schedule_data.keys())
    all_times = sorted({time for date_slots in schedule_data.values() for time in date_slots})
    day_names = ["الأحد", "الاثنين", "الثلاثاء", "الأربعاء", "الخميس", "الجمعة", "السبت"]
    
    for prof_name in all_professors:
        title = f"جدول الحراسة: {prof_name}"
        headers = ["اليوم/التاريخ"] + all_times
        
        # --- بناء هيكل الجدول مع عنوان سليم ---
        heading = doc.add_heading(level=2); heading.clear()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pPr = heading._p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi'); bidi.set(qn('w:val'), '1'); pPr.append(bidi)
        run = heading.add_run(title)
        font = run.font; font.rtl = True; font.name = 'Arial'

        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        table.autofit = False
        tbl_pr = table._element.xpath('w:tblPr')[0]
        bidi_visual_element = OxmlElement('w:bidiVisual')
        tbl_pr.append(bidi_visual_element)
        
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            p = hdr_cells[i].paragraphs[0]; p.text = ""
            run = p.add_run(header)
            font = run.font; font.rtl = True; font.name = 'Arial'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.rtl = True

        has_any_duty = False
        for date in all_dates:
            row_cells = table.add_row().cells
            day_name = day_names[datetime.strptime(date, '%Y-%m-%d').isoweekday() % 7]
            
            p = row_cells[0].paragraphs[0]; p.text = ""
            run = p.add_run(f"{day_name}\n{date}"); run.font.rtl = True; run.font.name = 'Arial'; run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.rtl = True

            for i, time in enumerate(all_times, 1):
                cell_content_parts = []
                is_teaching_and_guarding = False
                is_teaching_only = False
                
                exams_in_slot = schedule_data.get(date, {}).get(time, [])
                
                for exam in exams_in_slot:
                    is_guarding = prof_name in exam.get('guards', [])
                    is_owner = (exam['subject'], exam['level']) in prof_owned_subjects.get(prof_name, set())

                    if is_guarding or is_owner:
                        has_any_duty = True
                        if is_guarding:
                            if is_owner: is_teaching_and_guarding = True
                            # ✅ التعديل هنا: استبدال القاعات بـ (حراسة)
                            cell_content_parts.append(f"{exam['subject']} ({exam['level']})\n(حراسة)")
                        elif is_owner:
                            is_teaching_only = True
                            cell_content_parts.append(f"{exam['subject']} ({exam['level']})\n(دون حراسة)")
                
                p = row_cells[i].paragraphs[0]; p.text = ""
                lines = "\n---\n".join(cell_content_parts).split('\n')
                for idx, line in enumerate(lines):
                    if idx > 0: p.add_run().add_break()
                    run = p.add_run(line)
                    font = run.font; font.rtl = True; font.name = 'Arial'
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; p.paragraph_format.rtl = True
                
                shading_elm = OxmlElement('w:shd')
                if is_teaching_and_guarding:
                    shading_elm.set(qn('w:fill'), 'D4EDDA')
                    row_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
                elif is_teaching_only:
                    shading_elm.set(qn('w:fill'), 'FFF3CD')
                    row_cells[i]._tc.get_or_add_tcPr().append(shading_elm)

        if has_any_duty:
             doc.add_page_break()
        else:
            doc._body.remove(table._element)
            doc._body.remove(heading._element)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name="جداول_الحراسة_للأساتذة.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# أضف هذه الدالة الجديدة بالكامل في نهاية ملف app.py
@app.route('/api/export/word/all-profs-anonymous', methods=['POST'])
def export_profs_anonymous_word():
    schedule_data = request.get_json()
    if not schedule_data: return jsonify({"error": "No schedule data provided"}), 400

    conn = get_db_connection()
    all_professors = sorted([p['name'] for p in conn.execute("SELECT name FROM professors").fetchall()])
    assignments_rows = conn.execute('SELECT p.name as prof_name, s.name as subj_name, l.name as level_name FROM assignments a JOIN professors p ON a.professor_id = p.id JOIN subjects s ON a.subject_id = s.id JOIN levels l ON s.level_id = l.id').fetchall()
    conn.close()

    prof_owned_subjects = defaultdict(set)
    for row in assignments_rows:
        prof_owned_subjects[row['prof_name']].add((row['subj_name'], row['level_name']))

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    margin = Cm(0.5)
    section.top_margin, section.bottom_margin, section.left_margin, section.right_margin = margin, margin, margin, margin

    all_dates = sorted(schedule_data.keys())
    all_times = sorted({time for date_slots in schedule_data.values() for time in date_slots})
    day_names = ["الأحد", "الاثنين", "الثلاثاء", "الأربعاء", "الخميس", "الجمعة", "السبت"]
    
    for prof_name in all_professors:
        title = f"جدول الحراسة (مُبسَّط): {prof_name}"
        headers = ["اليوم/التاريخ"] + all_times
        
        heading = doc.add_heading(level=2); heading.clear()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pPr = heading._p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi'); bidi.set(qn('w:val'), '1'); pPr.append(bidi)
        run = heading.add_run(title)
        font = run.font; font.rtl = True; font.name = 'Arial'

        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        table.autofit = False
        tbl_pr = table._element.xpath('w:tblPr')[0]
        bidi_visual_element = OxmlElement('w:bidiVisual')
        tbl_pr.append(bidi_visual_element)
        
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            p = hdr_cells[i].paragraphs[0]; p.text = ""
            run = p.add_run(header)
            font = run.font; font.rtl = True; font.name = 'Arial'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.rtl = True

        has_any_duty = False
        for date in all_dates:
            row_cells = table.add_row().cells
            day_name = day_names[datetime.strptime(date, '%Y-%m-%d').isoweekday() % 7]
            
            p = row_cells[0].paragraphs[0]; p.text = ""
            run = p.add_run(f"{day_name}\n{date}"); run.font.rtl = True; run.font.name = 'Arial'; run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.rtl = True

            for i, time in enumerate(all_times, 1):
                cell_content_parts = []
                is_teaching_and_guarding = False
                is_teaching_only = False
                
                exams_in_slot = schedule_data.get(date, {}).get(time, [])
                
                for exam in exams_in_slot:
                    is_guarding = prof_name in exam.get('guards', [])
                    is_owner = (exam['subject'], exam['level']) in prof_owned_subjects.get(prof_name, set())

                    if is_guarding or is_owner:
                        has_any_duty = True
                        if is_guarding:
                            if is_owner:
                                is_teaching_and_guarding = True
                                cell_content_parts.append(f"{exam['subject']} ({exam['level']})\n(حراسة)")
                            else:
                                # ✅ التعديل هنا: إذا كان حارساً وليس صاحب المادة
                                cell_content_parts.append("(تكليف بحراسة)")
                        elif is_owner:
                            is_teaching_only = True
                            cell_content_parts.append(f"{exam['subject']} ({exam['level']})\n(دون حراسة)")
                
                p = row_cells[i].paragraphs[0]; p.text = ""
                lines = "\n---\n".join(cell_content_parts).split('\n')
                for idx, line in enumerate(lines):
                    if idx > 0: p.add_run().add_break()
                    run = p.add_run(line)
                    font = run.font; font.rtl = True; font.name = 'Arial'
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; p.paragraph_format.rtl = True
                
                shading_elm = OxmlElement('w:shd')
                if is_teaching_and_guarding:
                    shading_elm.set(qn('w:fill'), 'D4EDDA')
                    row_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
                elif is_teaching_only:
                    shading_elm.set(qn('w:fill'), 'FFF3CD')
                    row_cells[i]._tc.get_or_add_tcPr().append(shading_elm)

        if has_any_duty:
             doc.add_page_break()
        else:
            doc._body.remove(table._element)
            doc._body.remove(heading._element)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name="جداول_الحراسة_المبسطة.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.route('/api/export-manual-distribution-template', methods=['POST'])
def export_manual_distribution_template():
    try:
        settings = request.get_json()
        
        conn = get_db_connection()
        all_levels_list = [row['name'] for row in conn.execute("SELECT name FROM levels").fetchall()]
        all_subjects_rows = conn.execute("SELECT s.name, l.name as level FROM subjects s JOIN levels l ON s.level_id = l.id").fetchall()
        all_subjects = [dict(row) for row in all_subjects_rows]
        all_halls = [dict(row) for row in conn.execute("SELECT name, type FROM halls").fetchall()]
        assignments_rows = conn.execute('SELECT p.name as prof_name, s.name as subj_name, l.name as level_name FROM assignments a JOIN professors p ON a.professor_id = p.id JOIN subjects s ON a.subject_id = s.id JOIN levels l ON s.level_id = l.id').fetchall()
        conn.close()
        subject_owners = { (clean_string_for_matching(s['subj_name']), clean_string_for_matching(s['level_name'])): clean_string_for_matching(s['prof_name']) for s in assignments_rows }

        print("... [تصدير يدوي] بدء التوزيع الأولي للمواد...")
        initial_schedule = _run_initial_subject_placement(settings, all_subjects, all_levels_list, subject_owners, all_halls)
        print("... [تصدير يدوي] انتهى التوزيع الأولي، جاري إنشاء ملف الإكسل...")

        exam_schedule_settings = settings.get('examSchedule', {})
        all_dates = sorted(exam_schedule_settings.keys())
        all_times = sorted(list(set(time for slots in exam_schedule_settings.values() for slot in slots for time in [slot['time']])))

        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            for level_name in sorted(all_levels_list):
                df_level = pd.DataFrame(index=all_times, columns=all_dates)
                df_level.index.name = "الفترة"
                
                for date, slots in initial_schedule.items():
                    for time, exams in slots.items():
                        for exam in exams:
                            if exam['level'] == level_name:
                                cell_content = f"{exam['subject']} ::: {exam['professor']} ::: {exam['level']}"
                                df_level.at[time, date] = cell_content
                
                unplaced_subjects = [s for s in all_subjects if s['level'] == level_name and not any(e['subject'] == s['name'] and e['level'] == s['level'] for d in initial_schedule.values() for t in d.values() for e in t)]
                if unplaced_subjects:
                    unplaced_row_name = "--- مواد غير موزعة ---"
                    df_level.loc[unplaced_row_name] = ''
                    cell_texts = [f"{s['name']} ::: {subject_owners.get((s['name'], s['level']), 'N/A')} ::: {s['level']}" for s in unplaced_subjects]
                    if all_dates:
                        df_level.at[unplaced_row_name, all_dates[0]] = "\n".join(cell_texts)

                if not df_level.empty:
                    safe_sheet_name = level_name[:31]
                    df_level.to_excel(writer, sheet_name=safe_sheet_name)
                    worksheet = writer.sheets[safe_sheet_name]
                    worksheet.sheet_view.rightToLeft = True
                    worksheet.column_dimensions['A'].width = 20
                    for i in range(2, len(all_dates) + 2):
                        worksheet.column_dimensions[get_column_letter(i)].width = 35

                    # =================== بداية الكود الجديد ===================
                    # تطبيق التفاف النص على كل الخلايا لضمان عدم فيض النص
                    wrap_alignment = Alignment(wrap_text=True, horizontal='right', vertical='top')
                    for row in worksheet.iter_rows():
                        for cell in row:
                            cell.alignment = wrap_alignment
                    # =================== نهاية الكود الجديد ===================

        output.seek(0)
        return send_file(output, as_attachment=True, download_name='مخطط_توزيع_المواد_للتعديل.xlsx', mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"فشل إنشاء الملف: {e}"}), 500

@app.route('/api/import-manual-distribution', methods=['POST'])
def import_manual_distribution():
    if 'file' not in request.files: return jsonify({"error": "لم يتم العثور على ملف."}), 400
    file = request.files['file']
    
    try:
        xls = pd.read_excel(file, sheet_name=None, index_col=0, dtype=str)
        pinned_schedule = defaultdict(lambda: defaultdict(list))
        
        conn = get_db_connection()
        all_halls = [dict(row) for row in conn.execute("SELECT name, type FROM halls").fetchall()]
        settings_row = conn.execute("SELECT value FROM settings WHERE key = 'main_settings'").fetchone()
        conn.close()
        
        settings = json.loads(settings_row['value']) if settings_row else {}
        level_hall_assignments = settings.get('levelHallAssignments', {})

        pinned_count = 0
        for sheet_name, df in xls.items(): # اسم الورقة لم نعد نستخدمه للبيانات
            for date in df.columns:
                for time in df.index:
                    cell_value = df.at[time, date]
                    # معالجة الخلايا التي قد تحتوي على عدة مواد (مثل خانة المواد غير الموزعة)
                    if pd.notna(cell_value):
                        # تقسيم الخلية إلى أسطر منفصلة
                        subjects_in_cell = cell_value.strip().split('\n')
                        for subject_line in subjects_in_cell:
                            if ':::' in subject_line:
                                try:
                                    # ✅ --- التعديل هنا: قراءة الأجزاء الثلاثة ---
                                    subject_name, professor_name, level_name = [part.strip() for part in subject_line.split(' ::: ')]

                                    # تجاهل المواد التي لم يتم تحديد يوم أو وقت لها
                                    if not date or not time or "مواد غير موزعة" in time:
                                        continue

                                    halls_for_level = set(level_hall_assignments.get(level_name, []))
                                    halls_details = [h for h in all_halls if h['name'] in halls_for_level]
                                    
                                    exam = {
                                        "date": date.strip(), "time": time.strip(),
                                        "subject": subject_name, "level": level_name,
                                        "professor": professor_name, "halls": halls_details,
                                        "guards": []
                                    }
                                    pinned_schedule[exam['date']][exam['time']].append(exam)
                                    pinned_count += 1
                                except ValueError:
                                    # تجاهل الأسطر التي لا تتبع النمط الثلاثي
                                    continue
        
        conn = get_db_connection()
        conn.execute("INSERT OR REPLACE INTO settings (key, value) VALUES (?, ?)", 
                     ('pinned_subject_schedule', json.dumps(pinned_schedule)))
        conn.commit()
        conn.close()

        return jsonify({"success": True, "message": f"تم استيراد وتثبيت {pinned_count} مادة بنجاح."})

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"فشل تحليل الملف. تأكد من أن الهيئة متطابقة مع الملف المصدر. الخطأ: {e}"}), 500

@app.route('/api/clear-manual-distribution', methods=['POST'])
def clear_manual_distribution():
    try:
        conn = get_db_connection()
        conn.execute("DELETE FROM settings WHERE key = 'pinned_subject_schedule'")
        conn.commit()
        conn.close()
        return jsonify({"success": True, "message": "تم مسح الجدول اليدوي. سيعتمد التشغيل القادم على التوزيع التلقائي."})
    except Exception as e:
        return jsonify({"error": f"حدث خطأ أثناء المسح: {e}"}), 500

@app.route('/api/stop-algorithm', methods=['POST'])
def stop_algorithm():
    log_queue.put("--- [إشارة توقف] تم استلام طلب الإيقاف من المستخدم. جاري إنهاء العملية... ---")
    STOP_EVENT.set()
    return jsonify({"success": True, "message": "تم إرسال إشارة الإيقاف بنجاح."})

# ================== الجزء الرابع: تشغيل البرنامج ==================
if __name__ == '__main__':
    init_db()
    def open_browser():
          webbrowser.open_new("http://127.0.0.1:5000")
    Timer(1, open_browser).start()
    serve(app, host='127.0.0.1', port=5000)
